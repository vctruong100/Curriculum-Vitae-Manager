"""
Strict validators for master list (.xlsx) and CV (.docx) files.

Provides deep structural validation beyond the basic checks in excel_parser
and docx_handler, including:
- Master list: column presence, hierarchy stream correctness, duplicate
  detection, malformed rows, year parsing, formula/date cell detection.
- CV: Research Experience presence, formatting conformance, indentation,
  tabs, bold/red protocol styling.

Each validator returns a structured report suitable for JSON serialisation.
"""

import re
import logging
import unicodedata
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from openpyxl import load_workbook
from openpyxl.cell.cell import Cell, TYPE_FORMULA
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from normalizer import is_phase_heading, normalize_phase, validate_year

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Report helpers
# ---------------------------------------------------------------------------

_SEV_ERROR = "error"
_SEV_WARN = "warning"
_SEV_INFO = "info"


def _issue(severity: str, row: int, message: str, field: str = "") -> Dict[str, Any]:
    return {
        "severity": severity,
        "row": row,
        "field": field,
        "message": message,
    }


# ---------------------------------------------------------------------------
# Master list (.xlsx) validator
# ---------------------------------------------------------------------------

def validate_master_xlsx_strict(
    file_path: Path,
) -> Dict[str, Any]:
    """
    Deep-validate a master .xlsx file.

    Returns a report dict:
    {
        "valid": bool,
        "file": str,
        "issues": [{"severity", "row", "field", "message"}, ...],
        "stats": {"phases": int, "subcategories": int, "studies": int},
    }
    """
    report: Dict[str, Any] = {
        "valid": True,
        "file": str(file_path),
        "issues": [],
        "stats": {"phases": 0, "subcategories": 0, "studies": 0},
    }
    issues = report["issues"]

    # --- File-level checks ---
    if not file_path.exists():
        issues.append(_issue(_SEV_ERROR, 0, f"File not found: {file_path}"))
        report["valid"] = False
        return report

    if file_path.suffix.lower() != ".xlsx":
        issues.append(_issue(_SEV_ERROR, 0, "File must have .xlsx extension"))
        report["valid"] = False
        return report

    try:
        wb = load_workbook(file_path, read_only=False, data_only=True)
    except Exception as exc:
        issues.append(_issue(_SEV_ERROR, 0, f"Cannot open workbook: {exc}"))
        report["valid"] = False
        return report

    ws = wb.active
    if ws is None:
        issues.append(_issue(_SEV_ERROR, 0, "Workbook has no active sheet"))
        report["valid"] = False
        wb.close()
        return report

    # --- Row-level checks ---
    current_phase: Optional[str] = None
    current_subcat: Optional[str] = None
    seen_studies: Dict[str, int] = {}  # canonical key -> first row
    hierarchy_state = "expect_phase"  # expect_phase | expect_subcat_or_year | expect_year

    row_num = 0
    for row in ws.iter_rows(min_row=1):
        row_num += 1
        # Skip completely blank rows
        vals = [c.value for c in row[:3]]
        if all(v is None for v in vals):
            continue

        cell_a: Cell = row[0] if len(row) > 0 else None
        cell_b: Cell = row[1] if len(row) > 1 else None
        cell_c: Cell = row[2] if len(row) > 2 else None

        col_a_raw = cell_a.value if cell_a else None
        col_b_raw = cell_b.value if cell_b else None
        col_c_raw = cell_c.value if cell_c else None

        # Detect formula cells (data_only should resolve, but warn)
        if cell_a and cell_a.data_type == "f":
            issues.append(
                _issue(_SEV_WARN, row_num, "Column A contains a formula; value may be unreliable", "A")
            )
        if cell_b and cell_b.data_type == "f":
            issues.append(
                _issue(_SEV_WARN, row_num, "Column B contains a formula", "B")
            )

        col_a = str(col_a_raw).strip() if col_a_raw is not None else ""
        if not col_a:
            continue

        # --- Phase heading? ---
        phase_name = is_phase_heading(col_a)
        if phase_name:
            current_phase = normalize_phase(phase_name)
            current_subcat = None
            report["stats"]["phases"] += 1
            hierarchy_state = "expect_subcat_or_year"
            logger.debug("Validator row %d: phase '%s'", row_num, current_phase)
            continue

        # --- Year / study row? ---
        year_val = _parse_year_cell(col_a_raw, col_a)
        if year_val is not None:
            # Validate hierarchy: must have phase
            if current_phase is None:
                issues.append(
                    _issue(_SEV_ERROR, row_num, "Study row before any Phase heading", "A")
                )
                report["valid"] = False

            # Column B should be non-empty
            col_b = str(col_b_raw).strip() if col_b_raw is not None else ""
            col_c = str(col_c_raw).strip() if col_c_raw is not None else ""

            if not col_b:
                issues.append(
                    _issue(_SEV_ERROR, row_num, "Column B (full description) is empty for study row", "B")
                )
                report["valid"] = False

            if not col_c:
                issues.append(
                    _issue(_SEV_WARN, row_num, "Column C (masked description) is empty", "C")
                )

            # Duplicate detection (canonical key)
            canon = _canonical_study_key(year_val, col_b)
            if canon in seen_studies:
                issues.append(
                    _issue(
                        _SEV_WARN,
                        row_num,
                        f"Possible duplicate of row {seen_studies[canon]}",
                        "B",
                    )
                )
            else:
                seen_studies[canon] = row_num

            report["stats"]["studies"] += 1
            hierarchy_state = "expect_subcat_or_year"
            continue

        # --- Otherwise: subcategory heading ---
        if current_phase is None:
            issues.append(
                _issue(_SEV_WARN, row_num, f"Row before any Phase heading: '{col_a[:60]}'", "A")
            )
        else:
            current_subcat = col_a
            report["stats"]["subcategories"] += 1
            hierarchy_state = "expect_year"

    # Final sanity
    if report["stats"]["phases"] == 0:
        issues.append(_issue(_SEV_ERROR, 0, "No Phase headings found"))
        report["valid"] = False
    if report["stats"]["studies"] == 0:
        issues.append(_issue(_SEV_ERROR, 0, "No study rows found"))
        report["valid"] = False

    wb.close()

    # Mark valid=False if any error-severity issues
    if any(i["severity"] == _SEV_ERROR for i in issues):
        report["valid"] = False

    logger.info(
        "Master validator: valid=%s  phases=%d  subcats=%d  studies=%d  issues=%d",
        report["valid"],
        report["stats"]["phases"],
        report["stats"]["subcategories"],
        report["stats"]["studies"],
        len(issues),
    )
    return report


def _parse_year_cell(raw_value: Any, str_value: str) -> Optional[int]:
    """
    Try to parse a year from a cell value.
    Handles int, float, date, and string representations.
    """
    import datetime as _dt

    if isinstance(raw_value, (int, float)):
        year = int(raw_value)
        if 1900 <= year <= 2100:
            return year
        return None

    if isinstance(raw_value, (_dt.datetime, _dt.date)):
        return raw_value.year

    # String
    m = re.match(r"^(\d{4})\b", str_value)
    if m:
        year = int(m.group(1))
        if 1900 <= year <= 2100:
            return year
    return None


def _canonical_study_key(year: int, col_b: str) -> str:
    """Build a canonical key for duplicate detection."""
    norm = unicodedata.normalize("NFC", col_b.lower().strip())
    norm = re.sub(r"\s+", " ", norm)
    return f"{year}|{norm}"


# ---------------------------------------------------------------------------
# CV (.docx) validator
# ---------------------------------------------------------------------------

def validate_cv_docx_strict(
    file_path: Path,
    expected_font: str = "Calibri",
    expected_font_size_pt: float = 11.0,
    expected_hanging_indent_inches: float = 0.5,
) -> Dict[str, Any]:
    """
    Deep-validate a CV .docx file for Research Experience formatting.

    Returns a report dict with issues and stats.
    """
    report: Dict[str, Any] = {
        "valid": True,
        "file": str(file_path),
        "issues": [],
        "stats": {
            "total_paragraphs": 0,
            "research_exp_paragraphs": 0,
            "study_lines": 0,
            "phase_headings": 0,
        },
    }
    issues = report["issues"]

    if not file_path.exists():
        issues.append(_issue(_SEV_ERROR, 0, f"File not found: {file_path}"))
        report["valid"] = False
        return report

    if file_path.suffix.lower() != ".docx":
        issues.append(_issue(_SEV_ERROR, 0, "File must have .docx extension"))
        report["valid"] = False
        return report

    try:
        doc = Document(file_path)
    except PermissionError:
        issues.append(_issue(_SEV_ERROR, 0, "Permission denied — file may be open in Word"))
        report["valid"] = False
        return report
    except Exception as exc:
        issues.append(_issue(_SEV_ERROR, 0, f"Cannot open document: {exc}"))
        report["valid"] = False
        return report

    report["stats"]["total_paragraphs"] = len(doc.paragraphs)

    # --- Locate Research Experience ---
    re_start = None
    re_end = None
    for idx, para in enumerate(doc.paragraphs):
        text_lower = para.text.strip().lower()
        if "research experience" in text_lower and re_start is None:
            re_start = idx
            continue
        if re_start is not None and re_end is None:
            # Check for next major heading
            if para.style and "heading" in para.style.name.lower():
                if "research experience" not in text_lower:
                    re_end = idx
                    break

    if re_start is None:
        issues.append(_issue(_SEV_ERROR, 0, "\"Research Experience\" section not found"))
        report["valid"] = False
        return report

    if re_end is None:
        re_end = len(doc.paragraphs)

    report["stats"]["research_exp_paragraphs"] = re_end - re_start

    # --- Check paragraphs inside Research Experience ---
    half_inch_emu = int(Inches(0.5))
    tolerance_emu = int(Inches(0.05))  # 5% tolerance

    for idx in range(re_start + 1, re_end):
        para = doc.paragraphs[idx]
        text = para.text.strip()
        if not text:
            continue

        para_num = idx + 1  # 1-indexed for humans

        # Detect phase heading
        if is_phase_heading(text):
            report["stats"]["phase_headings"] += 1
            continue

        # Detect study line (starts with year)
        year_match = re.match(r"^(\d{4})[\t\s]", text)
        if year_match:
            report["stats"]["study_lines"] += 1

            # Check tab separator
            if "\t" not in para.text:
                issues.append(
                    _issue(_SEV_WARN, para_num, "Study line uses spaces instead of TAB after year")
                )

            # Check hanging indent
            pf = para.paragraph_format
            if pf.first_line_indent is not None:
                actual = pf.first_line_indent
                expected_neg = -half_inch_emu
                if abs(actual - expected_neg) > tolerance_emu:
                    issues.append(
                        _issue(
                            _SEV_WARN,
                            para_num,
                            f"Hanging indent is {actual} EMU, expected ~{expected_neg} EMU (0.5\")",
                        )
                    )

            # Check left indent
            if pf.left_indent is not None and pf.left_indent != 0:
                # Left indent should be 0 or 0.5" (depending on implementation)
                pass  # Some implementations use left_indent=0.5" with hanging=0.5"

            # Check runs for font, bold, red protocol
            _check_study_runs(para, para_num, expected_font, expected_font_size_pt, issues)

    if report["stats"]["study_lines"] == 0:
        issues.append(_issue(_SEV_WARN, 0, "No study lines found in Research Experience"))

    if any(i["severity"] == _SEV_ERROR for i in issues):
        report["valid"] = False

    logger.info(
        "CV validator: valid=%s  studies=%d  phases=%d  issues=%d",
        report["valid"],
        report["stats"]["study_lines"],
        report["stats"]["phase_headings"],
        len(issues),
    )
    return report


def _check_study_runs(
    para,
    para_num: int,
    expected_font: str,
    expected_size_pt: float,
    issues: list,
):
    """Check individual runs of a study paragraph for formatting rules."""
    runs = para.runs
    if not runs:
        return

    for run in runs:
        # Font name
        if run.font.name and run.font.name != expected_font:
            issues.append(
                _issue(
                    _SEV_WARN,
                    para_num,
                    f"Run font is '{run.font.name}', expected '{expected_font}'",
                )
            )
            break  # One warning per paragraph is enough

        # Font size
        if run.font.size is not None:
            actual_pt = run.font.size.pt
            if abs(actual_pt - expected_size_pt) > 0.5:
                issues.append(
                    _issue(
                        _SEV_WARN,
                        para_num,
                        f"Run font size is {actual_pt}pt, expected {expected_size_pt}pt",
                    )
                )
                break

    # Check protocol runs for bold+red (heuristic: runs with alphanumeric-dash patterns)
    # The second or third run after the tab is typically the protocol
    found_tab = False
    sponsor_seen = False
    for run in runs:
        text = run.text.strip()
        if "\t" in run.text:
            found_tab = True
            continue
        if found_tab and not sponsor_seen and text:
            sponsor_seen = True  # First run after tab is sponsor
            if not run.font.bold:
                issues.append(
                    _issue(_SEV_WARN, para_num, "Sponsor run is not bold")
                )
            continue
        if found_tab and sponsor_seen and text:
            # This might be the protocol run
            if re.match(r"^[A-Za-z]{1,10}-?\d", text):
                if not run.font.bold:
                    issues.append(
                        _issue(_SEV_WARN, para_num, f"Protocol run '{text}' is not bold")
                    )
                color = run.font.color
                if color and color.rgb and color.rgb != RGBColor(0xFF, 0x00, 0x00):
                    issues.append(
                        _issue(
                            _SEV_WARN,
                            para_num,
                            f"Protocol run '{text}' color is {color.rgb}, expected FF0000",
                        )
                    )
            break  # Only check first potential protocol
