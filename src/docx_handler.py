"""
Word document (.docx) handler for CV Research Experience section.
Handles parsing and writing with proper formatting.
"""

import re
from pathlib import Path
from typing import List, Optional, Tuple
from datetime import datetime
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HIGHLIGHT_COLOR_YELLOW = WD_COLOR_INDEX.YELLOW

import logging as _logging

from models import Study, ResearchExperience, Phase, Subcategory, LogEntry
from normalizer import (
    normalize_for_matching,
    normalize_heading_key,
    normalize_subcat_key,
    is_uncategorized_key,
    parse_study_line,
    is_phase_heading,
    is_year_line,
    normalize_phase,
    normalize_for_display,
    strip_role_label,
    parse_sponsor_protocol,
    contains_protocol_token,
    is_already_masked,
)


class CVDocxHandler:
    """Handler for CV .docx files."""
    
    FONT_NAME = "Calibri"
    FONT_SIZE = Pt(11)
    PROTOCOL_COLOR = RGBColor(255, 0, 0)  # Red for protocol
    HANGING_INDENT = Inches(0.5)
    
    def __init__(self, file_path: Path, font_name: Optional[str] = None,
                 font_size: Optional[int] = None):
        self.file_path = file_path
        self.document = None
        self.research_exp_start_idx = None
        self.research_exp_end_idx = None
        self.signature_paragraphs = []  # Store signature section paragraphs
        self.has_signature_section = False  # Track if signature section exists
        self.year_bound = None  # Year benchmark for injection (studies AFTER this year)
        # Paragraph position tracking for preserve-existing write path
        self._phase_heading_para = {}   # normalize_heading_key(name) -> para_idx
        self._subcat_heading_para = {}  # (phase_key, subcat_key) -> para_idx
        self._subcat_last_study_para = {}  # (phase_key, subcat_key) -> para_idx
        self._subcat_study_para_list = {}  # (phase_key, subcat_key) -> [para_idx, ...]
        self._phase_last_para = {}      # phase_key -> para_idx (last para in that phase)
        # Override class-level font defaults if config values provided
        if font_name:
            self.FONT_NAME = font_name
        if font_size:
            self.FONT_SIZE = Pt(font_size)
    
    def load(self) -> None:
        """Load the document."""
        from error_handler import FilePermissionError
        try:
            self.document = Document(self.file_path)
        except PermissionError:
            raise FilePermissionError(self.file_path, "open")
    
    def find_research_experience_section(self) -> Tuple[Optional[int], Optional[int]]:
        """
        Find the Research Experience section in the document.
        
        Returns: (start_index, end_index) of paragraphs, or (None, None) if not found.
        The start_index is the heading paragraph, end_index is the last paragraph of the section.
        """
        start_idx = None
        end_idx = None
        
        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip().lower()
            
            # Find "Research Experience" heading
            if "research experience" in text and start_idx is None:
                start_idx = i
                continue
            
            # If we found the start, look for the next major section heading
            if start_idx is not None:
                # Check if this is a new major section (another heading at same level)
                # Common CV sections: Education, Publications, Skills, etc.
                if self._is_major_section_heading(para, text):
                    end_idx = i - 1
                    break
                
                # Check if this is the signature section at the end
                if self._is_signature_section(text):
                    end_idx = i - 1
                    self.has_signature_section = True  # Mark that signature exists
                    break
        
        # If we found start but no end, search backwards from end to find signature section
        if start_idx is not None and end_idx is None:
            # Default to end of document
            end_idx = len(self.document.paragraphs) - 1
            
            # Search backwards for signature section
            for i in range(len(self.document.paragraphs) - 1, start_idx, -1):
                text = self.document.paragraphs[i].text.strip()
                if self._is_signature_section(text):
                    self.has_signature_section = True
                    # Find start of signature block
                    for j in range(i, start_idx, -1):
                        sig_text = self.document.paragraphs[j].text.strip()
                        if not self._is_signature_section(sig_text) and sig_text:
                            end_idx = j
                            break
                    if end_idx == len(self.document.paragraphs) - 1:
                        end_idx = i - 1
                    break
        
        self.research_exp_start_idx = start_idx
        self.research_exp_end_idx = end_idx
        
        return start_idx, end_idx
    
    def _is_major_section_heading(self, para, text: str) -> bool:
        """Check if paragraph is a major section heading.

        Detection is *text-based first*.  A paragraph with a Word heading
        style (Heading 1, Heading 2, …) is treated as a section boundary
        **only** when its text matches a known major-section name.  Phase
        headings, subcategory headings, and other Research Experience
        content that happens to use a heading style must NOT terminate the
        Research Experience section.
        """
        major_sections = [
            'education', 'publications', 'skills', 'certifications',
            'awards', 'honors', 'professional experience', 'work experience',
            'references', 'summary', 'objective', 'teaching', 'grants',
            'presentations', 'memberships', 'affiliations', 'licenses'
        ]

        # Text-based check — authoritative
        for section in major_sections:
            if text == section or text.startswith(section + ':'):
                return True

        # Heading-styled paragraphs: only treat as boundary when the
        # text matches a known section.  Phase headings ("Phase I") and
        # subcategory headings ("Healthy Adults") styled with Word
        # heading styles are NOT section boundaries.
        if para.style and 'heading' in para.style.name.lower():
            # Already checked text against major_sections above.
            # If it didn't match, this heading belongs to RE content.
            _logging.debug(
                "[CV Parse] Heading-styled para '%s' (style='%s') kept "
                "inside Research Experience section",
                text,
                para.style.name,
            )
            return False

        return False
    
    def _is_signature_section(self, text: str) -> bool:
        """Check if text is part of a signature/declaration section."""
        text_lower = text.lower()
        
        # Common signature/declaration patterns
        signature_patterns = [
            'by signing this form',
            'i confirm that',
            'i certify that',
            'signature',
            'date of signature',
            'the information provided is accurate',
            'reflects my current qualifications',
            '___',  # Signature line
        ]
        
        for pattern in signature_patterns:
            if pattern in text_lower:
                return True
        
        return False
    
    def _infer_phase_from_context(
        self,
        research_exp: ResearchExperience,
        current_para_idx: int,
    ) -> Optional[Phase]:
        """Scan upward from *current_para_idx* to find the nearest Phase heading.

        Used when a subcategory heading is encountered without an active
        Phase context (e.g. malformed documents).  If a Phase heading is
        found, the corresponding Phase node from *research_exp* is returned
        (creating it if necessary).  Returns ``None`` if no Phase heading
        is found within the Research Experience section.
        """
        if self.research_exp_start_idx is None:
            return None
        for j in range(current_para_idx - 1, self.research_exp_start_idx, -1):
            if j >= len(self.document.paragraphs):
                continue
            scan_text = self.document.paragraphs[j].text.strip()
            if not scan_text:
                continue
            phase_name = is_phase_heading(scan_text)
            if phase_name is not None:
                _logging.info(
                    "[CV Parse] Upward scan found phase '%s' at para %d",
                    phase_name,
                    j,
                )
                return research_exp.get_or_create_phase(phase_name)
        return None

    @staticmethod
    def _merge_paragraph_text(para) -> str:
        """Merge all runs in a paragraph into a single string.

        This handles the case where Word splits a heading like
        "PHASE I" across multiple runs (e.g. bold "PHASE " and
        non-bold "I").  ``para.text`` already concatenates runs,
        but we call it explicitly here for clarity and to allow
        future pre-processing if needed.
        """
        return para.text

    def parse_research_experience(self) -> ResearchExperience:
        """
        Parse the Research Experience section into a structured format.
        Handles CVs with or without Phase/Subcategory hierarchy.
        
        Returns: ResearchExperience object with parsed studies.
        """
        if self.document is None:
            self.load()
        
        start_idx, end_idx = self.find_research_experience_section()
        
        if start_idx is None:
            raise ValueError("Research Experience section not found in document")
        
        research_exp = ResearchExperience()
        current_phase = None
        current_subcategory = None
        current_sponsor_heading = None 
        current_year_bound = None 
        
        for i in range(start_idx + 1, end_idx + 1):
            para = self.document.paragraphs[i]
            raw_text = self._merge_paragraph_text(para)
            text = raw_text.strip()
            
            if not text:
                continue
            
            # FIRST: Check if this is a no-year study line (tab-indented with sponsor:description)
            # This must come BEFORE other checks to catch Uncategorized studies
            # Skip lines that start with role labels - these are handled by sponsor heading logic
            role_prefixes_check = ['Research Assistant', 'Laboratory Technician', 'Laboratory Manager', 'Lab Manager', 
                            'Lab Technician', 'Clinical Research', 'Study Coordinator', 'Research Coordinator',
                            'Research Associate', 'Research Scientist', 'Senior Research', 'Senior Laboratory',
                            'Project Manager', 'Clinical Trial Manager']
            starts_with_role = any(text.startswith(role) for role in role_prefixes_check)
            
            # Check for colon or semicolon delimiter
            has_delimiter = ':' in text or ';' in text
            if has_delimiter and not starts_with_role and (raw_text.startswith('\t') or (current_phase and current_phase.name == "Uncategorized" and current_subcategory)):
                # Find first delimiter (colon or semicolon)
                colon_idx = text.find(':')
                semicolon_idx = text.find(';')
                if colon_idx == -1:
                    delimiter_idx = semicolon_idx
                elif semicolon_idx == -1:
                    delimiter_idx = colon_idx
                else:
                    delimiter_idx = min(colon_idx, semicolon_idx)
                
                sponsor_part = text[:delimiter_idx].strip()
                description = text[delimiter_idx + 1:].strip()
                
                # Only process if it looks like a study (has sponsor and description)
                # Sponsor part should be short (company name + protocol), not a long sentence
                if sponsor_part and description and len(description) > 20 and len(sponsor_part) < 50:
                    sponsor, protocol = parse_sponsor_protocol(sponsor_part)
                    
                    if sponsor:
                        if current_phase is None:
                            current_phase = research_exp.get_or_create_phase("Uncategorized")
                        if current_subcategory is None:
                            current_subcategory = current_phase.get_or_create_subcategory("General")
                        
                        study = Study(
                            phase=current_phase.name,
                            subcategory=current_subcategory.name,
                            year=0,
                            sponsor=sponsor,
                            protocol=protocol,
                            description_full=description,
                            description_masked=description,
                        )
                        current_subcategory.studies.append(study)
                        continue
            
            # Skip sub-section headers like "Research Experience (2022-Pres)" or "Research Experience (pre-2021)"
            # But extract year bound from "Pre YYYY" format
            # IMPORTANT: Reset phase and subcategory when encountering new Research Experience section
            if text.lower().startswith("research exp"):
                # Check for "Pre YYYY" pattern - this indicates older studies section
                year_bound_match = re.search(r'pre[- ]?(\d{4})', text, re.IGNORECASE)
                if year_bound_match:
                    current_year_bound = int(year_bound_match.group(1))
                # Reset phase and subcategory for this new section
                # Studies below this header should NOT inherit previous subcategory
                current_phase = None
                current_subcategory = None
                current_sponsor_heading = None
                continue
            
            # Check if this is a standalone year benchmark (e.g., "2022" on its own line)
            # This indicates the cutoff year for injecting studies
            # Also check raw_text to ensure it's truly standalone (not a study line with tab)
            year_match = re.match(r'^(\d{4})$', text)
            if year_match:
                potential_year = int(year_match.group(1))
                if 1900 <= potential_year <= 2099:
                    # Only treat as benchmark if there's no tab (which would indicate a study line)
                    if '\t' not in raw_text:
                        current_year_bound = potential_year
                        import logging
                        logging.info(f"[CV Parse] Year benchmark detected: {potential_year}")
                        continue
            
            # Check if this is a phase heading (Phase I, Phase II, etc.)
            phase_name = is_phase_heading(text)
            if phase_name:
                current_phase = research_exp.get_or_create_phase(phase_name)
                current_subcategory = None
                current_sponsor_heading = None
                p_key = normalize_heading_key(phase_name)
                self._phase_heading_para[p_key] = i
                self._phase_last_para[p_key] = i
                _logging.info(
                    "[CV Parse] Phase heading detected: raw='%s' "
                    "canonical='%s' (para index=%d)",
                    text,
                    phase_name,
                    i,
                )
                continue
            
            # Check if this is a study line (starts with 4-digit year)
            if is_year_line(text):
                parsed = parse_study_line(text)
                if parsed:
                    year, sponsor, protocol, description = parsed
                    
                    # Create study - will be categorized later when matched to master
                    study = Study(
                        phase=current_phase.name if current_phase else "Uncategorized",
                        subcategory=current_subcategory.name if current_subcategory else "General",
                        year=year,
                        sponsor=sponsor,
                        protocol=protocol,
                        description_full=description,
                        description_masked=description,  # Both same for now, will be updated on match
                    )
                    
                    # Add to structure
                    if current_phase is None:
                        current_phase = research_exp.get_or_create_phase("Uncategorized")
                    if current_subcategory is None:
                        current_subcategory = current_phase.get_or_create_subcategory("General")
                    
                    current_subcategory.studies.append(study)
                    p_key = normalize_heading_key(current_phase.name)
                    s_key = normalize_subcat_key(current_subcategory.name)
                    self._subcat_last_study_para[(p_key, s_key)] = i
                    self._phase_last_para[p_key] = i
                    subcat_tuple = (p_key, s_key)
                    if subcat_tuple not in self._subcat_study_para_list:
                        self._subcat_study_para_list[subcat_tuple] = []
                    self._subcat_study_para_list[subcat_tuple].append(i)
                continue
            
            # Check for sponsor headings (company names without year prefix)
            # These are sponsor names as headings, followed by study descriptions
            # Look for lines that don't start with a year and look like company names
            # Must not be a role label
            role_prefixes = ['Research Assistant', 'Laboratory Technician', 'Laboratory Manager', 'Lab Manager', 
                            'Lab Technician', 'Clinical Research', 'Study Coordinator', 'Research Coordinator',
                            'Research Associate', 'Research Scientist', 'Senior Research', 'Senior Laboratory',
                            'Project Manager', 'Clinical Trial Manager']
            is_role_label = any(text.startswith(role) for role in role_prefixes)
            
            if not text[:4].isdigit() and len(text.split()) <= 5 and not is_role_label:
                force_subcategory = (
                    current_phase is not None
                    and current_subcategory is None
                )

                next_is_year_line = False
                for peek_idx in range(i + 1, end_idx + 1):
                    peek_para = self.document.paragraphs[peek_idx]
                    peek_text = self._merge_paragraph_text(peek_para).strip()
                    if peek_text:
                        next_is_year_line = is_year_line(peek_text)
                        break

                sponsor_keywords = [
                    'INC', 'LLC', 'CORP', 'LTD', 'PHARMA', 'BIO',
                    'PLC', 'THERAPEUTICS', 'LABORATORIES', 'MEDICAL',
                ]
                has_sponsor_keyword = any(
                    keyword in text.upper()
                    for keyword in sponsor_keywords
                )
                is_likely_sponsor = (
                    not force_subcategory
                    and not next_is_year_line
                    and (
                        has_sponsor_keyword
                        or (
                            len(text.split()) <= 3
                            and (text.isupper() or text.istitle())
                        )
                    )
                )

                _logging.debug(
                    "[CV Parse] Sponsor/subcat heuristic: text='%s' "
                    "force_subcat=%s next_is_year=%s has_kw=%s "
                    "-> is_sponsor=%s (para %d)",
                    text,
                    force_subcategory,
                    next_is_year_line,
                    has_sponsor_keyword,
                    is_likely_sponsor,
                    i,
                )
                
                if is_likely_sponsor:
                    # Store the sponsor heading as-is (preserve original capitalization)
                    current_sponsor_heading = text
                    continue
                else:
                    # Might be a subcategory heading
                    if current_phase is None:
                        _logging.warning(
                            "[CV Parse] Subcategory '%s' found without "
                            "a governing Phase heading (para index=%d). "
                            "Scanning upward for nearest Phase context.",
                            text,
                            i,
                        )
                        inferred_phase = self._infer_phase_from_context(
                            research_exp, i,
                        )
                        if inferred_phase is not None:
                            current_phase = inferred_phase
                            _logging.info(
                                "[CV Parse] Adopted orphan subcategory '%s' "
                                "under inferred phase '%s'",
                                text,
                                current_phase.name,
                            )
                        else:
                            current_phase = research_exp.get_or_create_phase(
                                "Uncategorized"
                            )
                            _logging.info(
                                "[CV Parse] No phase context found; placed "
                                "subcategory '%s' under Uncategorized",
                                text,
                            )
                    subcat_display = normalize_for_display(text)
                    current_subcategory = current_phase.get_or_create_subcategory(
                        subcat_display
                    )
                    p_key = normalize_heading_key(current_phase.name)
                    s_key = normalize_subcat_key(subcat_display)
                    self._subcat_heading_para[(p_key, s_key)] = i
                    self._phase_last_para[p_key] = i
                    _logging.info(
                        "[CV Parse] Subcategory heading: raw='%s' "
                        "display='%s' under phase='%s' (para index=%d)",
                        text,
                        subcat_display,
                        current_phase.name,
                        i,
                    )
                    current_sponsor_heading = None
                    continue
            
            # Skip signature/declaration sections
            if self._is_signature_section(text):
                current_sponsor_heading = None  # Reset sponsor to prevent capturing more
                continue
            
            # Handle studies under sponsor headings: lines that start with role or are descriptions
            # These follow a sponsor heading and don't have "Sponsor:" prefix
            if current_sponsor_heading:
                # Skip if it looks like a signature/declaration
                if self._is_signature_section(text):
                    current_sponsor_heading = None
                    continue
                
                # Check if this is a study line (has role label or is long description)
                role_prefixes = ['Research Assistant', 'Laboratory Technician', 'Laboratory Manager', 'Lab Manager', 
                                'Lab Technician', 'Clinical Research', 'Study Coordinator', 'Research Coordinator',
                                'Research Associate', 'Research Scientist', 'Senior Research', 'Senior Laboratory',
                                'Project Manager', 'Clinical Trial Manager']
                has_role_label = any(text.startswith(role) for role in role_prefixes)
                is_long_description = len(text) > 50 and not text[:4].isdigit()
                
                if has_role_label or is_long_description:
                    # Extract description and strip role label
                    description = strip_role_label(text)
                    
                    # Try to extract year from description or use 0 (unknown)
                    year_match = re.search(r'\b(19|20)\d{2}\b', text)
                    year = int(year_match.group()) if year_match else 0
                    
                    study = Study(
                        phase=current_phase.name if current_phase else "Uncategorized",
                        subcategory=current_subcategory.name if current_subcategory else "General",
                        year=year,
                        sponsor=current_sponsor_heading,
                        protocol="",
                        description_full=description,
                        description_masked=description,
                    )
                    
                    if current_phase is None:
                        current_phase = research_exp.get_or_create_phase("Uncategorized")
                    if current_subcategory is None:
                        current_subcategory = current_phase.get_or_create_subcategory("General")
                    
                    current_subcategory.studies.append(study)
                    continue
        
        # Store year bound for later use
        self.year_bound = current_year_bound
        return research_exp
    
    def _create_study_paragraph(
        self,
        study: Study,
        include_protocol: bool = True,
        protocol_red: bool = True,
        highlight: bool = False,
    ):
        """
        Create a formatted paragraph for a study entry.
        
        Format: {Year}<TAB>{Sponsor}{[ SPACE ]{Protocol}}: {Description}
        - Year: not bold
        - Sponsor: bold
        - Protocol: bold + red (Mode A only)
        - Font: Calibri 11
        - Paragraph: left indent 0", hanging 0.5"
        """
        para = self.document.add_paragraph()
        
        # Set paragraph formatting
        para_format = para.paragraph_format
        para_format.left_indent = Inches(0)
        para_format.first_line_indent = Inches(-0.5)  # Negative for hanging
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(0)
        
        # Year (not bold) - show blank if year is 0
        year_str = str(study.year) if study.year > 0 else ""
        run_year = para.add_run(year_str)
        run_year.font.name = self.FONT_NAME
        run_year.font.size = self.FONT_SIZE
        run_year.font.bold = False
        self._set_font_eastasia(run_year)
        
        # Tab
        run_tab = para.add_run('\t')
        
        # Sponsor (bold) — highlighted when highlight=True
        run_sponsor = para.add_run(study.sponsor)
        run_sponsor.font.name = self.FONT_NAME
        run_sponsor.font.size = self.FONT_SIZE
        run_sponsor.font.bold = True
        if highlight:
            run_sponsor.font.highlight_color = HIGHLIGHT_COLOR_YELLOW
        self._set_font_eastasia(run_sponsor)
        
        # Protocol (bold + red, if present and included) — highlighted when highlight=True
        if include_protocol and study.protocol:
            run_space = para.add_run(' ')
            run_space.font.name = self.FONT_NAME
            run_space.font.size = self.FONT_SIZE
            
            run_protocol = para.add_run(study.protocol)
            run_protocol.font.name = self.FONT_NAME
            run_protocol.font.size = self.FONT_SIZE
            run_protocol.font.bold = True
            if protocol_red:
                run_protocol.font.color.rgb = self.PROTOCOL_COLOR
            if highlight:
                run_protocol.font.highlight_color = HIGHLIGHT_COLOR_YELLOW
            self._set_font_eastasia(run_protocol)
        
        # Colon and description
        description = study.description_full if include_protocol else study.description_masked
        run_desc = para.add_run(f': {description}')
        run_desc.font.name = self.FONT_NAME
        run_desc.font.size = self.FONT_SIZE
        run_desc.font.bold = False
        self._set_font_eastasia(run_desc)
        
        return para
    
    def _set_font_eastasia(self, run):
        """Set East Asian font to match (for proper Calibri rendering)."""
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), self.FONT_NAME)
    
    def _create_heading_paragraph(self, text: str, is_phase: bool = False):
        """Create a heading paragraph for phase or subcategory."""
        para = self.document.add_paragraph()
        
        run = para.add_run(text)
        run.font.name = self.FONT_NAME
        run.font.size = self.FONT_SIZE
        run.font.bold = True
        self._set_font_eastasia(run)
        
        return para
    
    def inject_new_studies_only(
        self,
        studies_to_inject: list,
        include_protocol: bool = True,
        protocol_red: bool = True,
        highlight_inserted: bool = False,
    ) -> int:
        """Insert new study paragraphs using per-subcategory hybrid logic.

        For subcategories that receive new studies:
          - Existing study paragraph XML elements are MOVED (preserving
            formatting, runs, etc.) into sorted order together with
            newly created study elements.  The combined list is sorted
            by year descending, then sponsor, then protocol.
        For subcategories that receive NO new studies:
          - Nothing is touched at all.
        When the target subcategory or phase does not yet exist in the
        document, the necessary heading paragraphs are created.

        Args:
            studies_to_inject: list of (Study, phase_display_name,
                subcat_display_name)
            include_protocol: include protocol in output
            protocol_red: colour protocol red

        Returns:
            Number of paragraphs inserted (new elements only).
        """
        if not studies_to_inject:
            _logging.info(
                "[DocxHandler] inject_new_studies_only: nothing to inject"
            )
            return 0

        groups = {}
        for study, phase_name, subcat_name in studies_to_inject:
            p_key = normalize_heading_key(phase_name)
            s_key = normalize_subcat_key(subcat_name)
            key = (p_key, s_key)
            if key not in groups:
                groups[key] = {
                    "phase_name": phase_name,
                    "subcat_name": subcat_name,
                    "studies": [],
                }
            groups[key]["studies"].append(study)

        body_elem = self.document.element.body
        total_inserted = 0

        # --- Phase 1: pre-compute anchor element references while
        #     paragraph indices are still valid (before any mutation). ---
        injection_plan = []
        for (p_key, s_key), group in groups.items():
            subcat_tuple = (p_key, s_key)
            phase_name = group["phase_name"]
            subcat_name = group["subcat_name"]
            new_studies = group["studies"]

            need_phase_heading = False
            need_subcat_heading = False

            existing_para_indices = self._subcat_study_para_list.get(
                subcat_tuple, []
            )

            if existing_para_indices:
                anchor_idx = existing_para_indices[-1]
            elif subcat_tuple in self._subcat_heading_para:
                anchor_idx = self._subcat_heading_para[subcat_tuple]
            elif p_key in self._phase_last_para:
                anchor_idx = self._phase_last_para[p_key]
                need_subcat_heading = True
            elif self.research_exp_end_idx is not None:
                anchor_idx = self.research_exp_end_idx
                need_phase_heading = True
                need_subcat_heading = True
            else:
                _logging.warning(
                    "[DocxHandler] Cannot determine insertion point for "
                    "phase='%s' subcat='%s'; skipping",
                    phase_name,
                    subcat_name,
                )
                continue

            # Resolve element references NOW (indices are still valid)
            anchor_elem = self.document.paragraphs[anchor_idx]._element

            # Pre-resolve existing paragraph elements and year data
            existing_elems_with_year = []
            subcat_heading_elem = None
            first_existing_prev_elem = None

            if existing_para_indices:
                for pidx in existing_para_indices:
                    para = self.document.paragraphs[pidx]
                    elem = para._element
                    year_val = 0
                    raw = self._merge_paragraph_text(para).strip()
                    year_m = re.match(r'^(\d{4})', raw)
                    if year_m:
                        year_val = int(year_m.group(1))
                    existing_elems_with_year.append(
                        (elem, year_val, "", "", True)
                    )

                sh_idx = self._subcat_heading_para.get(subcat_tuple)
                if sh_idx is not None:
                    subcat_heading_elem = (
                        self.document.paragraphs[sh_idx]._element
                    )
                else:
                    first_elem = (
                        self.document.paragraphs[
                            existing_para_indices[0]
                        ]._element
                    )
                    first_existing_prev_elem = first_elem.getprevious()
                    if first_existing_prev_elem is None:
                        first_existing_prev_elem = anchor_elem

            _logging.info(
                "[DocxHandler] inject: phase='%s' subcat='%s' "
                "existing_paras=%d new_studies=%d anchor=%d "
                "need_phase=%s need_subcat=%s",
                phase_name,
                subcat_name,
                len(existing_para_indices),
                len(new_studies),
                anchor_idx,
                need_phase_heading,
                need_subcat_heading,
            )

            injection_plan.append({
                "p_key": p_key,
                "s_key": s_key,
                "anchor_idx": anchor_idx,
                "anchor_elem": anchor_elem,
                "need_phase": need_phase_heading,
                "need_subcat": need_subcat_heading,
                "existing_para_indices": existing_para_indices,
                "existing_elems_with_year": existing_elems_with_year,
                "subcat_heading_elem": subcat_heading_elem,
                "first_existing_prev_elem": first_existing_prev_elem,
                "group": group,
            })

        # Sort by anchor_idx descending so higher-index insertions
        # happen first; this prevents earlier insertions from shifting
        # the paragraph positions used by later groups.
        injection_plan.sort(key=lambda x: -x["anchor_idx"])

        # --- Phase 2: process each group using pre-computed
        #     element references (no index lookups). ---
        for plan in injection_plan:
            p_key = plan["p_key"]
            s_key = plan["s_key"]
            group = plan["group"]
            phase_name = group["phase_name"]
            subcat_name = group["subcat_name"]
            new_studies = group["studies"]
            anchor_elem = plan["anchor_elem"]

            if plan["existing_para_indices"]:
                existing_elems_with_year = plan["existing_elems_with_year"]

                for elem, _, _, _, _ in existing_elems_with_year:
                    parent = elem.getparent()
                    if parent is not None:
                        parent.remove(elem)

                new_elems_with_year = []
                for study in new_studies:
                    use_red = (
                        protocol_red
                        and not is_uncategorized_key(p_key)
                    )
                    elem = self._create_study_element(
                        study, include_protocol, use_red,
                        highlight=highlight_inserted,
                    )
                    new_elems_with_year.append(
                        (
                            elem,
                            study.year,
                            study.sponsor.lower(),
                            study.protocol.lower(),
                            False,
                        )
                    )
                    total_inserted += 1

                combined = existing_elems_with_year + new_elems_with_year
                combined.sort(
                    key=lambda t: (-t[1], t[2], t[3])
                )

                if plan["subcat_heading_elem"] is not None:
                    insert_after_elem = plan["subcat_heading_elem"]
                elif plan["first_existing_prev_elem"] is not None:
                    insert_after_elem = plan["first_existing_prev_elem"]
                else:
                    insert_after_elem = anchor_elem

                cursor = insert_after_elem
                for elem, _, _, _, _ in combined:
                    cursor.addnext(elem)
                    cursor = elem

            else:
                # Strip bottom border from anchor so it does not
                # visually separate existing content from the newly
                # inserted headings/studies.
                self._strip_paragraph_bottom_border(anchor_elem)

                cursor = anchor_elem

                if plan["need_phase"]:
                    phase_elem = self._create_paragraph_element(
                        phase_name, is_heading=True
                    )
                    cursor.addnext(phase_elem)
                    cursor = phase_elem
                    total_inserted += 1

                if plan["need_subcat"]:
                    subcat_elem = self._create_paragraph_element(
                        subcat_name, is_heading=True
                    )
                    cursor.addnext(subcat_elem)
                    cursor = subcat_elem
                    total_inserted += 1

                new_studies.sort(
                    key=lambda s: (
                        -s.year,
                        s.sponsor.lower(),
                        s.protocol.lower(),
                    )
                )
                for study in new_studies:
                    use_red = (
                        protocol_red
                        and not is_uncategorized_key(p_key)
                    )
                    elem = self._create_study_element(
                        study, include_protocol, use_red,
                        highlight=highlight_inserted,
                    )
                    cursor.addnext(elem)
                    cursor = elem
                    total_inserted += 1

        _logging.info(
            "[DocxHandler] inject_new_studies_only: inserted %d paragraphs",
            total_inserted,
        )
        return total_inserted

    def _strip_paragraph_bottom_border(self, para_element) -> None:
        """Remove any bottom border from a paragraph element's properties.

        Called before inserting new headings/studies after an anchor so
        that the existing section-end border does not visually separate
        old content from newly injected content.
        """
        pPr = para_element.find(qn('w:pPr'))
        if pPr is None:
            return
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            return
        bottom = pBdr.find(qn('w:bottom'))
        if bottom is not None:
            pBdr.remove(bottom)
            _logging.debug(
                "[DocxHandler] Stripped bottom border from anchor paragraph"
            )
        if len(pBdr) == 0:
            pPr.remove(pBdr)

    def _replace_paragraph_with_masked(
        self,
        para,
        masked_sponsor: str,
        masked_description: str,
        year: int,
    ) -> None:
        """Replace all runs of a study paragraph with masked content.

        Preserves the paragraph element's position in the XML tree
        (body or table cell) and its paragraph-level formatting
        (indentation, spacing). Only the run-level content is replaced.
        """
        elem = para._element
        for child in list(elem):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'r':
                elem.remove(child)

        def _add_run(text, bold=False):
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), self.FONT_NAME)
            rFonts.set(qn('w:hAnsi'), self.FONT_NAME)
            rFonts.set(qn('w:eastAsia'), self.FONT_NAME)
            rPr.append(rFonts)
            _hp = str(int(self.FONT_SIZE.pt * 2))
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), _hp)
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), _hp)
            rPr.append(szCs)
            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = text
            if text.startswith(' ') or text.endswith(' '):
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t)
            elem.append(r)

        year_str = str(year) if year > 0 else ""
        _add_run(year_str, bold=False)

        r_tab = OxmlElement('w:r')
        tab_elem = OxmlElement('w:tab')
        r_tab.append(tab_elem)
        elem.append(r_tab)

        _add_run(masked_sponsor, bold=True)

        desc_clean = masked_description
        sponsor_lower = masked_sponsor.lower().strip()
        desc_lower = desc_clean.lower().strip()
        if sponsor_lower and desc_lower.startswith(sponsor_lower):
            rest = desc_clean[len(masked_sponsor):].strip()
            if rest.startswith(':'):
                desc_clean = rest[1:].strip()

        _add_run(f': {desc_clean}', bold=False)

    def redact_studies_in_place(
        self,
        replacements: list,
    ) -> int:
        """Replace protocol-bearing study paragraphs in place.

        Each item in *replacements* is a dict::

            {
                "para_idx": int,
                "year": int,
                "masked_sponsor": str,
                "masked_description": str,
            }

        The paragraph at ``para_idx`` has its runs replaced with the
        masked content while preserving its XML position (body or
        table cell), paragraph-level formatting, and surrounding
        structure. No paragraphs are added or removed.

        Returns the number of paragraphs replaced.
        """
        count = 0
        for rep in replacements:
            para_idx = rep["para_idx"]
            if para_idx < 0 or para_idx >= len(self.document.paragraphs):
                _logging.warning(
                    "[DocxHandler] redact_in_place: invalid para_idx=%d",
                    para_idx,
                )
                continue
            para = self.document.paragraphs[para_idx]
            self._replace_paragraph_with_masked(
                para,
                rep["masked_sponsor"],
                rep["masked_description"],
                rep["year"],
            )
            count += 1
            _logging.info(
                "[DocxHandler] Replaced para %d with masked text "
                "(sponsor='%s', year=%d)",
                para_idx,
                rep["masked_sponsor"],
                rep["year"],
            )
        return count

    def sort_subcategory_in_place(
        self,
        phase_key: str,
        subcat_key: str,
    ) -> None:
        """Re-sort study paragraphs within one subcategory in place.

        Removes existing study paragraph XML elements from the
        subcategory, sorts them by (year desc, sponsor, protocol),
        and reinserts them in order after the subcategory heading.
        """
        subcat_tuple = (phase_key, subcat_key)
        para_indices = self._subcat_study_para_list.get(subcat_tuple, [])
        if len(para_indices) <= 1:
            return

        elems_with_sort_key = []
        for pidx in para_indices:
            para = self.document.paragraphs[pidx]
            elem = para._element
            raw = self._merge_paragraph_text(para).strip()
            year_val = 0
            sponsor_val = ""
            protocol_val = ""
            year_m = re.match(r'^(\d{4})', raw)
            if year_m:
                year_val = int(year_m.group(1))
            parsed = parse_study_line(raw)
            if parsed:
                _, sponsor_val, protocol_val, _ = parsed
            elems_with_sort_key.append(
                (elem, year_val, sponsor_val.lower(), protocol_val.lower())
            )

        for elem, _, _, _ in elems_with_sort_key:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)

        elems_with_sort_key.sort(
            key=lambda t: (-t[1], t[2], t[3])
        )

        heading_idx = self._subcat_heading_para.get(subcat_tuple)
        if heading_idx is not None:
            cursor = self.document.paragraphs[heading_idx]._element
        else:
            cursor = self.document.paragraphs[
                para_indices[0]
            ]._element.getprevious()
            if cursor is None:
                _logging.warning(
                    "[DocxHandler] sort_subcategory_in_place: no anchor "
                    "for (%s, %s)",
                    phase_key,
                    subcat_key,
                )
                return

        for elem, _, _, _ in elems_with_sort_key:
            cursor.addnext(elem)
            cursor = elem

        _logging.info(
            "[DocxHandler] Sorted %d studies in (%s, %s)",
            len(elems_with_sort_key),
            phase_key,
            subcat_key,
        )

    def write_research_experience(
        self,
        research_exp: ResearchExperience,
        include_protocol: bool = True,
        protocol_red: bool = True,
        highlight_new: bool = False,
        new_study_ids: set = None,
    ) -> None:
        """
        Write the Research Experience section to the document.
        Replaces existing content between start and end indices.

        Args:
            highlight_new: When True, newly inserted studies are highlighted yellow.
            new_study_ids: Set of python object id()s for studies that existed in
                the original CV.  Studies NOT in this set are considered new and
                will be highlighted when highlight_new is True.  If None and
                highlight_new is True, ALL studies are highlighted.
        """
        if self.research_exp_start_idx is None:
            raise ValueError("Research Experience section not found. Call find_research_experience_section first.")
        
        # Note: Do NOT sort here - sorting should be done by caller (processor)
        # to allow custom category ordering
        
        # Get the heading element to insert after
        heading_element = self.document.paragraphs[self.research_exp_start_idx]._element
        
        # Collect elements to remove (between heading and next section)
        elements_to_remove = []
        for i in range(self.research_exp_start_idx + 1, self.research_exp_end_idx + 1):
            if i < len(self.document.paragraphs):
                elements_to_remove.append(self.document.paragraphs[i]._element)
        
        # Remove the old content
        for elem in elements_to_remove:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
        
        # Build and insert new paragraphs after the heading
        insert_after = heading_element
        
        for phase in research_exp.phases:
            # Phase heading
            phase_para = self._create_paragraph_element(phase.name, is_heading=True)
            insert_after.addnext(phase_para)
            insert_after = phase_para
            
            for subcategory in phase.subcategories:
                # Subcategory heading
                subcat_para = self._create_paragraph_element(subcategory.name, is_heading=True)
                insert_after.addnext(subcat_para)
                insert_after = subcat_para
                
                # Studies
                for study in subcategory.studies:
                    # Disable red protocol coloring for Uncategorized phase
                    use_red = protocol_red and not is_uncategorized_key(phase.name)
                    # Determine if this study should be highlighted
                    should_highlight = False
                    if highlight_new:
                        if new_study_ids is not None:
                            should_highlight = id(study) not in new_study_ids
                        else:
                            should_highlight = True
                    study_para = self._create_study_element(
                        study, include_protocol, use_red,
                        highlight=should_highlight,
                    )
                    insert_after.addnext(study_para)
                    insert_after = study_para
        
        # Add horizontal line before signature section if it exists
        self._insert_horizontal_line(insert_after)
    
    def _insert_horizontal_line(self, insert_after):
        """Insert a horizontal line after the last study, before the signature section."""
        if not self.has_signature_section:
            return
        
        # Check if a horizontal line already exists (avoid duplicates)
        # Look at the next few paragraphs after insert_after
        next_elem = insert_after.getnext()
        for _ in range(3):  # Check next 3 paragraphs
            if next_elem is not None:
                # Get text from the paragraph
                text_elems = next_elem.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                text = ''.join([t.text or '' for t in text_elems])
                # If we find a line of underscores, don't add another
                if text.strip() and text.strip().replace('_', '') == '':
                    return
                next_elem = next_elem.getnext()
            else:
                break
        
        # Add empty paragraph for spacing
        empty_para = self._create_paragraph_element("", is_heading=False)
        insert_after.addnext(empty_para)
        insert_after = empty_para
        
        # Create horizontal line using underscores
        line_text = "_" * 85
        line_para = self._create_paragraph_element(line_text, is_heading=False)
        insert_after.addnext(line_para)
    
    def _add_top_border(self, paragraph_element):
        """Add a top border to a paragraph element."""
        # Get or create paragraph properties
        pPr = paragraph_element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            paragraph_element.insert(0, pPr)
        
        # Create border element
        pBdr = OxmlElement('w:pBdr')
        
        # Create top border - use thicker line for visibility
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')  # Single line
        top.set(qn('w:sz'), '12')  # 1.5pt (12 eighths of a point) - increased from 6
        top.set(qn('w:space'), '4')  # 4pt spacing - increased from 1
        top.set(qn('w:color'), '000000')  # Black color explicitly
        
        pBdr.append(top)
        pPr.append(pBdr)
    
    def _create_paragraph_element(self, text: str, is_heading: bool = False):
        """Create a paragraph XML element with text."""
        from docx.oxml.ns import nsmap
        
        # Create paragraph
        p = OxmlElement('w:p')
        
        # Create run
        r = OxmlElement('w:r')
        
        # Run properties
        rPr = OxmlElement('w:rPr')
        
        # Font
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), self.FONT_NAME)
        rFonts.set(qn('w:hAnsi'), self.FONT_NAME)
        rFonts.set(qn('w:eastAsia'), self.FONT_NAME)
        rPr.append(rFonts)
        
        # Size in half-points (e.g. 11pt = 22)
        _hp = str(int(self.FONT_SIZE.pt * 2))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), _hp)
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), _hp)
        rPr.append(szCs)
        
        # Bold for headings
        if is_heading:
            b = OxmlElement('w:b')
            rPr.append(b)
        
        r.append(rPr)
        
        # Text
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        
        p.append(r)
        return p
    
    def _create_study_element(self, study: Study, include_protocol: bool, protocol_red: bool,
                              highlight: bool = False):
        """Create a paragraph XML element for a study entry."""
        # Create paragraph
        p = OxmlElement('w:p')
        
        # Paragraph properties for hanging indent
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')  # 0.5 inch in twips
        ind.set(qn('w:hanging'), '720')  # Hanging indent
        pPr.append(ind)
        p.append(pPr)
        
        # Helper to create a run with formatting
        def create_run(text: str, bold: bool = False, red: bool = False,
                       highlight_run: bool = False):
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            # Font
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), self.FONT_NAME)
            rFonts.set(qn('w:hAnsi'), self.FONT_NAME)
            rFonts.set(qn('w:eastAsia'), self.FONT_NAME)
            rPr.append(rFonts)
            
            # Size in half-points
            _hp = str(int(self.FONT_SIZE.pt * 2))
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), _hp)
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), _hp)
            rPr.append(szCs)
            
            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)
            
            if red:
                color = OxmlElement('w:color')
                color.set(qn('w:val'), 'FF0000')
                rPr.append(color)
            
            if highlight_run:
                hl = OxmlElement('w:highlight')
                hl.set(qn('w:val'), 'yellow')
                rPr.append(hl)
            
            r.append(rPr)
            
            t = OxmlElement('w:t')
            t.text = text
            if text.startswith(' ') or text.endswith(' '):
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t)
            
            return r
        
        # Year (not bold) - show blank if year is 0
        year_str = str(study.year) if study.year > 0 else ""
        p.append(create_run(year_str, bold=False))
        
        # Tab
        r_tab = OxmlElement('w:r')
        tab = OxmlElement('w:tab')
        r_tab.append(tab)
        p.append(r_tab)
        
        # Sponsor (bold) — highlighted when highlight=True
        p.append(create_run(study.sponsor, bold=True, highlight_run=highlight))
        
        # Protocol (bold + red, if present and included) — highlighted when highlight=True
        if include_protocol and study.protocol:
            p.append(create_run(' ', bold=False))
            p.append(create_run(study.protocol, bold=True, red=protocol_red,
                               highlight_run=highlight))
        
        # Colon and description
        description = study.description_full if include_protocol else study.description_masked
        
        # Strip sponsor prefix from description_masked if present (handles legacy data)
        if not include_protocol and description:
            # Check if description starts with sponsor name followed by colon
            sponsor_lower = study.sponsor.lower().strip() if study.sponsor else ""
            desc_lower = description.lower().strip()
            if sponsor_lower and desc_lower.startswith(sponsor_lower):
                # Find the colon after sponsor and strip it
                rest = description[len(study.sponsor):].strip()
                if rest.startswith(':'):
                    description = rest[1:].strip()
        
        p.append(create_run(f': {description}', bold=False))
        
        return p
    
    def save(self, output_path: Optional[Path] = None) -> Path:
        """
        Save the document.
        
        If output_path is None, creates a new file with date suffix.
        Returns the path where the file was saved.
        """
        if output_path is None:
            date_str = datetime.now().strftime("%Y-%m-%d")
            stem = self.file_path.stem
            suffix = self.file_path.suffix
            output_path = self.file_path.parent / f"{stem} (Updated {date_str}){suffix}"
        
        self.document.save(output_path)
        return output_path
    
    def save_redacted(self, output_path: Optional[Path] = None) -> Path:
        """Save as redacted version."""
        if output_path is None:
            date_str = datetime.now().strftime("%Y-%m-%d")
            stem = self.file_path.stem
            suffix = self.file_path.suffix
            output_path = self.file_path.parent / f"{stem} (Redacted {date_str}){suffix}"
        
        self.document.save(output_path)
        return output_path


def validate_cv_docx(file_path: Path) -> Tuple[bool, str]:
    """
    Validate a CV .docx file.
    
    Returns: (is_valid, error_message)
    """
    try:
        if not file_path.exists():
            return False, f"File not found: {file_path}"
        
        if not file_path.suffix.lower() == '.docx':
            return False, "File must be a .docx file"
        
        doc = Document(file_path)
        
        # Look for Research Experience section
        found_research = False
        for para in doc.paragraphs:
            if "research experience" in para.text.strip().lower():
                found_research = True
                break
        
        if not found_research:
            return False, "Research Experience section not found in document"
        
        return True, ""
        
    except Exception as e:
        return False, f"Error reading document: {str(e)}"
