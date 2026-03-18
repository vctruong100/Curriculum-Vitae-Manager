import sys
import os
import re
import json
import shutil
from pathlib import Path
from datetime import datetime

import pytest
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from openpyxl import Workbook

APP_ROOT = Path(__file__).parent.parent.resolve()
if str(APP_ROOT) not in sys.path:
    sys.path.insert(0, str(APP_ROOT))

from config import AppConfig, set_config
from processor import CVProcessor
from normalizer import (
    contains_protocol_token,
    is_already_masked,
    normalize_heading_key,
    normalize_subcat_key,
    SPONSOR_PROTOCOL_RE,
)
from docx_handler import CVDocxHandler
from models import Study, ResearchExperience


def _make_master(path, rows=None):
    wb = Workbook()
    ws = wb.active
    ws.title = "Studies"
    if rows is None:
        rows = [
            ("Phase I", None, None),
            ("Oncology", None, None),
            (2024, "Pfizer PF-99999: A Phase 1 study of PF-99999 (pembrolizumab) in advanced lung cancer",
             "Pfizer: A Phase 1 study of XXX in advanced lung cancer"),
            (2023, "Novartis NVS-789: First-in-human study of NVS-789 for metastatic breast cancer",
             "Novartis: First-in-human study of XXX for metastatic breast cancer"),
            ("Cardiology", None, None),
            (2024, "AstraZeneca AZ-111: Phase 1 trial of AZ-111 (dapagliflozin) in heart failure",
             "AstraZeneca: Phase 1 trial of XXX in heart failure"),
            ("Phase II\u2013IV", None, None),
            ("Oncology", None, None),
            (2024, "Roche RO-777: Phase 3 study of RO-777 (atezolizumab) vs placebo in TNBC",
             "Roche: Phase 3 study of XXX vs placebo in TNBC"),
            (2023, "Roche RO-555: Phase 3 randomized trial of RO-555 vs standard of care in NSCLC",
             "Roche: Phase 3 randomized trial of XXX vs standard of care in NSCLC"),
        ]
    row_num = 1
    for item in rows:
        ws.cell(row=row_num, column=1, value=item[0])
        if len(item) > 1 and item[1] is not None:
            ws.cell(row=row_num, column=2, value=item[1])
        if len(item) > 2 and item[2] is not None:
            ws.cell(row=row_num, column=3, value=item[2])
        row_num += 1
    wb.save(path)
    wb.close()
    return path


def _make_cv_with_protocols(path, include_no_protocol=True, include_masked=False):
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_paragraph("Email: jane.doe@example.com")
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")
    doc.add_heading("Research Experience", level=1)

    doc.add_paragraph("Phase I")
    doc.add_paragraph("Oncology")

    p1 = doc.add_paragraph()
    p1.add_run("2024\t")
    run_s = p1.add_run("Pfizer")
    run_s.bold = True
    p1.add_run(" ")
    run_p = p1.add_run("PF-99999")
    run_p.bold = True
    run_p.font.color.rgb = RGBColor(0xFF, 0, 0)
    p1.add_run(": A Phase 1 study of PF-99999 (pembrolizumab) in advanced lung cancer")
    pf = p1.paragraph_format
    pf.left_indent = Inches(0)
    pf.first_line_indent = Inches(-0.5)

    p2 = doc.add_paragraph()
    p2.add_run("2023\t")
    run_s2 = p2.add_run("Novartis")
    run_s2.bold = True
    p2.add_run(" ")
    run_p2 = p2.add_run("NVS-789")
    run_p2.bold = True
    run_p2.font.color.rgb = RGBColor(0xFF, 0, 0)
    p2.add_run(": First-in-human study of NVS-789 for metastatic breast cancer")
    pf2 = p2.paragraph_format
    pf2.left_indent = Inches(0)
    pf2.first_line_indent = Inches(-0.5)

    if include_no_protocol:
        p3 = doc.add_paragraph()
        p3.add_run("2022\t")
        run_s3 = p3.add_run("GenericSponsor")
        run_s3.bold = True
        p3.add_run(": A generic study with no protocol identifier at all")
        pf3 = p3.paragraph_format
        pf3.left_indent = Inches(0)
        pf3.first_line_indent = Inches(-0.5)

    if include_masked:
        p4 = doc.add_paragraph()
        p4.add_run("2021\t")
        run_s4 = p4.add_run("Pfizer")
        run_s4.bold = True
        p4.add_run(": A Phase 1 study of XXX in advanced lung cancer")
        pf4 = p4.paragraph_format
        pf4.left_indent = Inches(0)
        pf4.first_line_indent = Inches(-0.5)

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")
    doc.save(path)
    return path


def _make_cv_two_subcats(path):
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Research Experience", level=1)

    doc.add_paragraph("Phase I")

    doc.add_paragraph("Oncology")
    p1 = doc.add_paragraph()
    p1.add_run("2024\t")
    r = p1.add_run("Pfizer")
    r.bold = True
    p1.add_run(" ")
    rp = p1.add_run("PF-99999")
    rp.bold = True
    rp.font.color.rgb = RGBColor(0xFF, 0, 0)
    p1.add_run(": A Phase 1 study of PF-99999 (pembrolizumab) in advanced lung cancer")
    p1.paragraph_format.left_indent = Inches(0)
    p1.paragraph_format.first_line_indent = Inches(-0.5)

    p2 = doc.add_paragraph()
    p2.add_run("2023\t")
    r2 = p2.add_run("Novartis")
    r2.bold = True
    p2.add_run(" ")
    rp2 = p2.add_run("NVS-789")
    rp2.bold = True
    rp2.font.color.rgb = RGBColor(0xFF, 0, 0)
    p2.add_run(": First-in-human study of NVS-789 for metastatic breast cancer")
    p2.paragraph_format.left_indent = Inches(0)
    p2.paragraph_format.first_line_indent = Inches(-0.5)

    doc.add_paragraph("Cardiology")
    p3 = doc.add_paragraph()
    p3.add_run("2022\t")
    r3 = p3.add_run("GenericSponsor")
    r3.bold = True
    p3.add_run(": A generic cardiology study with no protocol at all")
    p3.paragraph_format.left_indent = Inches(0)
    p3.paragraph_format.first_line_indent = Inches(-0.5)

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J. 2023.")
    doc.save(path)
    return path


def _make_cv_in_table(path):
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Research Experience", level=1)

    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    cell.paragraphs[0].text = "Phase I"

    p_sub = cell.add_paragraph("Oncology")

    p1 = cell.add_paragraph()
    p1.add_run("2024\t")
    r = p1.add_run("Pfizer")
    r.bold = True
    p1.add_run(" ")
    rp = p1.add_run("PF-99999")
    rp.bold = True
    rp.font.color.rgb = RGBColor(0xFF, 0, 0)
    p1.add_run(": A Phase 1 study of PF-99999 (pembrolizumab) in advanced lung cancer")

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("Doe J. 2023.")
    doc.save(path)
    return path


def _read_para_texts(docx_path):
    doc = Document(docx_path)
    return [p.text for p in doc.paragraphs]


class TestContainsProtocolToken:

    def test_with_protocol(self):
        assert contains_protocol_token("Pfizer PF-99999: A study of PF-99999")

    def test_without_protocol(self):
        assert not contains_protocol_token("GenericSponsor: A generic study with no protocol")

    def test_already_masked(self):
        assert not contains_protocol_token("Pfizer: A study of XXX in cancer")

    def test_split_run_protocol(self):
        assert contains_protocol_token("Novartis NVS-789: First-in-human study")

    def test_dash_variants(self):
        assert contains_protocol_token("Roche RO\u2013777: A study")

    def test_short_prefix_rejected(self):
        assert not contains_protocol_token("A PF-99999: study")


class TestIsAlreadyMasked:

    def test_masked_line(self):
        assert is_already_masked("Pfizer: A study of XXX in cancer")

    def test_not_masked(self):
        assert not is_already_masked("Pfizer PF-99999: A study of PF-99999 in cancer")

    def test_no_xxx_not_masked(self):
        assert not is_already_masked("Pfizer: A study in cancer")


class TestSponsorProtocolRegex:

    def test_basic_match(self):
        m = SPONSOR_PROTOCOL_RE.match("Pfizer PF-99999")
        assert m is not None
        assert m.group("sponsor").strip() == "Pfizer"
        assert m.group("protocol") == "PF-99999"

    def test_multi_word_sponsor(self):
        m = SPONSOR_PROTOCOL_RE.match("Eli Lilly LY246736")
        assert m is not None
        assert "Lilly" in m.group("sponsor")


class TestOnlyProtocolBearingRedacted:

    def test_protocol_study_is_redacted(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=True)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        texts = _read_para_texts(Path(result.output_path))
        has_redacted = False
        for t in texts:
            if "XXX" in t:
                has_redacted = True
                break
        assert has_redacted

    def test_no_protocol_study_unchanged(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=True)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        texts = _read_para_texts(Path(result.output_path))
        generic_lines = [t for t in texts if "GenericSponsor" in t]
        assert len(generic_lines) == 1
        assert "generic study with no protocol" in generic_lines[0].lower()

    def test_log_operations(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=True)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        ops = [e.operation for e in result.log_entries]
        assert "replaced" in ops
        assert "skipped-no-protocol" in ops
        config_entries = [e for e in result.log_entries if e.operation == "config"]
        assert len(config_entries) >= 1
        assert "sort_and_format=False" in config_entries[0].details


class TestOrderUnchangedWhenUnchecked:

    def test_paragraph_order_preserved(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=True)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        original_texts = _read_para_texts(cv_path)
        original_count = len(original_texts)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        output_texts = _read_para_texts(Path(result.output_path))
        assert len(output_texts) == original_count

    def test_heading_positions_preserved(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=True)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        original_texts = _read_para_texts(cv_path)
        phase_positions = [
            i for i, t in enumerate(original_texts)
            if t.strip() == "Phase I"
        ]

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        output_texts = _read_para_texts(Path(result.output_path))
        new_phase_positions = [
            i for i, t in enumerate(output_texts)
            if t.strip() == "Phase I"
        ]
        assert phase_positions == new_phase_positions


class TestSortAndFormatChecked:

    def test_affected_category_resorted(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=False)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=True,
        )
        assert result.success

        ops = [e.operation for e in result.log_entries]
        assert "sort-category" in ops

    def test_unaffected_category_untouched(self, tmp_path):
        cv_path = _make_cv_two_subcats(tmp_path / "cv.docx")
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        original_texts = _read_para_texts(cv_path)
        cardiology_idx = None
        for i, t in enumerate(original_texts):
            if t.strip() == "Cardiology":
                cardiology_idx = i
                break

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=True,
        )
        assert result.success

        output_texts = _read_para_texts(Path(result.output_path))
        new_cardiology_idx = None
        for i, t in enumerate(output_texts):
            if t.strip() == "Cardiology":
                new_cardiology_idx = i
                break

        assert cardiology_idx == new_cardiology_idx
        card_study = output_texts[new_cardiology_idx + 1]
        assert "GenericSponsor" in card_study
        assert "generic cardiology" in card_study.lower()


class TestIdempotency:

    def test_second_run_no_changes(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=False)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result1 = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result1.success
        first_output = Path(result1.output_path)

        result2 = processor.mode_b_redact_protocols(
            first_output, master_path=master_path, sort_and_format=False,
        )
        assert result2.success

        ops = [e.operation for e in result2.log_entries]
        assert "replaced" not in ops

        masked_ops = [
            e for e in result2.log_entries
            if e.operation == "skipped-already-masked"
        ]
        assert len(masked_ops) >= 1

    def test_idempotent_with_sort(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=False)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result1 = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=True,
        )
        assert result1.success
        first_output = Path(result1.output_path)

        result2 = processor.mode_b_redact_protocols(
            first_output, master_path=master_path, sort_and_format=True,
        )
        assert result2.success

        replaced_ops = [
            e for e in result2.log_entries if e.operation == "replaced"
        ]
        assert len(replaced_ops) == 0


class TestPreviewModeB:

    def test_preview_lists_anchors_and_masked_text(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=True)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        changes, error = processor.preview_changes(
            cv_path, master_path, mode="redact_protocols",
            sort_and_format=False,
        )
        assert error == ""

        redact_changes = [c for c in changes if c["action"] == "redact"]
        assert len(redact_changes) >= 1
        for c in redact_changes:
            assert "anchor_para_idx" in c
            assert "new_description" in c
            assert "new_sponsor" in c
            assert "XXX" in c["new_description"]

    def test_preview_shows_skipped_no_protocol(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=True)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        changes, error = processor.preview_changes(
            cv_path, master_path, mode="redact_protocols",
        )
        assert error == ""

        skip_changes = [
            c for c in changes if c["action"] == "skipped-no-protocol"
        ]
        assert len(skip_changes) >= 1

    def test_preview_shows_resort_indicator(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=False)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        changes, error = processor.preview_changes(
            cv_path, master_path, mode="redact_protocols",
            sort_and_format=True,
        )
        assert error == ""

        redact_changes = [c for c in changes if c["action"] == "redact"]
        assert len(redact_changes) >= 1
        for c in redact_changes:
            assert c.get("would_resort_category") is True
            assert c.get("sort_and_format") is True

    def test_preview_no_resort_when_unchecked(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=False)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        changes, error = processor.preview_changes(
            cv_path, master_path, mode="redact_protocols",
            sort_and_format=False,
        )
        assert error == ""

        redact_changes = [c for c in changes if c["action"] == "redact"]
        for c in redact_changes:
            assert c.get("sort_and_format") is False

    def test_preview_does_not_modify_document(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx")
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        original_texts = _read_para_texts(cv_path)

        processor = CVProcessor(config)
        processor.preview_changes(
            cv_path, master_path, mode="redact_protocols",
        )

        after_texts = _read_para_texts(cv_path)
        assert original_texts == after_texts


class TestOutputRouting:

    def test_result_folder_structure(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "Jane Doe CV.docx")
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        output_path = Path(result.output_path)
        assert output_path.exists()
        assert output_path.suffix == ".docx"
        assert output_path.parent.name == "Jane Doe CV"
        assert output_path.parent.parent.name == "result"

    def test_no_logs_in_result_folder(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "Jane Doe CV.docx")
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        output_path = Path(result.output_path)
        result_folder = output_path.parent
        for f in result_folder.iterdir():
            assert f.suffix in (".docx", ".xlsx"), (
                f"Unexpected file in result folder: {f.name}"
            )


class TestAlreadyMaskedStudy:

    def test_masked_study_skipped(self, tmp_path):
        cv_path = _make_cv_with_protocols(
            tmp_path / "cv.docx",
            include_no_protocol=False,
            include_masked=True,
        )
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        masked_ops = [
            e for e in result.log_entries
            if e.operation == "skipped-already-masked"
        ]
        assert len(masked_ops) >= 1


class TestCustomDocProperty:

    def test_original_cv_name_set(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "Jane Doe CV.docx")
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        output_path = Path(result.output_path)
        doc = Document(output_path)
        kw = doc.core_properties.keywords or ""
        assert "_original_cv_name:" in kw
        assert "Jane Doe CV" in kw


class TestCleanup:

    def test_tmp_path_isolated(self, tmp_path):
        marker = tmp_path / "marker.txt"
        marker.write_text("test")
        assert marker.exists()

    def test_no_leftover_from_previous(self, tmp_path):
        children = list(tmp_path.iterdir())
        assert len(children) == 0


class TestSortAndFormatLogging:

    def test_sort_flag_logged(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=False)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=True,
        )
        assert result.success

        config_entries = [
            e for e in result.log_entries if e.operation == "config"
        ]
        assert any("sort_and_format=True" in e.details for e in config_entries)

    def test_sort_false_logged(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "cv.docx", include_no_protocol=False)
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        config_entries = [
            e for e in result.log_entries if e.operation == "config"
        ]
        assert any("sort_and_format=False" in e.details for e in config_entries)


class TestSplitRunRedaction:

    def test_split_run_protocol_detected_and_redacted(self, tmp_path):
        doc = Document()
        doc.add_heading("Research Experience", level=1)
        doc.add_paragraph("Phase I")
        doc.add_paragraph("Oncology")

        p = doc.add_paragraph()
        p.add_run("2024\t")
        r1 = p.add_run("Pfizer")
        r1.bold = True
        p.add_run(" ")
        r2 = p.add_run("PF-")
        r2.bold = True
        r2.font.color.rgb = RGBColor(0xFF, 0, 0)
        r3 = p.add_run("99999")
        r3.bold = True
        r3.font.color.rgb = RGBColor(0xFF, 0, 0)
        p.add_run(": A Phase 1 study of PF-99999 (pembrolizumab) in advanced lung cancer")

        doc.add_heading("Publications", level=1)

        cv_path = tmp_path / "split_run_cv.docx"
        doc.save(cv_path)

        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        texts = _read_para_texts(Path(result.output_path))
        study_lines = [t for t in texts if "lung cancer" in t.lower()]
        assert len(study_lines) >= 1
        for line in study_lines:
            assert "PF-99999" not in line
            assert "XXX" in line


class TestNoColorReliance:

    def test_redaction_without_red_color(self, tmp_path):
        doc = Document()
        doc.add_heading("Research Experience", level=1)
        doc.add_paragraph("Phase I")
        doc.add_paragraph("Oncology")

        p = doc.add_paragraph()
        p.add_run("2024\t")
        r = p.add_run("Pfizer")
        r.bold = True
        p.add_run(" ")
        rp = p.add_run("PF-99999")
        rp.bold = True
        p.add_run(": A Phase 1 study of PF-99999 (pembrolizumab) in advanced lung cancer")

        doc.add_heading("Publications", level=1)

        cv_path = tmp_path / "no_color_cv.docx"
        doc.save(cv_path)

        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path, sort_and_format=False,
        )
        assert result.success

        texts = _read_para_texts(Path(result.output_path))
        study_lines = [t for t in texts if "lung cancer" in t.lower()]
        assert len(study_lines) >= 1
        for line in study_lines:
            assert "PF-99999" not in line
            assert "XXX" in line


class TestModeAtoModeBSameFolder:

    def test_mode_b_after_mode_a_writes_to_same_folder(self, tmp_path):
        cv_path = _make_cv_with_protocols(tmp_path / "Jane Doe CV.docx")
        master_path = _make_master(tmp_path / "master.xlsx")
        config = AppConfig(data_root=str(tmp_path / "data"))
        config.ensure_user_directories()
        set_config(config)

        processor = CVProcessor(config)
        result_a = processor.mode_a_update_inject(
            cv_path, master_path=master_path, enable_sort_existing=True,
        )
        assert result_a.success

        result_b = processor.mode_b_redact_protocols(
            Path(result_a.output_path),
            master_path=master_path,
            sort_and_format=False,
        )
        assert result_b.success

        path_a = Path(result_a.output_path)
        path_b = Path(result_b.output_path)
        assert path_a.parent == path_b.parent
