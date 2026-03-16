"""
Tests for the hierarchy parsing and insertion deduplication fix.

Reproduces the bug where new studies were inserted under a duplicate
Phase/Subcategory block instead of the existing one, and verifies
the fix across all required edge cases.

Covers:
- Phase casing variants: "PHASE I", "Phase I", "Phase 1", "phase i"
- Subcategory casing and spacing: "Healthy Adults", "HEALTHY ADULTS", "Healthy  Adults"
- Headings split across runs in the .docx
- Extra whitespace, smart quotes, long dashes
- Documents where Phase and Subcategory are present but out of canonical sort order
- Idempotent reruns
- Both global-sort and no-sort-existing paths
- Preview JSON includes container matching info
- Normalizer unit tests for normalize_heading_key, normalize_subcat_key
- PHASE_SYNONYMS coverage
"""

import sys
import json
from pathlib import Path

import pytest
from docx import Document
from docx.shared import RGBColor, Inches, Pt

sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from config import AppConfig, set_config
from processor import CVProcessor
from models import Study, ResearchExperience, Phase, Subcategory
from normalizer import (
    normalize_heading_key,
    normalize_subcat_key,
    is_phase_heading,
    PHASE_SYNONYMS,
)
from tests.conftest import _make_master_xlsx


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_study_paragraph(doc, year, sponsor, protocol, description):
    """Add a formatted study paragraph to a document."""
    p = doc.add_paragraph()
    p.add_run(f"{year}\t")
    run_s = p.add_run(sponsor)
    run_s.bold = True
    if protocol:
        p.add_run(" ")
        run_p = p.add_run(protocol)
        run_p.bold = True
        run_p.font.color.rgb = RGBColor(0xFF, 0, 0)
    p.add_run(f": {description}")
    pf = p.paragraph_format
    pf.left_indent = Inches(0)
    pf.first_line_indent = Inches(-0.5)
    return p


def _make_bug_cv(path: Path, phase_text="PHASE I", subcat_text="Healthy Adults"):
    """
    Create a CV that reproduces the original bug scenario:
    Research Experience section with the given phase/subcat headings
    and two 2025 studies.
    """
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_paragraph("Email: jane.doe@example.com")
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")
    doc.add_heading("Research Experience", level=1)

    doc.add_paragraph(phase_text)
    doc.add_paragraph(subcat_text)

    _add_study_paragraph(
        doc, 2025, "ELI LILLY AND COMPANY", "",
        "A single-ascending and multiple-ascending dose study of XXXX "
        "in healthy participants, participants with obesity and "
        "hypertension, and participants with decreased estimated "
        "glomerular filtration rate"
    )
    _add_study_paragraph(
        doc, 2025, "VISTERRA", "",
        "A phase 1, randomized, placebo-controlled, double blind, "
        "multiple ascending dose trial to assess the safety, "
        "tolerability, pharmacokinetics, and pharmacodynamics of XXXX "
        "in healthy participants"
    )

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")
    doc.save(path)
    return path


def _make_bug_master(path: Path):
    """
    Master that contains the same studies as the bug CV plus two new
    studies that should be injected into the SAME Phase I > Healthy Adults.
    """
    studies_data = [
        ("Phase I", None, None),
        ("Healthy Adults", None, None),
        (2026,
         "MERCK MK4082-002: A Multiple-Ascending Dose Study to Evaluate "
         "the Safety, Tolerability and Pharmacokinetics of MK-4082 in "
         "Healthy Overweight/Obese Participants",
         "MERCK: A Multiple-Ascending Dose Study to Evaluate the Safety, "
         "Tolerability and Pharmacokinetics of XXX in Healthy "
         "Overweight/Obese Participants"),
        (2025,
         "ABBVIE M24-920: A Phase I Pharmacokinetic Study in Healthy "
         "Subjects to Evaluate the Relative Bioavailability of "
         "Risankizumab Following Subcutaneous Administration with "
         "On-Body Injector",
         "ABBVIE: A Phase I Pharmacokinetic Study in Healthy Subjects "
         "to Evaluate the Relative Bioavailability of XXX Following "
         "Subcutaneous Administration with On-Body Injector"),
        (2025,
         "ELI LILLY AND COMPANY: A single-ascending and "
         "multiple-ascending dose study of XXXX in healthy participants, "
         "participants with obesity and hypertension, and participants "
         "with decreased estimated glomerular filtration rate",
         "ELI LILLY AND COMPANY: A single-ascending and "
         "multiple-ascending dose study of XXXX in healthy participants, "
         "participants with obesity and hypertension, and participants "
         "with decreased estimated glomerular filtration rate"),
        (2025,
         "VISTERRA: A phase 1, randomized, placebo-controlled, double "
         "blind, multiple ascending dose trial to assess the safety, "
         "tolerability, pharmacokinetics, and pharmacodynamics of XXXX "
         "in healthy participants",
         "VISTERRA: A phase 1, randomized, placebo-controlled, double "
         "blind, multiple ascending dose trial to assess the safety, "
         "tolerability, pharmacokinetics, and pharmacodynamics of XXXX "
         "in healthy participants"),
    ]
    return _make_master_xlsx(path, studies_data)


def _extract_headings_and_studies(docx_path: Path):
    """
    Read a .docx and return a structured list of
    (type, text) tuples for the Research Experience section.
    type is 'phase', 'subcat', or 'study'.
    """
    doc = Document(docx_path)
    in_section = False
    items = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text.lower().startswith("research experience"):
            in_section = True
            continue
        if in_section:
            if para.style and para.style.name and "heading" in para.style.name.lower():
                if not text.lower().startswith("research experience"):
                    break
            phase = is_phase_heading(text)
            if phase is not None:
                items.append(("phase", phase))
                continue
            parts = text.replace("\t", " ").split()
            if parts and parts[0].isdigit() and len(parts[0]) == 4:
                items.append(("study", text))
                continue
            if text and not text.startswith("_"):
                items.append(("subcat", text))
    return items


def _count_phase_occurrences(items, canonical_name):
    """Count how many times a phase heading appears in the items list."""
    key = normalize_heading_key(canonical_name)
    return sum(
        1
        for typ, txt in items
        if typ == "phase" and normalize_heading_key(txt) == key
    )


def _count_subcat_occurrences(items, canonical_name):
    """Count how many times a subcategory heading appears."""
    key = normalize_subcat_key(canonical_name)
    return sum(
        1
        for typ, txt in items
        if typ == "subcat" and normalize_subcat_key(txt) == key
    )


# ===========================================================================
# Normalizer unit tests
# ===========================================================================

class TestNormalizeHeadingKey:
    """Unit tests for normalize_heading_key."""

    def test_phase_i_variants(self):
        assert normalize_heading_key("Phase I") == "Phase I"
        assert normalize_heading_key("PHASE I") == "Phase I"
        assert normalize_heading_key("phase i") == "Phase I"
        assert normalize_heading_key("Phase 1") == "Phase I"
        assert normalize_heading_key("PHASE 1") == "Phase I"
        assert normalize_heading_key("Phase I.") == "Phase I"

    def test_phase_ii_iv_dash_variants(self):
        assert normalize_heading_key("Phase II\u2013IV") == "Phase II\u2013IV"
        assert normalize_heading_key("Phase II-IV") == "Phase II\u2013IV"
        assert normalize_heading_key("Phase II\u2014IV") == "Phase II\u2013IV"
        assert normalize_heading_key("Phase 2-4") == "Phase II\u2013IV"
        assert normalize_heading_key("PHASE II-IV") == "Phase II\u2013IV"

    def test_whitespace_collapse(self):
        assert normalize_heading_key("Phase  I") == "Phase I"
        assert normalize_heading_key("  Phase I  ") == "Phase I"
        assert normalize_heading_key("Phase\tI") == "Phase I"

    def test_non_phase_passthrough(self):
        result = normalize_heading_key("Oncology")
        assert result == "oncology"

    def test_uncategorized(self):
        assert normalize_heading_key("Uncategorized") == "Uncategorized"
        assert normalize_heading_key("UNCATEGORIZED") == "Uncategorized"
        assert normalize_heading_key("uncategorized") == "Uncategorized"


class TestNormalizeSubcatKey:
    """Unit tests for normalize_subcat_key."""

    def test_casefolding(self):
        assert normalize_subcat_key("Healthy Adults") == "healthy adults"
        assert normalize_subcat_key("HEALTHY ADULTS") == "healthy adults"
        assert normalize_subcat_key("healthy adults") == "healthy adults"

    def test_whitespace_collapse(self):
        assert normalize_subcat_key("Healthy  Adults") == "healthy adults"
        assert normalize_subcat_key("  Healthy Adults  ") == "healthy adults"
        assert normalize_subcat_key("Healthy\tAdults") == "healthy adults"

    def test_dash_normalization(self):
        assert normalize_subcat_key("Auto\u2013Immune") == "auto-immune"
        assert normalize_subcat_key("Auto\u2014Immune") == "auto-immune"
        assert normalize_subcat_key("Auto-Immune") == "auto-immune"

    def test_smart_quotes(self):
        assert normalize_subcat_key("Parkinson\u2019s Disease") == "parkinson's disease"


class TestIsPhaseHeading:
    """is_phase_heading must recognize all variants."""

    def test_all_caps(self):
        assert is_phase_heading("PHASE I") == "Phase I"

    def test_title_case(self):
        assert is_phase_heading("Phase I") == "Phase I"

    def test_arabic_numeral(self):
        assert is_phase_heading("Phase 1") == "Phase I"

    def test_lowercase(self):
        assert is_phase_heading("phase i") == "Phase I"

    def test_phase_ii_iv_endash(self):
        assert is_phase_heading("Phase II\u2013IV") == "Phase II\u2013IV"

    def test_phase_ii_iv_hyphen(self):
        assert is_phase_heading("Phase II-IV") == "Phase II\u2013IV"

    def test_non_phase(self):
        assert is_phase_heading("Healthy Adults") is None
        assert is_phase_heading("2025") is None
        assert is_phase_heading("") is None


# ===========================================================================
# Model-level dedup tests
# ===========================================================================

class TestPhaseDedup:
    """get_or_create_phase must coalesce all phase name variants."""

    def test_case_variants_coalesce(self):
        re_exp = ResearchExperience()
        p1 = re_exp.get_or_create_phase("PHASE I")
        p2 = re_exp.get_or_create_phase("Phase I")
        p3 = re_exp.get_or_create_phase("phase i")
        assert p1 is p2
        assert p2 is p3
        assert len(re_exp.phases) == 1

    def test_arabic_roman_coalesce(self):
        re_exp = ResearchExperience()
        p1 = re_exp.get_or_create_phase("Phase 1")
        p2 = re_exp.get_or_create_phase("Phase I")
        assert p1 is p2
        assert len(re_exp.phases) == 1

    def test_dash_variants_coalesce(self):
        re_exp = ResearchExperience()
        p1 = re_exp.get_or_create_phase("Phase II-IV")
        p2 = re_exp.get_or_create_phase("Phase II\u2013IV")
        assert p1 is p2
        assert len(re_exp.phases) == 1

    def test_preserves_first_name(self):
        re_exp = ResearchExperience()
        re_exp.get_or_create_phase("PHASE I")
        re_exp.get_or_create_phase("Phase I")
        assert re_exp.phases[0].name == "PHASE I"


class TestSubcategoryDedup:
    """get_or_create_subcategory must coalesce subcategory name variants."""

    def test_case_variants(self):
        phase = Phase(name="Phase I")
        sc1 = phase.get_or_create_subcategory("Healthy Adults")
        sc2 = phase.get_or_create_subcategory("HEALTHY ADULTS")
        sc3 = phase.get_or_create_subcategory("healthy adults")
        assert sc1 is sc2
        assert sc2 is sc3
        assert len(phase.subcategories) == 1

    def test_whitespace_variants(self):
        phase = Phase(name="Phase I")
        sc1 = phase.get_or_create_subcategory("Healthy Adults")
        sc2 = phase.get_or_create_subcategory("Healthy  Adults")
        assert sc1 is sc2
        assert len(phase.subcategories) == 1

    def test_preserves_first_name(self):
        phase = Phase(name="Phase I")
        phase.get_or_create_subcategory("HEALTHY ADULTS")
        phase.get_or_create_subcategory("Healthy Adults")
        assert phase.subcategories[0].name == "HEALTHY ADULTS"


class TestIdentityTupleNormalization:
    """get_identity_tuple must normalize phase/subcat for dedup."""

    def test_phase_variants_produce_same_tuple(self):
        s1 = Study(
            phase="Phase I", subcategory="Oncology", year=2025,
            sponsor="Pfizer", protocol="PF-1",
            description_full="desc", description_masked="masked",
        )
        s2 = Study(
            phase="PHASE I", subcategory="Oncology", year=2025,
            sponsor="Pfizer", protocol="PF-1",
            description_full="desc", description_masked="masked",
        )
        assert s1.get_identity_tuple() == s2.get_identity_tuple()

    def test_subcat_variants_produce_same_tuple(self):
        s1 = Study(
            phase="Phase I", subcategory="Healthy Adults", year=2025,
            sponsor="X", protocol="",
            description_full="d", description_masked="m",
        )
        s2 = Study(
            phase="Phase I", subcategory="HEALTHY ADULTS", year=2025,
            sponsor="X", protocol="",
            description_full="d", description_masked="m",
        )
        assert s1.get_identity_tuple() == s2.get_identity_tuple()

    def test_equality_across_variants(self):
        s1 = Study(
            phase="Phase 1", subcategory="Healthy  Adults", year=2025,
            sponsor="X", protocol="",
            description_full="d", description_masked="m",
        )
        s2 = Study(
            phase="Phase I", subcategory="Healthy Adults", year=2025,
            sponsor="X", protocol="",
            description_full="d", description_masked="m",
        )
        assert s1 == s2


# ===========================================================================
# End-to-end: exact bug reproduction
# ===========================================================================

class TestBugReproduction:
    """Reproduce the original bug and verify the fix."""

    def test_no_duplicate_phase_or_subcat(self, app_config, tmp_dir):
        """
        CV has "PHASE I" -> "Healthy Adults" with 2025 studies.
        Master has "Phase I" -> "Healthy Adults" with same + new studies.
        After Update/Inject, there must be exactly ONE Phase I and
        ONE Healthy Adults — no duplicates.
        """
        cv_path = _make_bug_cv(tmp_dir / "cv.docx")
        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
            enable_sort_existing=True,
        )
        assert result.success is True, result.error_message

        items = _extract_headings_and_studies(output)
        phase_count = _count_phase_occurrences(items, "Phase I")
        subcat_count = _count_subcat_occurrences(items, "Healthy Adults")

        assert phase_count == 1, (
            f"Expected 1 Phase I heading, found {phase_count}. Items: {items}"
        )
        assert subcat_count == 1, (
            f"Expected 1 Healthy Adults heading, found {subcat_count}. Items: {items}"
        )

    def test_new_studies_under_existing_block(self, app_config, tmp_dir):
        """New studies must appear under the same Phase I > Healthy Adults.

        The CV has 2025 studies so year_bound = 2025.  Only master studies
        with year > 2025 are injected (MERCK 2026).  The ABBVIE 2025 study
        is correctly skipped because year <= year_bound.
        """
        cv_path = _make_bug_cv(tmp_dir / "cv.docx")
        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
            enable_sort_existing=True,
        )
        assert result.success is True

        items = _extract_headings_and_studies(output)
        study_texts = [txt for typ, txt in items if typ == "study"]

        assert len(study_texts) >= 3, (
            f"Expected at least 3 studies (2 existing + 1 new 2026), "
            f"got {len(study_texts)}"
        )

        merck_found = any("MERCK" in s for s in study_texts)
        lilly_found = any("ELI LILLY" in s or "LILLY" in s for s in study_texts)
        visterra_found = any("VISTERRA" in s for s in study_texts)

        assert merck_found, "MERCK 2026 study not found in output"
        assert lilly_found, "ELI LILLY study not found in output"
        assert visterra_found, "VISTERRA study not found in output"


# ===========================================================================
# Phase casing variants
# ===========================================================================

class TestPhaseCasingVariants:
    """CV phase heading in various casings must all coalesce."""

    @pytest.mark.parametrize("phase_text", [
        "PHASE I", "Phase I", "Phase 1", "phase i",
    ])
    def test_no_duplicate_with_variant(self, app_config, tmp_dir, phase_text):
        cv_path = _make_bug_cv(tmp_dir / "cv.docx", phase_text=phase_text)
        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
            enable_sort_existing=True,
        )
        assert result.success is True

        items = _extract_headings_and_studies(output)
        phase_count = _count_phase_occurrences(items, "Phase I")
        assert phase_count == 1, (
            f"phase_text='{phase_text}': expected 1 Phase I, got {phase_count}. "
            f"Items: {items}"
        )


# ===========================================================================
# Subcategory casing and spacing variants
# ===========================================================================

class TestSubcatVariants:
    """Subcategory variants must coalesce."""

    @pytest.mark.parametrize("subcat_text", [
        "Healthy Adults", "HEALTHY ADULTS", "Healthy  Adults",
    ])
    def test_no_duplicate_subcat(self, app_config, tmp_dir, subcat_text):
        cv_path = _make_bug_cv(
            tmp_dir / "cv.docx", subcat_text=subcat_text
        )
        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
            enable_sort_existing=True,
        )
        assert result.success is True

        items = _extract_headings_and_studies(output)
        subcat_count = _count_subcat_occurrences(items, "Healthy Adults")
        assert subcat_count == 1, (
            f"subcat_text='{subcat_text}': expected 1 subcat, got {subcat_count}. "
            f"Items: {items}"
        )


# ===========================================================================
# Split runs
# ===========================================================================

class TestSplitRuns:
    """Phase/subcategory headings split across multiple runs."""

    def test_phase_split_across_runs(self, app_config, tmp_dir):
        doc = Document()
        doc.add_heading("Jane Doe, MD", level=0)
        doc.add_heading("Research Experience", level=1)

        p = doc.add_paragraph()
        r1 = p.add_run("PHASE ")
        r1.bold = True
        r2 = p.add_run("I")
        r2.bold = True

        doc.add_paragraph("Healthy Adults")

        _add_study_paragraph(
            doc, 2025, "ACME", "",
            "A study of XXXX in healthy participants"
        )

        doc.add_heading("Publications", level=1)
        cv_path = tmp_dir / "split_cv.docx"
        doc.save(cv_path)

        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
            enable_sort_existing=True,
        )
        assert result.success is True

        items = _extract_headings_and_studies(output)
        phase_count = _count_phase_occurrences(items, "Phase I")
        assert phase_count == 1, (
            f"Split-run Phase: expected 1, got {phase_count}. Items: {items}"
        )

    def test_subcat_split_across_runs(self, app_config, tmp_dir):
        doc = Document()
        doc.add_heading("Jane Doe, MD", level=0)
        doc.add_heading("Research Experience", level=1)

        doc.add_paragraph("Phase I")

        p = doc.add_paragraph()
        r1 = p.add_run("Healthy ")
        r1.bold = True
        r2 = p.add_run("Adults")
        r2.bold = True

        _add_study_paragraph(
            doc, 2025, "ACME", "",
            "A study of XXXX in healthy participants"
        )

        doc.add_heading("Publications", level=1)
        cv_path = tmp_dir / "split_subcat_cv.docx"
        doc.save(cv_path)

        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
            enable_sort_existing=True,
        )
        assert result.success is True

        items = _extract_headings_and_studies(output)
        subcat_count = _count_subcat_occurrences(items, "Healthy Adults")
        assert subcat_count == 1, (
            f"Split-run Subcat: expected 1, got {subcat_count}. Items: {items}"
        )


# ===========================================================================
# Idempotency
# ===========================================================================

class TestIdempotency:
    """Running Update/Inject twice must produce no additional changes."""

    def test_idempotent_run(self, app_config, tmp_dir):
        cv_path = _make_bug_cv(tmp_dir / "cv.docx")
        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)

        out1 = tmp_dir / "out1.docx"
        r1 = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=out1,
            enable_sort_existing=True,
        )
        assert r1.success is True

        out2 = tmp_dir / "out2.docx"
        r2 = processor.mode_a_update_inject(
            out1,
            master_path=master_path,
            output_path=out2,
            enable_sort_existing=True,
        )
        assert r2.success is True

        items1 = _extract_headings_and_studies(out1)
        items2 = _extract_headings_and_studies(out2)

        studies1 = [t for typ, t in items1 if typ == "study"]
        studies2 = [t for typ, t in items2 if typ == "study"]

        assert len(studies1) == len(studies2), (
            f"Idempotency violated: {len(studies1)} vs {len(studies2)} studies"
        )

        assert _count_phase_occurrences(items2, "Phase I") == 1
        assert _count_subcat_occurrences(items2, "Healthy Adults") == 1


# ===========================================================================
# Sort-existing paths
# ===========================================================================

class TestSortExistingPaths:
    """Both enable_sort_existing=True and False must avoid duplicates."""

    def test_sort_true_no_duplicates(self, app_config, tmp_dir):
        cv_path = _make_bug_cv(tmp_dir / "cv.docx")
        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "sorted.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
            enable_sort_existing=True,
        )
        assert result.success is True

        items = _extract_headings_and_studies(output)
        assert _count_phase_occurrences(items, "Phase I") == 1
        assert _count_subcat_occurrences(items, "Healthy Adults") == 1

    def test_sort_false_no_duplicates(self, app_config, tmp_dir):
        cv_path = _make_bug_cv(tmp_dir / "cv.docx")
        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "unsorted.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
            enable_sort_existing=False,
        )
        assert result.success is True

        items = _extract_headings_and_studies(output)
        assert _count_phase_occurrences(items, "Phase I") == 1
        assert _count_subcat_occurrences(items, "Healthy Adults") == 1


# ===========================================================================
# Preview JSON container info
# ===========================================================================

class TestPreviewContainerInfo:
    """Preview JSON must include matched container information."""

    def test_preview_includes_container_keys(self, app_config, tmp_dir):
        cv_path = _make_bug_cv(tmp_dir / "cv.docx")
        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        changes, error = processor.preview_changes(
            cv_path,
            master_path=master_path,
            mode="update_inject",
        )
        assert error == ""

        for change in changes:
            assert "phase_key" in change
            assert "subcat_key" in change
            assert "matched_phase_container" in change
            assert "matched_subcat_container" in change
            assert change["phase_key"] == "Phase I"
            assert change["subcat_key"] == "healthy adults"


# ===========================================================================
# Logging verification
# ===========================================================================

class TestLoggingContainerMatch:
    """Log entries must include phase_key and subcat_key."""

    def test_inserted_log_has_keys(self, app_config, tmp_dir):
        cv_path = _make_bug_cv(tmp_dir / "cv.docx")
        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
        )
        assert result.success is True

        inserted = [
            e for e in result.log_entries
            if e.operation == "inserted"
        ]
        for entry in inserted:
            assert "phase_key=" in entry.details
            assert "subcat_key=" in entry.details
            assert "container_phase=" in entry.details
            assert "container_subcat=" in entry.details

    def test_matched_log_has_keys(self, app_config, tmp_dir):
        cv_path = _make_bug_cv(tmp_dir / "cv.docx")
        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
        )
        assert result.success is True

        matched = [
            e for e in result.log_entries
            if e.operation == "matched-existing"
        ]
        for entry in matched:
            assert "phase_key=" in entry.details
            assert "subcat_key=" in entry.details


# ===========================================================================
# Out-of-order phases in document
# ===========================================================================

class TestOutOfOrderPhases:
    """Phases present but not in canonical sort order still coalesce."""

    def test_out_of_order_inserts_correctly(self, app_config, tmp_dir):
        doc = Document()
        doc.add_heading("Jane Doe, MD", level=0)
        doc.add_heading("Research Experience", level=1)

        doc.add_paragraph("Phase II\u2013IV")
        doc.add_paragraph("Oncology")
        _add_study_paragraph(
            doc, 2024, "Roche", "RO-777",
            "Phase 3 study of RO-777 in TNBC"
        )

        doc.add_paragraph("Phase I")
        doc.add_paragraph("Healthy Adults")
        _add_study_paragraph(
            doc, 2025, "ACME", "",
            "A study of XXXX in healthy participants"
        )

        doc.add_heading("Publications", level=1)
        cv_path = tmp_dir / "ooo_cv.docx"
        doc.save(cv_path)

        master_path = _make_bug_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
            enable_sort_existing=True,
        )
        assert result.success is True

        items = _extract_headings_and_studies(output)
        assert _count_phase_occurrences(items, "Phase I") == 1
        assert _count_subcat_occurrences(items, "Healthy Adults") == 1
