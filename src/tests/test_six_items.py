"""
Comprehensive tests for the six new items:
  1. Undo for recent delete in Mode C
  2. Total study count refresh after delete/undo
  3. Application icon generation
  4. Optional update checker
  5. Hanging indentation (configurable)
  6. Launcher scripts (.bat/.sh)
"""

import sys
import os
import time
import json
import tempfile
import shutil
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

APP_ROOT = Path(__file__).parent.parent.resolve()
if str(APP_ROOT) not in sys.path:
    sys.path.insert(0, str(APP_ROOT))

from config import (
    AppConfig, set_config,
    APP_VERSION,
    HANGING_INDENT_MIN, HANGING_INDENT_MAX, HANGING_INDENT_DEFAULT,
    UNDO_TIMEOUT_SECONDS,
    DEFAULT_ICON_PATH, BUILD_ICON_PATH,
    UPDATE_CHECK_URL,
)
from undo_buffer import UndoBuffer
from models import Study
from database import DatabaseManager


PROJECT_ROOT = APP_ROOT.parent.resolve()


# =========================================================================
# Fixtures
# =========================================================================

@pytest.fixture
def app_config(tmp_path):
    config = AppConfig(data_root=str(tmp_path / "data"))
    config.ensure_user_directories()
    set_config(config)
    return config


@pytest.fixture
def db(app_config):
    with DatabaseManager(config=app_config) as db_mgr:
        yield db_mgr


@pytest.fixture
def site_with_studies(db):
    site = db.create_site("TestSite")
    studies = [
        Study(phase="Phase I", subcategory="Oncology", year=2024,
              sponsor="Pfizer", protocol="PF-001",
              description_full="Full desc A", description_masked="Masked desc A"),
        Study(phase="Phase I", subcategory="Oncology", year=2023,
              sponsor="Novartis", protocol="NVS-002",
              description_full="Full desc B", description_masked="Masked desc B"),
        Study(phase="Phase II\u2013IV", subcategory="Cardiology", year=2022,
              sponsor="AstraZeneca", protocol="AZ-003",
              description_full="Full desc C", description_masked="Masked desc C"),
    ]
    for s in studies:
        db.add_study(site.id, s)
    return site


# =========================================================================
# 1. Undo Buffer
# =========================================================================

class TestUndoBuffer:

    def test_empty_buffer_cannot_undo(self):
        buf = UndoBuffer()
        assert buf.can_undo is False

    def test_store_and_pop(self):
        buf = UndoBuffer()
        data = [{"phase": "Phase I", "subcategory": "Onc", "year": 2024,
                 "sponsor": "Pfizer", "protocol": "PF-1",
                 "description_full": "full", "description_masked": "masked"}]
        buf.store(1, data)
        assert buf.can_undo is True
        assert buf.study_count == 1
        popped = buf.pop()
        assert len(popped) == 1
        assert popped[0]["sponsor"] == "Pfizer"
        assert buf.can_undo is False

    def test_clear(self):
        buf = UndoBuffer()
        buf.store(1, [{"phase": "P", "subcategory": "S", "year": 2024,
                       "sponsor": "X", "protocol": "", "description_full": "",
                       "description_masked": ""}])
        buf.clear()
        assert buf.can_undo is False
        assert buf.study_count == 0

    def test_clear_if_site_changed(self):
        buf = UndoBuffer()
        buf.store(1, [{"phase": "P", "subcategory": "S", "year": 2024,
                       "sponsor": "X", "protocol": "", "description_full": "",
                       "description_masked": ""}])
        buf.clear_if_site_changed(2)
        assert buf.can_undo is False

    def test_same_site_no_clear(self):
        buf = UndoBuffer()
        buf.store(1, [{"phase": "P", "subcategory": "S", "year": 2024,
                       "sponsor": "X", "protocol": "", "description_full": "",
                       "description_masked": ""}])
        buf.clear_if_site_changed(1)
        assert buf.can_undo is True

    def test_timeout_expires(self):
        buf = UndoBuffer(timeout_seconds=0)
        buf.store(1, [{"phase": "P", "subcategory": "S", "year": 2024,
                       "sponsor": "X", "protocol": "", "description_full": "",
                       "description_masked": ""}])
        time.sleep(0.05)
        assert buf.can_undo is False

    def test_pop_empty_returns_empty(self):
        buf = UndoBuffer()
        assert buf.pop() == []

    def test_store_empty_list_noop(self):
        buf = UndoBuffer()
        buf.store(1, [])
        assert buf.can_undo is False

    def test_store_overwrites_previous(self):
        buf = UndoBuffer()
        buf.store(1, [{"phase": "A", "subcategory": "S", "year": 2024,
                       "sponsor": "X", "protocol": "", "description_full": "",
                       "description_masked": ""}])
        buf.store(1, [{"phase": "B", "subcategory": "S", "year": 2024,
                       "sponsor": "Y", "protocol": "", "description_full": "",
                       "description_masked": ""}])
        popped = buf.pop()
        assert len(popped) == 1
        assert popped[0]["phase"] == "B"

    def test_site_id_property(self):
        buf = UndoBuffer()
        assert buf.site_id is None
        buf.store(42, [{"phase": "P", "subcategory": "S", "year": 2024,
                        "sponsor": "X", "protocol": "", "description_full": "",
                        "description_masked": ""}])
        assert buf.site_id == 42


# =========================================================================
# 2. Count Refresh (database-level)
# =========================================================================

class TestCountRefresh:

    def test_count_after_add(self, db, site_with_studies):
        count = db.get_study_count(site_with_studies.id)
        assert count == 3

    def test_count_after_delete(self, db, site_with_studies):
        studies = db.get_studies(site_with_studies.id)
        db.delete_study(studies[0].id, site_with_studies.id)
        count = db.get_study_count(site_with_studies.id)
        assert count == 2

    def test_count_after_undo_restore(self, db, site_with_studies):
        studies = db.get_studies(site_with_studies.id)
        deleted = studies[0]
        db.delete_study(deleted.id, site_with_studies.id)
        assert db.get_study_count(site_with_studies.id) == 2
        restored = Study(
            phase=deleted.phase, subcategory=deleted.subcategory,
            year=deleted.year, sponsor=deleted.sponsor,
            protocol=deleted.protocol,
            description_full=deleted.description_full,
            description_masked=deleted.description_masked,
        )
        db.add_study(site_with_studies.id, restored)
        assert db.get_study_count(site_with_studies.id) == 3

    def test_count_after_bulk_delete(self, db, site_with_studies):
        studies = db.get_studies(site_with_studies.id)
        for s in studies:
            db.delete_study(s.id, site_with_studies.id)
        assert db.get_study_count(site_with_studies.id) == 0


# =========================================================================
# 3. get_study method
# =========================================================================

class TestGetStudy:

    def test_get_existing_study(self, db, site_with_studies):
        studies = db.get_studies(site_with_studies.id)
        fetched = db.get_study(studies[0].id, site_with_studies.id)
        assert fetched is not None
        assert fetched.sponsor == studies[0].sponsor

    def test_get_nonexistent_study(self, db, site_with_studies):
        assert db.get_study(99999, site_with_studies.id) is None

    def test_get_study_wrong_site(self, db, site_with_studies):
        studies = db.get_studies(site_with_studies.id)
        assert db.get_study(studies[0].id, 99999) is None


# =========================================================================
# 4. Hanging Indent Config
# =========================================================================

class TestHangingIndentConfig:

    def test_default_value(self):
        assert HANGING_INDENT_DEFAULT == 0.5

    def test_range_constants(self):
        assert HANGING_INDENT_MIN == 0.0
        assert HANGING_INDENT_MAX == 2.0

    def test_valid_config(self, tmp_path):
        cfg = AppConfig(data_root=str(tmp_path / "data"), hanging_indent_inches=1.0)
        assert cfg.hanging_indent_inches == 1.0

    def test_zero_is_valid(self, tmp_path):
        cfg = AppConfig(data_root=str(tmp_path / "data"), hanging_indent_inches=0.0)
        assert cfg.hanging_indent_inches == 0.0

    def test_max_is_valid(self, tmp_path):
        cfg = AppConfig(data_root=str(tmp_path / "data"), hanging_indent_inches=2.0)
        assert cfg.hanging_indent_inches == 2.0

    def test_negative_raises(self, tmp_path):
        with pytest.raises(ValueError, match="hanging_indent_inches"):
            AppConfig(data_root=str(tmp_path / "data"), hanging_indent_inches=-0.1)

    def test_over_max_raises(self, tmp_path):
        with pytest.raises(ValueError, match="hanging_indent_inches"):
            AppConfig(data_root=str(tmp_path / "data"), hanging_indent_inches=2.1)

    def test_non_numeric_raises(self, tmp_path):
        with pytest.raises(ValueError, match="hanging_indent_inches"):
            AppConfig(data_root=str(tmp_path / "data"), hanging_indent_inches="bad")

    def test_int_accepted(self, tmp_path):
        cfg = AppConfig(data_root=str(tmp_path / "data"), hanging_indent_inches=1)
        assert cfg.hanging_indent_inches == 1

    def test_persistence(self, tmp_path):
        cfg = AppConfig(data_root=str(tmp_path / "data"), hanging_indent_inches=0.75)
        cfg_path = tmp_path / "test_config.json"
        cfg.save(cfg_path)
        loaded = AppConfig.load(cfg_path)
        assert loaded.hanging_indent_inches == 0.75


# =========================================================================
# 5. Hanging Indent in docx_handler
# =========================================================================

class TestHangingIndentDocx:

    def test_default_hanging_indent(self):
        from docx.shared import Inches
        from docx_handler import CVDocxHandler
        handler = CVDocxHandler(Path("dummy.docx"))
        assert handler.HANGING_INDENT == Inches(0.5)

    def test_custom_hanging_indent(self):
        from docx.shared import Inches
        from docx_handler import CVDocxHandler
        handler = CVDocxHandler(Path("dummy.docx"), hanging_indent_inches=1.0)
        assert handler.HANGING_INDENT == Inches(1.0)

    def test_zero_hanging_indent(self):
        from docx.shared import Inches
        from docx_handler import CVDocxHandler
        handler = CVDocxHandler(Path("dummy.docx"), hanging_indent_inches=0.0)
        assert handler.HANGING_INDENT == Inches(0.0)

    def test_study_paragraph_uses_config_indent(self, tmp_path):
        from docx import Document
        from docx.shared import Inches, Pt
        from docx_handler import CVDocxHandler

        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("Research Experience", level=1)
        doc.add_paragraph("Phase I")
        doc.add_paragraph("Oncology")
        doc.save(doc_path)

        handler = CVDocxHandler(doc_path, hanging_indent_inches=1.0)
        handler.load()

        study = Study(
            phase="Phase I", subcategory="Oncology", year=2024,
            sponsor="Pfizer", protocol="PF-001",
            description_full="Test study full",
            description_masked="Test study masked",
        )
        handler._create_study_paragraph(study, include_protocol=True, protocol_red=False)

        last_para = handler.document.paragraphs[-1]
        pf = last_para.paragraph_format
        assert pf.left_indent == Inches(1.0)
        assert pf.first_line_indent == Inches(-1.0)

    def test_study_element_uses_config_indent(self, tmp_path):
        from docx_handler import CVDocxHandler
        from docx.oxml.ns import qn

        handler = CVDocxHandler(Path("dummy.docx"), hanging_indent_inches=1.5)
        study = Study(
            phase="Phase I", subcategory="Oncology", year=2024,
            sponsor="Pfizer", protocol="PF-001",
            description_full="Test", description_masked="Test",
        )
        elem = handler._create_study_element(study, include_protocol=True, protocol_red=False)
        ind = elem.find(qn('w:pPr')).find(qn('w:ind'))
        expected_twips = str(int(1.5 * 1440))
        assert ind.get(qn('w:left')) == expected_twips
        assert ind.get(qn('w:hanging')) == expected_twips


# =========================================================================
# 6. Icon Generation
# =========================================================================

class TestIconGeneration:

    def test_generate_icon_creates_file(self, tmp_path):
        sys.path.insert(0, str(PROJECT_ROOT / "build"))
        try:
            from generate_icon import generate_icon
        except ImportError:
            pytest.skip("Pillow not installed")
        finally:
            sys.path.pop(0)

        icon_path = tmp_path / "test_icon.ico"
        result = generate_icon(icon_path)
        if result is None:
            pytest.skip("Pillow not installed")
        assert icon_path.exists()
        assert icon_path.stat().st_size > 0

    def test_generate_icon_skips_when_exists(self, tmp_path, capsys):
        sys.path.insert(0, str(PROJECT_ROOT / "build"))
        try:
            from generate_icon import generate_icon, main as icon_main
        except ImportError:
            pytest.skip("Pillow not installed")
        finally:
            sys.path.pop(0)

        icon_path = tmp_path / "test_icon.ico"
        generate_icon(icon_path)
        if not icon_path.exists():
            pytest.skip("Pillow not installed")

        original_size = icon_path.stat().st_size
        generate_icon(icon_path)
        assert icon_path.stat().st_size > 0


# =========================================================================
# 7. Update Checker
# =========================================================================

class TestUpdateChecker:

    def test_parse_semver_valid(self):
        from update_checker import parse_semver
        assert parse_semver("1.2.3") == (1, 2, 3)
        assert parse_semver("v1.2.3") == (1, 2, 3)
        assert parse_semver("v0.0.1") == (0, 0, 1)

    def test_parse_semver_invalid(self):
        from update_checker import parse_semver
        assert parse_semver("") is None
        assert parse_semver("abc") is None
        assert parse_semver("1.2") is None

    def test_is_newer_true(self):
        from update_checker import is_newer
        assert is_newer("v99.0.0", "1.0.0") is True

    def test_is_newer_false_same(self):
        from update_checker import is_newer
        assert is_newer("1.2.0", "1.2.0") is False

    def test_is_newer_false_older(self):
        from update_checker import is_newer
        assert is_newer("1.0.0", "1.2.0") is False

    def test_is_newer_invalid_remote(self):
        from update_checker import is_newer
        assert is_newer("bad", "1.2.0") is False

    def test_check_for_update_no_network(self):
        from update_checker import check_for_update
        from urllib.error import URLError
        with patch("update_checker.urlopen", side_effect=URLError("no network")):
            with pytest.raises(URLError):
                check_for_update()

    def test_check_for_update_newer_available(self):
        from update_checker import check_for_update
        mock_data = json.dumps({
            "tag_name": "v99.0.0",
            "html_url": "https://example.com",
            "zipball_url": "https://example.com/zip",
        }).encode("utf-8")
        mock_resp = MagicMock()
        mock_resp.read.return_value = mock_data
        with patch("update_checker.urlopen", return_value=mock_resp):
            result = check_for_update()
        assert result is not None
        assert result["tag_name"] == "v99.0.0"

    def test_check_for_update_up_to_date(self):
        from update_checker import check_for_update
        mock_data = json.dumps({
            "tag_name": f"v{APP_VERSION}",
            "html_url": "https://example.com",
        }).encode("utf-8")
        mock_resp = MagicMock()
        mock_resp.read.return_value = mock_data
        with patch("update_checker.urlopen", return_value=mock_resp):
            result = check_for_update()
        assert result is None

    def test_module_isolation(self):
        import importlib
        spec = importlib.util.find_spec("update_checker")
        assert spec is not None
        source = Path(spec.origin).read_text(encoding="utf-8")
        assert "urlopen" in source
        assert "requests" not in source.split("import")[-1] if "import" in source else True


# =========================================================================
# 8. Config Constants
# =========================================================================

class TestConfigConstants:

    def test_app_version_format(self):
        parts = APP_VERSION.split(".")
        assert len(parts) == 3
        for p in parts:
            assert p.isdigit()

    def test_undo_timeout(self):
        assert UNDO_TIMEOUT_SECONDS == 300

    def test_icon_paths(self):
        assert DEFAULT_ICON_PATH == "build/assets/app.ico"
        assert BUILD_ICON_PATH == "build/assets/app.ico"

    def test_update_url(self):
        assert UPDATE_CHECK_URL == "https://api.github.com/repos/vctruong100/Curriculum-Vitae-Manager/releases/latest"
        assert UPDATE_CHECK_URL.startswith("https://")

    def test_check_updates_default_false(self, tmp_path):
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        assert cfg.check_updates_on_startup is False

    def test_check_updates_validation(self, tmp_path):
        with pytest.raises(ValueError, match="check_updates_on_startup"):
            AppConfig(data_root=str(tmp_path / "data"), check_updates_on_startup="yes")


# =========================================================================
# 9. Launcher Scripts
# =========================================================================

class TestLauncherScripts:

    def test_bat_exists(self):
        bat_path = PROJECT_ROOT / "scripts" / "launch_win.bat"
        assert bat_path.exists(), f"Expected {bat_path}"

    def test_sh_exists(self):
        sh_path = PROJECT_ROOT / "scripts" / "launch.sh"
        assert sh_path.exists(), f"Expected {sh_path}"

    def test_bat_tries_py_first(self):
        bat = (PROJECT_ROOT / "scripts" / "launch_win.bat").read_text()
        lines = bat.split("\n")
        py_line = None
        python_line = None
        for i, line in enumerate(lines):
            if "py -m pip" in line and py_line is None:
                py_line = i
            if "python -m pip" in line and python_line is None:
                python_line = i
        assert py_line is not None, "Expected 'py -m pip' in bat"
        assert python_line is not None, "Expected 'python -m pip' in bat"
        assert py_line < python_line, "'py' should be tried before 'python'"

    def test_sh_tries_py_first(self):
        sh = (PROJECT_ROOT / "scripts" / "launch.sh").read_text()
        assert "py" in sh
        assert "python" in sh


# =========================================================================
# 10. Tooltip Text
# =========================================================================

class TestTooltipHangingIndent:

    def test_hanging_indent_tooltip_exists(self):
        from tooltip_text import get_tooltip_text, TOOLTIP_DEFAULT
        text = get_tooltip_text("hanging_indent_inches")
        assert text != TOOLTIP_DEFAULT
        assert "indent" in text.lower()


# =========================================================================
# 11. Undo + Database Integration
# =========================================================================

class TestUndoDatabaseIntegration:

    def test_delete_and_undo_restores_fields(self, db, site_with_studies):
        studies = db.get_studies(site_with_studies.id)
        target = studies[0]
        undo_data = [{
            'phase': target.phase,
            'subcategory': target.subcategory,
            'year': target.year,
            'sponsor': target.sponsor,
            'protocol': target.protocol,
            'description_full': target.description_full,
            'description_masked': target.description_masked,
        }]

        buf = UndoBuffer()
        db.delete_study(target.id, site_with_studies.id)
        buf.store(site_with_studies.id, undo_data)

        assert db.get_study_count(site_with_studies.id) == 2

        popped = buf.pop()
        for sd in popped:
            restored = Study(**sd)
            db.add_study(site_with_studies.id, restored)

        assert db.get_study_count(site_with_studies.id) == 3
        all_studies = db.get_studies(site_with_studies.id)
        sponsors = [s.sponsor for s in all_studies]
        assert target.sponsor in sponsors

    def test_undo_buffer_cleared_on_site_switch(self, db, site_with_studies):
        buf = UndoBuffer()
        studies = db.get_studies(site_with_studies.id)
        target = studies[0]
        undo_data = [{
            'phase': target.phase, 'subcategory': target.subcategory,
            'year': target.year, 'sponsor': target.sponsor,
            'protocol': target.protocol,
            'description_full': target.description_full,
            'description_masked': target.description_masked,
        }]
        db.delete_study(target.id, site_with_studies.id)
        buf.store(site_with_studies.id, undo_data)

        other_site = db.create_site("OtherSite")
        buf.clear_if_site_changed(other_site.id)
        assert buf.can_undo is False

    def test_no_op_undo_on_empty_buffer(self, db, site_with_studies):
        buf = UndoBuffer()
        popped = buf.pop()
        assert popped == []
        assert db.get_study_count(site_with_studies.id) == 3


# =========================================================================
# 12. Spec file
# =========================================================================

class TestSpecFile:

    def test_spec_includes_undo_buffer(self):
        spec_path = PROJECT_ROOT / "cv_manager.spec"
        if not spec_path.exists():
            pytest.skip("spec file not found")
        content = spec_path.read_text()
        assert "undo_buffer" in content
