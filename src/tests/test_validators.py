"""
Tests for the validators module.

Covers: strict master xlsx validation, strict CV docx validation,
edge cases (formulas, dates, duplicates, missing columns, formatting).
"""

import sys
from pathlib import Path

import pytest
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from validators import validate_master_xlsx_strict, validate_cv_docx_strict


class TestValidateMasterXlsxStrict:
    def test_valid_file(self, sample_master_xlsx):
        report = validate_master_xlsx_strict(sample_master_xlsx)
        assert report["valid"] is True
        assert report["stats"]["phases"] > 0
        assert report["stats"]["studies"] > 0

    def test_nonexistent_file(self, tmp_dir):
        report = validate_master_xlsx_strict(tmp_dir / "nope.xlsx")
        assert report["valid"] is False

    def test_empty_file(self, empty_master_xlsx):
        report = validate_master_xlsx_strict(empty_master_xlsx)
        assert report["valid"] is False

    def test_no_phase_headings(self, malformed_master_xlsx):
        report = validate_master_xlsx_strict(malformed_master_xlsx)
        # Should have error about missing phase or study before phase
        errors = [i for i in report["issues"] if i["severity"] == "error"]
        assert len(errors) > 0

    def test_duplicate_detection(self, tmp_dir):
        """Two identical study rows should trigger a warning."""
        path = tmp_dir / "dupes.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.cell(row=1, column=1, value="Phase I")
        ws.cell(row=2, column=1, value="Oncology")
        ws.cell(row=3, column=1, value=2024)
        ws.cell(row=3, column=2, value="Pfizer PF-123: Same study")
        ws.cell(row=3, column=3, value="Pfizer: Same study")
        ws.cell(row=4, column=1, value=2024)
        ws.cell(row=4, column=2, value="Pfizer PF-123: Same study")
        ws.cell(row=4, column=3, value="Pfizer: Same study")
        wb.save(path)
        wb.close()

        report = validate_master_xlsx_strict(path)
        warnings = [i for i in report["issues"] if i["severity"] == "warning"]
        dupe_warns = [w for w in warnings if "duplicate" in w["message"].lower()]
        assert len(dupe_warns) >= 1

    def test_empty_column_b(self, tmp_dir):
        """Study row with empty Column B should be an error."""
        path = tmp_dir / "no_col_b.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.cell(row=1, column=1, value="Phase I")
        ws.cell(row=2, column=1, value="Oncology")
        ws.cell(row=3, column=1, value=2024)
        # Column B intentionally empty
        ws.cell(row=3, column=3, value="Pfizer: masked")
        wb.save(path)
        wb.close()

        report = validate_master_xlsx_strict(path)
        errors = [i for i in report["issues"] if i["severity"] == "error"]
        assert any("column b" in e["message"].lower() for e in errors)

    def test_wrong_extension(self, tmp_dir):
        path = tmp_dir / "test.csv"
        path.write_text("data")
        report = validate_master_xlsx_strict(path)
        assert report["valid"] is False

    def test_stats_counting(self, sample_master_xlsx):
        report = validate_master_xlsx_strict(sample_master_xlsx)
        assert report["stats"]["phases"] >= 2
        assert report["stats"]["subcategories"] >= 2
        assert report["stats"]["studies"] >= 5


class TestValidateCvDocxStrict:
    def test_valid_cv(self, sample_cv_docx):
        report = validate_cv_docx_strict(sample_cv_docx)
        assert report["valid"] is True
        assert report["stats"]["study_lines"] >= 2

    def test_no_research_section(self, cv_no_research):
        report = validate_cv_docx_strict(cv_no_research)
        assert report["valid"] is False
        errors = [i for i in report["issues"] if i["severity"] == "error"]
        assert any("research experience" in e["message"].lower() for e in errors)

    def test_nonexistent_file(self, tmp_dir):
        report = validate_cv_docx_strict(tmp_dir / "nope.docx")
        assert report["valid"] is False

    def test_wrong_font_detection(self, tmp_dir):
        """CV with wrong font should produce warnings."""
        doc = Document()
        doc.add_heading("Research Experience", level=1)
        p = doc.add_paragraph()
        run = p.add_run("2024\t")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run_s = p.add_run("Pfizer")
        run_s.bold = True
        run_s.font.name = "Times New Roman"
        run_s.font.size = Pt(12)
        p.add_run(": A study")
        doc.add_heading("Publications", level=1)
        path = tmp_dir / "wrong_font.docx"
        doc.save(path)

        report = validate_cv_docx_strict(path)
        warnings = [i for i in report["issues"] if i["severity"] == "warning"]
        font_warns = [w for w in warnings if "font" in w["message"].lower()]
        assert len(font_warns) >= 1

    def test_stats_populated(self, sample_cv_docx):
        report = validate_cv_docx_strict(sample_cv_docx)
        assert report["stats"]["total_paragraphs"] > 0
        assert report["stats"]["research_exp_paragraphs"] > 0
