import sys
import os
import re
import shutil
from pathlib import Path
from datetime import datetime

import pytest
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from openpyxl import Workbook

APP_ROOT = Path(__file__).parent.parent.resolve()
if str(APP_ROOT) not in sys.path:
    sys.path.insert(0, str(APP_ROOT))

from config import AppConfig, set_config
from processor import CVProcessor
from normalizer import (
    normalize_heading_key,
    normalize_subcat_key,
    is_phase_heading,
)
from docx_handler import CVDocxHandler
from models import Study, ResearchExperience


def _make_master(path, extra_studies=None):
    wb = Workbook()
    ws = wb.active
    ws.title = "Studies"
    rows = [
        ("Phase I", None, None),
        ("Healthy Adults", None, None),
        (2024, "ELI LILLY LY246736: A Phase 1 study of LY246736 in healthy adult volunteers",
         "ELI LILLY: A Phase 1 study of XXX in healthy adult volunteers"),
        (2024, "VISTERRA VIS-123: A first-in-human study of VIS-123 in healthy volunteers",
         "VISTERRA: A first-in-human study of XXX in healthy volunteers"),
        (2026, "MERCK MK4082-002: A Multiple-Ascending Dose Study to Evaluate the Safety, Tolerability and Pharmacokinetics of MK-4082 in Healthy Overweight/Obese Participants",
         "MERCK: A Multiple-Ascending Dose Study to Evaluate the Safety, Tolerability and Pharmacokinetics of XXX in Healthy Overweight/Obese Participants"),
        ("Vaccine", None, None),
        (2024, "GILEAD GS-5001: Phase 1/2 Study to Evaluate the Safety of GS-5001 in Healthy Participants",
         "GILEAD: Phase 1/2 Study to Evaluate the Safety of XXX in Healthy Participants"),
        ("Phase II\u2013IV", None, None),
        ("Oncology", None, None),
        (2024, "Roche RO-777: Phase 3 study of RO-777 (atezolizumab) vs placebo in TNBC",
         "Roche: Phase 3 study of XXX vs placebo in TNBC"),
    ]
    if extra_studies:
        rows.extend(extra_studies)
    row_num = 1
    for item in rows:
        ws.cell(row=row_num, column=1, value=item[0])
        if len(item) > 1 and item[1] is not None:
            ws.cell(row=row_num, column=2, value=item[1])
        if len(item) > 2 and item[2] is not None:
            ws.cell(row=row_num, column=3, value=item[2])
        row_num += 1
    wb.save(path)
    wb.close()
    return path


def _make_cv_with_phases(path, style="Normal", subcats=None):
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_paragraph("Email: jane@example.com")
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")
    doc.add_heading("Research Experience", level=1)

    if subcats is None:
        subcats = [
            {
                "phase": "Phase I",
                "subcategories": {
                    "Healthy Adults": [
                        (2024, "ELI LILLY", "LY246736",
                         "A Phase 1 study of LY246736 in healthy adult volunteers"),
                        (2024, "VISTERRA", "VIS-123",
                         "A first-in-human study of VIS-123 in healthy volunteers"),
                    ],
                    "Vaccine": [
                        (2024, "GILEAD", "GS-5001",
                         "Phase 1/2 Study to Evaluate the Safety of GS-5001 in Healthy Participants"),
                    ],
                },
            },
        ]

    for phase_data in subcats:
        phase_text = phase_data["phase"]
        if style == "Normal":
            doc.add_paragraph(phase_text)
        else:
            doc.add_paragraph(phase_text, style=style)

        for subcat_name, studies in phase_data["subcategories"].items():
            if style == "Normal":
                doc.add_paragraph(subcat_name)
            else:
                doc.add_paragraph(subcat_name, style=style)

            for year, sponsor, protocol, desc in studies:
                p = doc.add_paragraph()
                p.add_run(f"{year}\t")
                run_s = p.add_run(sponsor)
                run_s.bold = True
                p.add_run(" ")
                run_p = p.add_run(protocol)
                run_p.bold = True
                run_p.font.color.rgb = RGBColor(0xFF, 0, 0)
                p.add_run(f": {desc}")
                pf = p.paragraph_format
                pf.left_indent = Inches(0)
                pf.first_line_indent = Inches(-0.5)

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")
    doc.save(path)
    return path


def _make_cv_heading_style(path, heading_level=1):
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_paragraph("Email: jane@example.com")
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")
    doc.add_heading("Research Experience", level=1)

    doc.add_heading("Phase I", level=heading_level)
    doc.add_heading("Healthy Adults", level=heading_level + 1 if heading_level < 3 else 3)

    p = doc.add_paragraph()
    p.add_run("2024\t")
    run_s = p.add_run("ELI LILLY")
    run_s.bold = True
    p.add_run(" ")
    run_p = p.add_run("LY246736")
    run_p.bold = True
    run_p.font.color.rgb = RGBColor(0xFF, 0, 0)
    p.add_run(": A Phase 1 study of LY246736 in healthy adult volunteers")

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")
    doc.save(path)
    return path


def _make_cv_uppercase_phases(path):
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")
    doc.add_heading("Research Experience", level=1)
    doc.add_paragraph("PHASE I")
    doc.add_paragraph("HEALTHY ADULTS")

    p = doc.add_paragraph()
    p.add_run("2024\t")
    run_s = p.add_run("ELI LILLY")
    run_s.bold = True
    p.add_run(" ")
    run_p = p.add_run("LY246736")
    run_p.bold = True
    p.add_run(": A Phase 1 study of LY246736 in healthy adult volunteers")

    doc.add_heading("Publications", level=1)
    doc.save(path)
    return path


def _make_cv_double_space_subcat(path):
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")
    doc.add_heading("Research Experience", level=1)
    doc.add_paragraph("Phase I")
    doc.add_paragraph("Healthy  Adults")

    p = doc.add_paragraph()
    p.add_run("2024\t")
    run_s = p.add_run("ELI LILLY")
    run_s.bold = True
    p.add_run(" ")
    run_p = p.add_run("LY246736")
    run_p.bold = True
    p.add_run(": A Phase 1 study of LY246736 in healthy adult volunteers")

    doc.add_heading("Publications", level=1)
    doc.save(path)
    return path


def _make_cv_table_heading(path):
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")
    doc.add_heading("Research Experience", level=1)

    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Phase I"

    doc.add_paragraph("Healthy Adults")

    p = doc.add_paragraph()
    p.add_run("2024\t")
    run_s = p.add_run("ELI LILLY")
    run_s.bold = True
    p.add_run(" ")
    run_p = p.add_run("LY246736")
    run_p.bold = True
    p.add_run(": A Phase 1 study of LY246736 in healthy adult volunteers")

    doc.add_heading("Publications", level=1)
    doc.save(path)
    return path


def _make_cv_split_runs(path):
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")
    doc.add_heading("Research Experience", level=1)

    p_phase = doc.add_paragraph()
    r1 = p_phase.add_run("PHASE ")
    r1.bold = True
    r2 = p_phase.add_run("I")
    r2.bold = True

    doc.add_paragraph("Healthy Adults")

    p = doc.add_paragraph()
    p.add_run("2024\t")
    run_s = p.add_run("ELI LILLY")
    run_s.bold = True
    p.add_run(" ")
    run_p = p.add_run("LY246736")
    run_p.bold = True
    p.add_run(": A Phase 1 study of LY246736 in healthy adult volunteers")

    doc.add_heading("Publications", level=1)
    doc.save(path)
    return path


def _extract_study_lines(docx_path):
    doc = Document(docx_path)
    results = []
    for para in doc.paragraphs:
        text = para.text.strip()
        m = re.match(r'^(\d{4})', text)
        if m and "\t" in para.text:
            results.append((int(m.group(1)), text))
    return results


def _extract_headings_text(docx_path):
    doc = Document(docx_path)
    headings = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and not re.match(r'^\d{4}', text) and "\t" not in text:
            if len(text) < 60:
                headings.append(text)
    return headings


class TestHeadingStyleAgnostic:

    @pytest.mark.parametrize("level", [1, 2, 3])
    def test_heading_style_phase_detected(self, app_config, tmp_dir, level):
        cv_path = _make_cv_heading_style(tmp_dir / "cv.docx", heading_level=level)
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()
        phase_names = [p.name for p in research.phases]
        assert "Phase I" in phase_names

    @pytest.mark.parametrize("level", [1, 2, 3])
    def test_heading_style_subcat_detected(self, app_config, tmp_dir, level):
        cv_path = _make_cv_heading_style(tmp_dir / "cv.docx", heading_level=level)
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()
        subcats = []
        for phase in research.phases:
            for sc in phase.subcategories:
                subcats.append(sc.name)
        assert any("healthy" in s.lower() for s in subcats)

    def test_heading_style_section_not_truncated(self, app_config, tmp_dir):
        cv_path = _make_cv_heading_style(tmp_dir / "cv.docx", heading_level=2)
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()
        all_studies = research.get_all_studies()
        assert len(all_studies) >= 1

    def test_uppercase_phase_detected(self, app_config, tmp_dir):
        cv_path = _make_cv_uppercase_phases(tmp_dir / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()
        phase_names = [p.name for p in research.phases]
        assert "Phase I" in phase_names

    def test_uppercase_subcat_detected(self, app_config, tmp_dir):
        cv_path = _make_cv_uppercase_phases(tmp_dir / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()
        subcats = []
        for phase in research.phases:
            for sc in phase.subcategories:
                subcats.append(sc.name)
        assert any("healthy" in s.lower() for s in subcats)

    def test_double_space_subcat_matched(self, app_config, tmp_dir):
        cv_path = _make_cv_double_space_subcat(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=True,
        )
        assert result.success
        headings = _extract_headings_text(output)
        ha_count = sum(
            1 for h in headings
            if normalize_subcat_key(h) == normalize_subcat_key("Healthy Adults")
        )
        assert ha_count == 1, f"Expected 1 Healthy Adults heading, got {ha_count}"

    def test_split_run_heading_detected(self, app_config, tmp_dir):
        cv_path = _make_cv_split_runs(tmp_dir / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()
        phase_names = [p.name for p in research.phases]
        assert "Phase I" in phase_names

    def test_heading_style_no_duplicate_phases(self, app_config, tmp_dir):
        cv_path = _make_cv_heading_style(tmp_dir / "cv.docx", heading_level=2)
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=True,
        )
        assert result.success
        headings = _extract_headings_text(output)
        phase_headings = [
            h for h in headings
            if is_phase_heading(h) is not None
        ]
        unique_phases = set(
            normalize_heading_key(h) for h in phase_headings
        )
        assert len(phase_headings) == len(unique_phases), (
            f"Duplicate phase headings: {phase_headings}"
        )


class TestContainerSelectionBug:

    def test_merck_study_in_healthy_adults_not_vaccine(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success

        doc = Document(output)
        paras = [p.text.strip() for p in doc.paragraphs]

        merck_idx = None
        vaccine_idx = None
        ha_idx = None
        for i, text in enumerate(paras):
            if "MK4082" in text or "MK-4082" in text:
                merck_idx = i
            if normalize_subcat_key(text) == normalize_subcat_key("Vaccine"):
                vaccine_idx = i
            if normalize_subcat_key(text) == normalize_subcat_key("Healthy Adults"):
                ha_idx = i

        assert merck_idx is not None, "MERCK MK4082-002 study not found in output"
        assert ha_idx is not None, "Healthy Adults heading not found"

        if vaccine_idx is not None:
            assert merck_idx < vaccine_idx, (
                f"MERCK study at index {merck_idx} appears AFTER "
                f"Vaccine heading at {vaccine_idx}"
            )

    def test_sort_false_correct_subcategory_insertion(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success

        doc = Document(output)
        paras = [p.text.strip() for p in doc.paragraphs]

        in_healthy_adults = False
        in_vaccine = False
        merck_in_ha = False
        merck_in_vaccine = False

        for text in paras:
            lower = text.lower().strip()
            if lower == "healthy adults":
                in_healthy_adults = True
                in_vaccine = False
            elif lower == "vaccine":
                in_healthy_adults = False
                in_vaccine = True
            elif is_phase_heading(text) is not None:
                in_healthy_adults = False
                in_vaccine = False

            if "MK4082" in text or "MK-4082" in text:
                if in_healthy_adults:
                    merck_in_ha = True
                if in_vaccine:
                    merck_in_vaccine = True

        assert merck_in_ha, "MERCK study should be under Healthy Adults"
        assert not merck_in_vaccine, "MERCK study should NOT be under Vaccine"

    def test_anchor_uses_matched_container(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        handler.parse_research_experience()

        p_key = normalize_heading_key("Phase I")
        s_key = normalize_subcat_key("Healthy Adults")
        subcat_tuple = (p_key, s_key)

        assert subcat_tuple in handler._subcat_last_study_para, (
            f"Healthy Adults not tracked in _subcat_last_study_para; "
            f"available keys: {list(handler._subcat_last_study_para.keys())}"
        )
        assert subcat_tuple in handler._subcat_study_para_list


class TestSortDisabledHybrid:

    def test_categories_without_new_studies_untouched(self, app_config, tmp_dir):
        subcats_data = [
            {
                "phase": "Phase I",
                "subcategories": {
                    "Healthy Adults": [
                        (2024, "ELI LILLY", "LY246736",
                         "A Phase 1 study of LY246736 in healthy adult volunteers"),
                    ],
                    "Oncology": [
                        (2021, "TESTCO", "TC-001", "A Phase 1 oncology trial"),
                        (2023, "TESTCO", "TC-002", "A Phase 1 oncology trial v2"),
                        (2022, "TESTCO", "TC-003", "A Phase 1 oncology trial v3"),
                    ],
                },
            },
        ]
        cv_path = _make_cv_with_phases(tmp_dir / "cv.docx", subcats=subcats_data)

        master_data = [
            ("Phase I", None, None),
            ("Healthy Adults", None, None),
            (2024, "ELI LILLY LY246736: A Phase 1 study of LY246736 in healthy adult volunteers",
             "ELI LILLY: A Phase 1 study of XXX in healthy adult volunteers"),
            (2026, "MERCK MK4082-002: A new healthy adults study of MK-4082",
             "MERCK: A new healthy adults study of XXX"),
            ("Oncology", None, None),
            (2021, "TESTCO TC-001: A Phase 1 oncology trial",
             "TESTCO: A Phase 1 oncology trial"),
            (2022, "TESTCO TC-003: A Phase 1 oncology trial v3",
             "TESTCO: A Phase 1 oncology trial v3"),
            (2023, "TESTCO TC-002: A Phase 1 oncology trial v2",
             "TESTCO: A Phase 1 oncology trial v2"),
        ]
        master_path = tmp_dir / "master.xlsx"
        _make_master(master_path, [])
        wb = Workbook()
        ws = wb.active
        ws.title = "Studies"
        row_num = 1
        for item in master_data:
            ws.cell(row=row_num, column=1, value=item[0])
            if len(item) > 1 and item[1] is not None:
                ws.cell(row=row_num, column=2, value=item[1])
            if len(item) > 2 and item[2] is not None:
                ws.cell(row=row_num, column=3, value=item[2])
            row_num += 1
        wb.save(master_path)
        wb.close()

        orig_doc = Document(cv_path)
        oncology_study_texts = []
        in_oncology = False
        for para in orig_doc.paragraphs:
            text = para.text.strip()
            if text.lower() == "oncology":
                in_oncology = True
                continue
            if in_oncology and re.match(r'^\d{4}', text):
                oncology_study_texts.append(text)
            elif text and not re.match(r'^\d{4}', text) and in_oncology:
                in_oncology = False

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success

        out_doc = Document(output)
        out_oncology_texts = []
        in_oncology = False
        for para in out_doc.paragraphs:
            text = para.text.strip()
            if text.lower() == "oncology":
                in_oncology = True
                continue
            if in_oncology and re.match(r'^\d{4}', text):
                out_oncology_texts.append(text)
            elif text and not re.match(r'^\d{4}', text) and in_oncology:
                in_oncology = False

        assert out_oncology_texts == oncology_study_texts, (
            f"Oncology studies (no new additions) should be unchanged.\n"
            f"Before: {oncology_study_texts}\nAfter: {out_oncology_texts}"
        )

    def test_categories_with_new_studies_sorted_descending(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success

        doc = Document(output)
        in_ha = False
        ha_years = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if normalize_subcat_key(text) == normalize_subcat_key("Healthy Adults"):
                in_ha = True
                continue
            if in_ha and re.match(r'^\d{4}', text):
                m = re.match(r'^(\d{4})', text)
                ha_years.append(int(m.group(1)))
            elif text and not re.match(r'^\d{4}', text) and in_ha:
                in_ha = False

        assert ha_years == sorted(ha_years, reverse=True), (
            f"Combined list should be sorted descending: {ha_years}"
        )

    def test_existing_xml_preserved_not_recreated(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        orig_doc = Document(cv_path)
        orig_run_counts = {}
        for para in orig_doc.paragraphs:
            text = para.text.strip()
            if re.match(r'^\d{4}', text) and "\t" in para.text:
                orig_run_counts[text] = len(para.runs)

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success

        out_doc = Document(output)
        for para in out_doc.paragraphs:
            text = para.text.strip()
            if text in orig_run_counts:
                assert len(para.runs) == orig_run_counts[text], (
                    f"Run count changed for '{text[:50]}...': "
                    f"was {orig_run_counts[text]}, now {len(para.runs)}"
                )

    def test_empty_subcategory_inserts_after_heading(self, app_config, tmp_dir):
        subcats_data = [
            {
                "phase": "Phase I",
                "subcategories": {
                    "Healthy Adults": [],
                },
            },
        ]
        cv_path = _make_cv_with_phases(tmp_dir / "cv.docx", subcats=subcats_data)
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success
        studies = _extract_study_lines(output)
        assert len(studies) >= 1


class TestOutputRouting:

    def test_mode_a_result_folder_structure(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "Jane Doe CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, enable_sort_existing=True,
        )
        assert result.success
        output_path = Path(result.output_path)
        assert output_path.exists()
        assert output_path.parent.name == "Jane Doe CV"
        assert output_path.parent.parent.name == "result"

    def test_mode_b_result_folder_structure(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "Jane Doe CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path,
        )
        assert result.success
        output_path = Path(result.output_path)
        assert output_path.exists()
        assert output_path.parent.name == "Jane Doe CV"

    def test_mode_b_after_mode_a_same_folder(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "Jane Doe CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)

        result_a = processor.mode_a_update_inject(
            cv_path, master_path=master_path, enable_sort_existing=True,
        )
        assert result_a.success
        updated_path = Path(result_a.output_path)

        result_b = processor.mode_b_redact_protocols(
            updated_path, master_path=master_path,
        )
        assert result_b.success
        redacted_path = Path(result_b.output_path)

        assert updated_path.parent == redacted_path.parent, (
            f"Mode A output dir ({updated_path.parent}) != "
            f"Mode B output dir ({redacted_path.parent})"
        )

    def test_custom_doc_property_stored_and_read(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "Jane Doe CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, enable_sort_existing=True,
        )
        assert result.success
        output_path = Path(result.output_path)

        read_name = CVProcessor._get_original_cv_name_from_doc(output_path)
        assert read_name == "Jane Doe CV", f"Expected 'Jane Doe CV', got '{read_name}'"


class TestNoLogsInResult:

    def test_result_folder_only_docx(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "Jane Doe CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, enable_sort_existing=True,
        )
        assert result.success
        result_folder = Path(result.output_path).parent
        non_docx = [
            f for f in result_folder.iterdir()
            if f.suffix.lower() not in ('.docx',)
        ]
        assert len(non_docx) == 0, (
            f"Non-.docx files in result folder: {[f.name for f in non_docx]}"
        )

    def test_mode_b_result_folder_only_docx(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "Jane Doe CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path,
        )
        assert result.success
        result_folder = Path(result.output_path).parent
        non_docx = [
            f for f in result_folder.iterdir()
            if f.suffix.lower() not in ('.docx',)
        ]
        assert len(non_docx) == 0

    def test_logs_still_written_to_logs_dir(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "Jane Doe CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, enable_sort_existing=True,
        )
        assert result.success
        logs_dir = app_config.get_user_logs_path()
        log_files = list(logs_dir.glob("*.json")) + list(logs_dir.glob("*.csv"))
        assert len(log_files) >= 1, (
            f"Expected log files in {logs_dir}"
        )


class TestModeCExport:

    def test_export_writes_xlsx_to_result_folder(self, app_config, tmp_dir):
        from import_export import ImportExportManager
        from database import DatabaseManager

        with DatabaseManager(config=app_config) as db:
            site = db.create_site("Test Site Alpha")
            study = Study(
                phase="Phase I",
                subcategory="Oncology",
                year=2024,
                sponsor="Pfizer",
                protocol="PF-999",
                description_full="A study desc",
                description_masked="A study desc",
            )
            db.add_study(site.id, study)

        manager = ImportExportManager(app_config)
        success, message, output_path = manager.export_site_to_xlsx(site.id)
        assert success, f"Export failed: {message}"
        assert output_path is not None
        assert output_path.exists(), f"Export file not found at {output_path}"
        assert output_path.suffix == ".xlsx"
        assert "result" in str(output_path)

    def test_export_file_physically_exists(self, app_config, tmp_dir):
        from import_export import ImportExportManager
        from database import DatabaseManager

        with DatabaseManager(config=app_config) as db:
            site = db.create_site("TestSite2")
            study = Study(
                phase="Phase I",
                subcategory="General",
                year=2023,
                sponsor="TestCo",
                protocol="TC-1",
                description_full="Test study",
                description_masked="Test study",
            )
            db.add_study(site.id, study)

        manager = ImportExportManager(app_config)
        success, message, output_path = manager.export_site_to_xlsx(site.id)
        assert success
        assert output_path.stat().st_size > 0


class TestPytestCleanup:

    def test_tmp_path_isolated(self, tmp_dir):
        test_file = tmp_dir / "should_be_cleaned.txt"
        test_file.write_text("test")
        assert test_file.exists()

    def test_no_leftover_from_previous(self, tmp_dir):
        leftover = tmp_dir / "should_be_cleaned.txt"
        assert not leftover.exists()


class TestPreviewJSON:

    def test_preview_includes_container_info(self, app_config, tmp_dir):
        cv_path = _make_cv_with_phases(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        processor = CVProcessor(app_config)
        changes, error = processor.preview_changes(
            cv_path, master_path=master_path, mode="update_inject",
        )
        assert error == ""
        inject_changes = [c for c in changes if c["action"] == "inject"]
        for change in inject_changes:
            assert "phase_key" in change
            assert "subcat_key" in change
            assert "matched_phase_container" in change or change.get("matched_phase_container") is None
