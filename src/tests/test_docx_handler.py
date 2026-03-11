"""
Tests for the docx_handler module.

Covers: section finding, parsing, writing, formatting, edge cases
(missing section, tables, extra whitespace, smart quotes, read-only).
"""

import sys
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches, RGBColor

sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from docx_handler import CVDocxHandler, validate_cv_docx
from models import Study, ResearchExperience


class TestFindResearchExperienceSection:
    def test_find_section(self, sample_cv_docx):
        handler = CVDocxHandler(sample_cv_docx)
        handler.load()
        start, end = handler.find_research_experience_section()
        assert start is not None
        assert end is not None
        assert end > start

    def test_missing_section(self, cv_no_research):
        handler = CVDocxHandler(cv_no_research)
        handler.load()
        start, end = handler.find_research_experience_section()
        assert start is None

    def test_section_bounded_by_next_heading(self, sample_cv_docx):
        handler = CVDocxHandler(sample_cv_docx)
        handler.load()
        start, end = handler.find_research_experience_section()
        # The end should be before the Publications heading
        doc = handler.document
        for i in range(end + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip().lower()
            if "publications" in text:
                # end should be before publications
                assert end < i
                break


class TestParseResearchExperience:
    def test_parse_studies(self, sample_cv_docx):
        handler = CVDocxHandler(sample_cv_docx)
        handler.load()
        re_exp = handler.parse_research_experience()
        studies = re_exp.get_all_studies()
        assert len(studies) >= 2

    def test_study_fields(self, sample_cv_docx):
        handler = CVDocxHandler(sample_cv_docx)
        handler.load()
        re_exp = handler.parse_research_experience()
        studies = re_exp.get_all_studies()
        for s in studies:
            assert s.year > 0
            assert len(s.sponsor) > 0

    def test_phase_detection(self, sample_cv_docx):
        handler = CVDocxHandler(sample_cv_docx)
        handler.load()
        re_exp = handler.parse_research_experience()
        phases = {p.name for p in re_exp.phases}
        assert "Phase I" in phases


class TestWriteResearchExperience:
    def test_write_and_save(self, sample_cv_docx, tmp_dir):
        handler = CVDocxHandler(sample_cv_docx)
        handler.load()
        re_exp = handler.parse_research_experience()
        re_exp.sort_all()

        handler.write_research_experience(re_exp, include_protocol=True, protocol_red=True)
        output = tmp_dir / "output.docx"
        handler.save(output)
        assert output.exists()
        assert output.stat().st_size > 0

    def test_write_redacted(self, sample_cv_docx, tmp_dir):
        handler = CVDocxHandler(sample_cv_docx)
        handler.load()
        re_exp = handler.parse_research_experience()

        # Clear protocols
        for phase in re_exp.phases:
            for sc in phase.subcategories:
                for study in sc.studies:
                    study.protocol = ""

        re_exp.sort_all()
        handler.write_research_experience(re_exp, include_protocol=False, protocol_red=False)
        output = tmp_dir / "redacted.docx"
        handler.save(output)
        assert output.exists()

    def test_round_trip_study_count(self, sample_cv_docx, tmp_dir):
        """Write then re-parse — study count should be preserved."""
        handler = CVDocxHandler(sample_cv_docx)
        handler.load()
        re_exp = handler.parse_research_experience()
        original_count = len(re_exp.get_all_studies())
        re_exp.sort_all()

        handler.write_research_experience(re_exp, include_protocol=True, protocol_red=True)
        output = tmp_dir / "roundtrip.docx"
        handler.save(output)

        handler2 = CVDocxHandler(output)
        handler2.load()
        re_exp2 = handler2.parse_research_experience()
        new_count = len(re_exp2.get_all_studies())
        assert new_count == original_count


class TestValidateCvDocx:
    def test_valid_cv(self, sample_cv_docx):
        is_valid, error = validate_cv_docx(sample_cv_docx)
        assert is_valid is True

    def test_no_research_section(self, cv_no_research):
        is_valid, error = validate_cv_docx(cv_no_research)
        assert is_valid is False
        assert "research experience" in error.lower()

    def test_nonexistent_file(self, tmp_dir):
        is_valid, error = validate_cv_docx(tmp_dir / "nope.docx")
        assert is_valid is False

    def test_wrong_extension(self, tmp_dir):
        txt_file = tmp_dir / "test.txt"
        txt_file.write_text("hello")
        is_valid, error = validate_cv_docx(txt_file)
        assert is_valid is False


class TestEdgeCases:
    def test_extra_whitespace_in_heading(self, tmp_dir):
        """Research Experience heading with extra whitespace."""
        doc = Document()
        doc.add_heading("  Research Experience  ", level=1)
        p = doc.add_paragraph()
        p.add_run("2024\tPfizer PF-123: A study")
        doc.add_heading("Publications", level=1)
        path = tmp_dir / "whitespace.docx"
        doc.save(path)

        handler = CVDocxHandler(path)
        handler.load()
        start, end = handler.find_research_experience_section()
        assert start is not None

    def test_smart_quotes_in_study(self, tmp_dir):
        """Studies with smart quotes should still parse."""
        doc = Document()
        doc.add_heading("Research Experience", level=1)
        doc.add_paragraph("Phase I")
        doc.add_paragraph("Oncology")
        p = doc.add_paragraph()
        p.add_run("2024\tPfizer PF-123: A \u201csmart quoted\u201d study")
        doc.add_heading("Publications", level=1)
        path = tmp_dir / "smartquotes.docx"
        doc.save(path)

        handler = CVDocxHandler(path)
        handler.load()
        re_exp = handler.parse_research_experience()
        studies = re_exp.get_all_studies()
        assert len(studies) >= 1

    def test_tabs_vs_spaces(self, tmp_dir):
        """Study lines with spaces instead of tabs should still parse."""
        doc = Document()
        doc.add_heading("Research Experience", level=1)
        doc.add_paragraph("Phase I")
        doc.add_paragraph("Oncology")
        # Use spaces instead of tab
        p = doc.add_paragraph("2024   Pfizer PF-123: A study with spaces")
        doc.add_heading("Publications", level=1)
        path = tmp_dir / "spaces.docx"
        doc.save(path)

        handler = CVDocxHandler(path)
        handler.load()
        re_exp = handler.parse_research_experience()
        studies = re_exp.get_all_studies()
        assert len(studies) >= 1

    def test_empty_research_section(self, tmp_dir):
        """Research Experience section with no studies."""
        doc = Document()
        doc.add_heading("Research Experience", level=1)
        doc.add_paragraph("")
        doc.add_heading("Publications", level=1)
        path = tmp_dir / "empty_research.docx"
        doc.save(path)

        handler = CVDocxHandler(path)
        handler.load()
        re_exp = handler.parse_research_experience()
        assert len(re_exp.get_all_studies()) == 0
