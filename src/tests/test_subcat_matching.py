"""
Tests for subcategory matching fixes:

1. Slash-space normalization in normalize_subcat_key / normalize_heading_key
2. Combined "Phase I Endocrinology / Metabolic" heading parsing
3. Fuzzy matching fallback in get_or_create_subcategory
"""

import sys
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, Inches, RGBColor

# Ensure app root is importable
APP_ROOT = Path(__file__).parent.parent.resolve()
if str(APP_ROOT) not in sys.path:
    sys.path.insert(0, str(APP_ROOT))

from normalizer import normalize_subcat_key, normalize_heading_key
from models import Phase, Subcategory, ResearchExperience, Study
from docx_handler import CVDocxHandler


# =========================================================================
# 1. Slash-space normalization
# =========================================================================

class TestSlashNormalization:
    """normalize_subcat_key and normalize_heading_key must collapse
    spaces around '/' so that 'A / B' and 'A/B' produce the same key."""

    def test_subcat_slash_with_spaces(self):
        assert normalize_subcat_key("Endocrinology / Metabolic") == \
               normalize_subcat_key("Endocrinology/Metabolic")

    def test_subcat_slash_left_space(self):
        assert normalize_subcat_key("Endocrinology /Metabolic") == \
               normalize_subcat_key("Endocrinology/Metabolic")

    def test_subcat_slash_right_space(self):
        assert normalize_subcat_key("Endocrinology/ Metabolic") == \
               normalize_subcat_key("Endocrinology/Metabolic")

    def test_subcat_multiple_slashes(self):
        assert normalize_subcat_key("A / B / C") == \
               normalize_subcat_key("A/B/C")

    def test_heading_slash_with_spaces(self):
        assert normalize_heading_key("Endocrinology / Metabolic") == \
               normalize_heading_key("Endocrinology/Metabolic")

    def test_slash_normalization_value(self):
        """Verify the normalized key actually has no spaces around '/'."""
        key = normalize_subcat_key("Endocrinology / Metabolic")
        assert "/" in key
        assert " / " not in key
        assert " /" not in key
        assert "/ " not in key

    def test_case_and_slash_combined(self):
        assert normalize_subcat_key("ENDOCRINOLOGY / METABOLIC") == \
               normalize_subcat_key("endocrinology/metabolic")


# =========================================================================
# 2. Combined phase + subcategory heading
# =========================================================================

class TestSplitPhaseSubcatHeading:
    """_split_phase_subcat_heading must extract both phase and subcategory
    from a combined heading paragraph."""

    def test_phase_i_with_subcat(self):
        phase, subcat = CVDocxHandler._split_phase_subcat_heading(
            "Phase I Endocrinology / Metabolic"
        )
        assert phase == "Phase I"
        assert subcat == "Endocrinology / Metabolic"

    def test_phase_i_standalone(self):
        phase, subcat = CVDocxHandler._split_phase_subcat_heading("Phase I")
        assert phase == "Phase I"
        assert subcat is None

    def test_phase_ii_iv_with_subcat(self):
        phase, subcat = CVDocxHandler._split_phase_subcat_heading(
            "Phase II-IV Oncology"
        )
        assert phase == "Phase II\u2013IV"
        assert subcat == "Oncology"

    def test_phase_iii_with_subcat(self):
        phase, subcat = CVDocxHandler._split_phase_subcat_heading(
            "Phase III Healthy Adults"
        )
        assert phase == "Phase III"
        assert subcat == "Healthy Adults"

    def test_phase_1_numeric_with_subcat(self):
        phase, subcat = CVDocxHandler._split_phase_subcat_heading(
            "Phase 1 Cardiology"
        )
        assert phase == "Phase I"
        assert subcat == "Cardiology"

    def test_not_a_phase_heading(self):
        phase, subcat = CVDocxHandler._split_phase_subcat_heading(
            "Endocrinology / Metabolic"
        )
        assert phase is None
        assert subcat is None

    def test_phase_ii_standalone(self):
        phase, subcat = CVDocxHandler._split_phase_subcat_heading("Phase II")
        assert phase == "Phase II"
        assert subcat is None

    def test_case_insensitive(self):
        phase, subcat = CVDocxHandler._split_phase_subcat_heading(
            "PHASE I Endocrinology"
        )
        assert phase == "Phase I"
        assert subcat == "Endocrinology"

    def test_phase_iv_with_subcat(self):
        phase, subcat = CVDocxHandler._split_phase_subcat_heading(
            "Phase IV Neurology"
        )
        assert phase == "Phase IV"
        assert subcat == "Neurology"


# =========================================================================
# 3. Fuzzy matching in get_or_create_subcategory
# =========================================================================

class TestFuzzySubcatMatching:
    """get_or_create_subcategory must fall back to fuzzy matching when
    exact normalized key match fails."""

    def test_exact_match_still_works(self):
        phase = Phase(name="Phase I")
        sc1 = phase.get_or_create_subcategory("Oncology")
        sc2 = phase.get_or_create_subcategory("Oncology")
        assert sc1 is sc2
        assert len(phase.subcategories) == 1

    def test_slash_spacing_exact_after_normalization(self):
        """With the slash normalization, these now match exactly."""
        phase = Phase(name="Phase I")
        sc1 = phase.get_or_create_subcategory("Endocrinology / Metabolic")
        sc2 = phase.get_or_create_subcategory("Endocrinology/Metabolic")
        assert sc1 is sc2
        assert len(phase.subcategories) == 1

    def test_fuzzy_match_plural(self):
        """'Infectious Disease' vs 'Infectious Diseases' should fuzzy match."""
        phase = Phase(name="Phase I")
        sc1 = phase.get_or_create_subcategory("Infectious Disease")
        sc2 = phase.get_or_create_subcategory("Infectious Diseases")
        assert sc1 is sc2
        assert len(phase.subcategories) == 1

    def test_fuzzy_no_false_positive(self):
        """Completely different subcategories must NOT match."""
        phase = Phase(name="Phase I")
        sc1 = phase.get_or_create_subcategory("Oncology")
        sc2 = phase.get_or_create_subcategory("Neurology")
        assert sc1 is not sc2
        assert len(phase.subcategories) == 2

    def test_fuzzy_disabled(self):
        """With fuzzy_threshold=0, only exact normalized match works."""
        phase = Phase(name="Phase I")
        sc1 = phase.get_or_create_subcategory(
            "Infectious Disease", fuzzy_threshold=0
        )
        sc2 = phase.get_or_create_subcategory(
            "Infectious Diseases", fuzzy_threshold=0
        )
        assert sc1 is not sc2
        assert len(phase.subcategories) == 2

    def test_preserves_first_name(self):
        """The display name from the first creation is preserved."""
        phase = Phase(name="Phase I")
        phase.get_or_create_subcategory("Endocrinology / Metabolic")
        sc = phase.get_or_create_subcategory("Endocrinology/Metabolic")
        assert sc.name == "Endocrinology / Metabolic"

    def test_fuzzy_case_variants_already_handled(self):
        """Case variants are handled by normalization, not fuzzy."""
        phase = Phase(name="Phase I")
        sc1 = phase.get_or_create_subcategory("healthy adults")
        sc2 = phase.get_or_create_subcategory("HEALTHY ADULTS")
        assert sc1 is sc2


# =========================================================================
# 4. Parser integration — combined heading in .docx
# =========================================================================

class TestParserCombinedHeading:
    """Verify the parser correctly handles combined phase+subcat headings."""

    def _make_cv_with_combined_heading(self, tmp_path, heading_text, studies):
        """Create a CV .docx with a combined heading and study lines."""
        doc = Document()
        doc.add_heading("Test Subject", level=0)
        doc.add_heading("Research Experience", level=1)

        # Add the combined heading (bold + italic)
        p = doc.add_paragraph()
        run = p.add_run(heading_text)
        run.bold = True
        run.italic = True

        # Add study lines
        for year, sponsor, protocol, desc in studies:
            sp = doc.add_paragraph()
            sp.add_run(f"{year}\t")
            r_sponsor = sp.add_run(sponsor)
            r_sponsor.bold = True
            if protocol:
                sp.add_run(" ")
                r_proto = sp.add_run(protocol)
                r_proto.bold = True
                r_proto.font.color.rgb = RGBColor(0xFF, 0, 0)
            sp.add_run(f": {desc}")

        doc.add_heading("Publications", level=1)
        doc.add_paragraph("1. Example reference")

        path = tmp_path / "cv_combined.docx"
        doc.save(path)
        return path

    def test_combined_heading_parsed_correctly(self, tmp_path):
        """A 'Phase I Endocrinology / Metabolic' heading should set both
        phase and subcategory."""
        path = self._make_cv_with_combined_heading(
            tmp_path,
            "Phase I Endocrinology / Metabolic",
            [
                (2024, "Acme", "AC-001",
                 "A study of AC-001 in type 2 diabetes"),
            ],
        )
        handler = CVDocxHandler(path)
        handler.load()
        handler.find_research_experience_section()
        re_exp = handler.parse_research_experience()

        assert len(re_exp.phases) == 1
        phase = re_exp.phases[0]
        assert phase.name == "Phase I"
        assert len(phase.subcategories) == 1
        # normalize_for_display preserves original spacing for display
        assert "Endocrinology" in phase.subcategories[0].name
        assert "Metabolic" in phase.subcategories[0].name
        assert len(phase.subcategories[0].studies) == 1

    def test_standalone_italic_subcat_parsed(self, tmp_path):
        """A standalone italic subcategory (no phase prefix) after a
        Phase heading should still be recognized."""
        doc = Document()
        doc.add_heading("Test Subject", level=0)
        doc.add_heading("Research Experience", level=1)
        doc.add_paragraph("Phase I")

        # Italic-only subcategory heading
        p = doc.add_paragraph()
        run = p.add_run("Endocrinology / Metabolic")
        run.bold = True
        run.italic = True

        sp = doc.add_paragraph()
        sp.add_run("2024\t")
        r_s = sp.add_run("Acme")
        r_s.bold = True
        sp.add_run(" ")
        r_p = sp.add_run("AC-001")
        r_p.bold = True
        sp.add_run(": A study of AC-001 in type 2 diabetes")

        doc.add_heading("Publications", level=1)
        doc.add_paragraph("1. Example reference")

        path = tmp_path / "cv_italic_subcat.docx"
        doc.save(path)

        handler = CVDocxHandler(path)
        handler.load()
        handler.find_research_experience_section()
        re_exp = handler.parse_research_experience()

        assert len(re_exp.phases) == 1
        phase = re_exp.phases[0]
        assert phase.name == "Phase I"
        assert len(phase.subcategories) == 1
        # normalize_for_display normalizes slashes
        subcat_name = phase.subcategories[0].name
        assert "Endocrinology" in subcat_name
        assert "Metabolic" in subcat_name
        assert len(phase.subcategories[0].studies) == 1

    def test_master_study_merges_into_parsed_combined_heading(self, tmp_path):
        """A master study with subcategory 'Endocrinology/Metabolic'
        should merge into a CV-parsed 'Endocrinology / Metabolic'."""
        # Build a model from the CV parsing
        re_exp = ResearchExperience()
        phase = re_exp.get_or_create_phase("Phase I")
        # Simulates what the parser yields from "Phase I Endocrinology / Metabolic"
        sc = phase.get_or_create_subcategory("Endocrinology / Metabolic")
        sc.studies.append(Study(
            phase="Phase I",
            subcategory="Endocrinology / Metabolic",
            year=2023,
            sponsor="Acme",
            protocol="AC-001",
            description_full="Study of AC-001",
            description_masked="Study of XXX",
        ))

        # Now a master study arrives with "Endocrinology/Metabolic" (no spaces)
        master_subcat_name = "Endocrinology/Metabolic"
        sc2 = phase.get_or_create_subcategory(master_subcat_name)

        # Must be the SAME subcategory
        assert sc2 is sc
        assert len(phase.subcategories) == 1
        assert sc2.name == "Endocrinology / Metabolic"  # first-seen preserved


# =========================================================================
# 5. End-to-end injection test
# =========================================================================

class TestInjectionSlashSubcat:
    """End-to-end: CV with 'Phase I Endocrinology / Metabolic' heading
    receives a new study from master with subcategory
    'Endocrinology/Metabolic'. The study must go under the existing
    subcategory, NOT create a duplicate."""

    @staticmethod
    def _make_master_xlsx_legacy(path, studies_data):
        """Create a legacy 3-column master .xlsx.

        *studies_data* is a list of tuples.  Phase/subcategory rows have
        a string in column A only.  Study rows have (year, col_b, col_c).
        """
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Studies"
        row_num = 1
        for item in studies_data:
            ws.cell(row=row_num, column=1, value=item[0])
            if len(item) > 1 and item[1] is not None:
                ws.cell(row=row_num, column=2, value=item[1])
            if len(item) > 2 and item[2] is not None:
                ws.cell(row=row_num, column=3, value=item[2])
            row_num += 1
        wb.save(path)
        wb.close()
        return path

    def test_injection_no_duplicate_subcat(self, tmp_path, app_config):
        # 1. Create CV with combined heading
        doc = Document()
        doc.add_heading("Jane Doe", level=0)
        doc.add_heading("Research Experience", level=1)

        # Combined heading: bold+italic "Phase I Endocrinology / Metabolic"
        p = doc.add_paragraph()
        run = p.add_run("Phase I Endocrinology / Metabolic")
        run.bold = True
        run.italic = True

        sp = doc.add_paragraph()
        sp.add_run("2023\t")
        rs = sp.add_run("Acme")
        rs.bold = True
        sp.add_run(" ")
        rp = sp.add_run("AC-001")
        rp.bold = True
        rp.font.color.rgb = RGBColor(0xFF, 0, 0)
        sp.add_run(": A Phase 1 study of AC-001 in type 2 diabetes")

        doc.add_heading("Publications", level=1)
        doc.add_paragraph("1. Example")

        cv_path = tmp_path / "cv_combined_e2e.docx"
        doc.save(cv_path)

        # 2. Create master (legacy 3-col) with same study + a NEW study
        master_path = tmp_path / "master_e2e.xlsx"
        self._make_master_xlsx_legacy(master_path, studies_data=[
            ("Phase I", None, None),
            ("Endocrinology/Metabolic", None, None),
            # Existing study (should match CV study)
            (2023,
             "Acme AC-001: A Phase 1 study of AC-001 in type 2 diabetes",
             "Acme: A Phase 1 study of XXX in type 2 diabetes"),
            # New study to inject
            (2024,
             "Beta BT-002: A study of BT-002 in metabolic syndrome",
             "Beta: A study of XXX in metabolic syndrome"),
        ])

        # 3. Run processor
        from processor import CVProcessor
        processor = CVProcessor(config=app_config)
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=tmp_path / "output.docx",
        )

        assert result.success, result.error_message

        # 4. Parse output and verify single subcategory
        out_handler = CVDocxHandler(Path(result.output_path))
        out_handler.load()
        out_handler.find_research_experience_section()
        out_re = out_handler.parse_research_experience()

        # Collect all subcategory names under Phase I
        phase_i_subcats = []
        for phase in out_re.phases:
            if "Phase I" in phase.name or "phase i" in phase.name.lower():
                for sc in phase.subcategories:
                    phase_i_subcats.append(sc.name)

        # There must be exactly ONE subcategory containing "Endocrinology"
        endo_subcats = [
            n for n in phase_i_subcats
            if "endocrinology" in n.lower()
        ]
        assert len(endo_subcats) == 1, (
            f"Expected 1 Endocrinology subcategory, got {len(endo_subcats)}: "
            f"{endo_subcats}"
        )
