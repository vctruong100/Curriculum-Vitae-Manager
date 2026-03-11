"""
Shared fixtures for the CV Research Experience Manager test suite.

All fixtures generate synthetic data on the fly — no external files needed.
Tests are hermetic: each test gets its own temp directory.
"""

import sys
import os
import tempfile
import shutil
from pathlib import Path

import pytest

# Ensure app root is importable
APP_ROOT = Path(__file__).parent.parent.resolve()
if str(APP_ROOT) not in sys.path:
    sys.path.insert(0, str(APP_ROOT))

from openpyxl import Workbook
from docx import Document
from docx.shared import Pt, Inches, RGBColor

from config import AppConfig, set_config
from models import Study, ResearchExperience, Phase, Subcategory


# ---------------------------------------------------------------------------
# Temp directory fixture
# ---------------------------------------------------------------------------

@pytest.fixture
def tmp_dir(tmp_path):
    """Provide a clean temporary directory (pytest built-in tmp_path)."""
    return tmp_path


@pytest.fixture
def app_config(tmp_dir):
    """Provide an AppConfig rooted in a temp directory."""
    config = AppConfig(data_root=str(tmp_dir / "data"))
    config.ensure_user_directories()
    set_config(config)
    return config


# ---------------------------------------------------------------------------
# Sample master .xlsx
# ---------------------------------------------------------------------------

def _make_master_xlsx(path: Path, studies_data=None):
    """
    Create a master .xlsx file at *path* with the given study data.
    Default data has 2 phases, 3 subcategories, and 6 studies.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Studies"

    if studies_data is None:
        studies_data = [
            ("Phase I", None, None),
            ("Oncology", None, None),
            (2024, "Pfizer PF-99999: A Phase 1 study of PF-99999 (pembrolizumab) in advanced lung cancer",
                   "Pfizer: A Phase 1 study of XXX in advanced lung cancer"),
            (2023, "Novartis NVS-789: First-in-human study of NVS-789 for metastatic breast cancer",
                   "Novartis: First-in-human study of XXX for metastatic breast cancer"),
            ("Cardiology", None, None),
            (2024, "AstraZeneca AZ-111: Phase 1 trial of AZ-111 (dapagliflozin) in heart failure",
                   "AstraZeneca: Phase 1 trial of XXX in heart failure"),
            ("Phase II\u2013IV", None, None),
            ("Oncology", None, None),
            (2024, "Roche RO-777: Phase 3 study of RO-777 (atezolizumab) vs placebo in TNBC",
                   "Roche: Phase 3 study of XXX vs placebo in TNBC"),
            (2023, "Roche RO-555: Phase 3 randomized trial of RO-555 vs standard of care in NSCLC",
                   "Roche: Phase 3 randomized trial of XXX vs standard of care in NSCLC"),
            (2020, "BMS BMS-222: Phase 2 study of BMS-222 immunotherapy in melanoma",
                   "BMS: Phase 2 study of XXX immunotherapy in melanoma"),
        ]

    row = 1
    for item in studies_data:
        ws.cell(row=row, column=1, value=item[0])
        if len(item) > 1 and item[1] is not None:
            ws.cell(row=row, column=2, value=item[1])
        if len(item) > 2 and item[2] is not None:
            ws.cell(row=row, column=3, value=item[2])
        row += 1

    wb.save(path)
    wb.close()
    return path


@pytest.fixture
def sample_master_xlsx(tmp_dir):
    """Create and return a path to a sample master .xlsx file."""
    path = tmp_dir / "master.xlsx"
    return _make_master_xlsx(path)


@pytest.fixture
def empty_master_xlsx(tmp_dir):
    """Create a master xlsx with no data."""
    path = tmp_dir / "empty_master.xlsx"
    wb = Workbook()
    wb.save(path)
    wb.close()
    return path


@pytest.fixture
def malformed_master_xlsx(tmp_dir):
    """Create a master xlsx with studies but no phase headings."""
    path = tmp_dir / "malformed_master.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.cell(row=1, column=1, value=2024)
    ws.cell(row=1, column=2, value="Pfizer PF-123: Some study")
    ws.cell(row=1, column=3, value="Pfizer: Some study")
    wb.save(path)
    wb.close()
    return path


# ---------------------------------------------------------------------------
# Sample CV .docx
# ---------------------------------------------------------------------------

def _make_cv_docx(path: Path, include_research_exp=True, studies=None):
    """
    Create a synthetic CV .docx at *path*.
    """
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_paragraph("Email: jane.doe@example.com")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")

    if include_research_exp:
        doc.add_heading("Research Experience", level=1)

        if studies is None:
            # Default: one phase, two studies
            doc.add_paragraph("Phase I")
            doc.add_paragraph("Oncology")

            p1 = doc.add_paragraph()
            p1.add_run("2023\t")
            run_s = p1.add_run("Pfizer")
            run_s.bold = True
            p1.add_run(" ")
            run_p = p1.add_run("PF-12345")
            run_p.bold = True
            run_p.font.color.rgb = RGBColor(0xFF, 0, 0)
            p1.add_run(": A Phase 1 study of PF-12345 in patients with advanced solid tumors")
            pf = p1.paragraph_format
            pf.left_indent = Inches(0)
            pf.first_line_indent = Inches(-0.5)

            p2 = doc.add_paragraph()
            p2.add_run("2022\t")
            run_s2 = p2.add_run("Novartis")
            run_s2.bold = True
            p2.add_run(" ")
            run_p2 = p2.add_run("NVS-789")
            run_p2.bold = True
            run_p2.font.color.rgb = RGBColor(0xFF, 0, 0)
            p2.add_run(": First-in-human study of NVS-789 for metastatic breast cancer")
            pf2 = p2.paragraph_format
            pf2.left_indent = Inches(0)
            pf2.first_line_indent = Inches(-0.5)
        else:
            for text in studies:
                doc.add_paragraph(text)

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")

    doc.save(path)
    return path


@pytest.fixture
def sample_cv_docx(tmp_dir):
    """Create and return path to a sample CV .docx with Research Experience."""
    path = tmp_dir / "sample_cv.docx"
    return _make_cv_docx(path)


@pytest.fixture
def cv_no_research(tmp_dir):
    """Create a CV .docx without a Research Experience section."""
    path = tmp_dir / "cv_no_research.docx"
    return _make_cv_docx(path, include_research_exp=False)


# ---------------------------------------------------------------------------
# Study helpers
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_studies():
    """Return a list of sample Study objects."""
    return [
        Study(
            phase="Phase I",
            subcategory="Oncology",
            year=2024,
            sponsor="Pfizer",
            protocol="PF-99999",
            description_full="A Phase 1 study of PF-99999 (pembrolizumab) in advanced lung cancer",
            description_masked="A Phase 1 study of XXX in advanced lung cancer",
        ),
        Study(
            phase="Phase I",
            subcategory="Oncology",
            year=2023,
            sponsor="Novartis",
            protocol="NVS-789",
            description_full="First-in-human study of NVS-789 for metastatic breast cancer",
            description_masked="First-in-human study of XXX for metastatic breast cancer",
        ),
        Study(
            phase="Phase II\u2013IV",
            subcategory="Oncology",
            year=2024,
            sponsor="Roche",
            protocol="RO-777",
            description_full="Phase 3 study of RO-777 (atezolizumab) vs placebo in TNBC",
            description_masked="Phase 3 study of XXX vs placebo in TNBC",
        ),
    ]


@pytest.fixture
def sample_research_exp(sample_studies):
    """Return a populated ResearchExperience object."""
    re_exp = ResearchExperience()
    for s in sample_studies:
        phase = re_exp.get_or_create_phase(s.phase)
        subcat = phase.get_or_create_subcategory(s.subcategory)
        subcat.studies.append(s)
    return re_exp
