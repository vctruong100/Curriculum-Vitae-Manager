"""
Comprehensive pytest suite for the three major fixes:

Fix 1 — Style-agnostic heading detection
  Phase/Subcategory headings styled as Word Heading 1/2/3 must be
  detected by text content, not style name.  The Research Experience
  section must NOT be prematurely terminated by heading-styled phase
  or subcategory headings.

Fix 2 — Sort-disabled preserves existing paragraphs
  When enable_sort_existing=False, existing study paragraphs are left
  completely untouched (formatting, runs, ordering).  Only newly
  injected studies are added.

Fix 3 — Per-CV result folder output routing
  Output files go into  result/<original_cv_name>/  folders.
  Mode B writing into the same folder as Mode A when processing
  that Mode A output.  Logs are copied into the per-CV folder.

All tests are hermetic — synthetic .docx and .xlsx files are created
on the fly in temporary directories.
"""

import sys
import os
import re
from pathlib import Path

import pytest

APP_ROOT = Path(__file__).parent.parent.resolve()
if str(APP_ROOT) not in sys.path:
    sys.path.insert(0, str(APP_ROOT))

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from openpyxl import Workbook

from config import AppConfig, set_config
from processor import CVProcessor, _RESULT_SUFFIX_RE
from models import Study, ResearchExperience
from docx_handler import CVDocxHandler
from normalizer import normalize_heading_key, normalize_subcat_key


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def tmp_dir(tmp_path):
    return tmp_path


@pytest.fixture
def app_config(tmp_dir):
    config = AppConfig(data_root=str(tmp_dir / "data"))
    config.ensure_user_directories()
    set_config(config)
    return config


# ---------------------------------------------------------------------------
# Helpers — CV builders
# ---------------------------------------------------------------------------

def _make_master(path: Path, studies_data=None):
    """Create a master xlsx.  Default: Phase I > Healthy Adults with 3 studies."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Studies"

    if studies_data is None:
        studies_data = [
            ("Phase I", None, None),
            ("Healthy Adults", None, None),
            (2025, "ELI LILLY AND COMPANY LY246736: A single-ascending and "
                   "multiple-ascending dose study of LY246736 in healthy participants",
                   "ELI LILLY AND COMPANY: A single-ascending and multiple-ascending "
                   "dose study of XXXX in healthy participants"),
            (2025, "VISTERRA VIS-001: A Phase I, randomized, placebo-controlled, "
                   "double blind, multiple ascending dose trial",
                   "VISTERRA: A Phase I, randomized, placebo-controlled, double blind, "
                   "multiple ascending dose trial"),
            (2026, "MERCK MK4082-002: A Multiple-Ascending Dose Study to Evaluate "
                   "the Safety of MK-4082 in Healthy Participants",
                   "MERCK: A Multiple-Ascending Dose Study to Evaluate the Safety "
                   "of XXXX in Healthy Participants"),
        ]

    row = 1
    for item in studies_data:
        ws.cell(row=row, column=1, value=item[0])
        if len(item) > 1 and item[1] is not None:
            ws.cell(row=row, column=2, value=item[1])
        if len(item) > 2 and item[2] is not None:
            ws.cell(row=row, column=3, value=item[2])
        row += 1

    wb.save(path)
    wb.close()
    return path


def _add_study_para(doc, year, sponsor, protocol, description,
                    font_name="Calibri", font_size=11):
    """Add a study paragraph with proper formatting."""
    p = doc.add_paragraph()
    run_y = p.add_run(f"{year}\t")
    run_y.font.name = font_name
    run_y.font.size = Pt(font_size)
    run_s = p.add_run(sponsor)
    run_s.bold = True
    run_s.font.name = font_name
    run_s.font.size = Pt(font_size)
    if protocol:
        run_sp = p.add_run(" ")
        run_sp.font.name = font_name
        run_p = p.add_run(protocol)
        run_p.bold = True
        run_p.font.color.rgb = RGBColor(0xFF, 0, 0)
        run_p.font.name = font_name
        run_p.font.size = Pt(font_size)
    run_d = p.add_run(f": {description}")
    run_d.font.name = font_name
    run_d.font.size = Pt(font_size)
    pf = p.paragraph_format
    pf.left_indent = Inches(0)
    pf.first_line_indent = Inches(-0.5)
    return p


def _make_cv_normal_style(path: Path):
    """CV where Phase I and Healthy Adults use Normal style (baseline)."""
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")

    doc.add_heading("Research Experience", level=1)
    doc.add_paragraph("Phase I")
    doc.add_paragraph("Healthy Adults")

    _add_study_para(doc, 2025, "ELI LILLY AND COMPANY", "LY246736",
                    "A single-ascending and multiple-ascending dose study "
                    "of XXXX in healthy participants")
    _add_study_para(doc, 2025, "VISTERRA", "VIS-001",
                    "A Phase I, randomized, placebo-controlled, double blind, "
                    "multiple ascending dose trial")

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")
    doc.save(path)
    return path


def _make_cv_heading_style(path: Path, phase_level=2, subcat_level=3):
    """CV where Phase I and Healthy Adults use Word Heading styles."""
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")

    doc.add_heading("Research Experience", level=1)
    doc.add_heading("Phase I", level=phase_level)
    doc.add_heading("Healthy Adults", level=subcat_level)

    _add_study_para(doc, 2025, "ELI LILLY AND COMPANY", "LY246736",
                    "A single-ascending and multiple-ascending dose study "
                    "of XXXX in healthy participants")
    _add_study_para(doc, 2025, "VISTERRA", "VIS-001",
                    "A Phase I, randomized, placebo-controlled, double blind, "
                    "multiple ascending dose trial")

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")
    doc.save(path)
    return path


def _make_cv_mixed_heading(path: Path):
    """CV with UPPERCASE phase in Heading 2 and subcategory in Heading 3."""
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")

    doc.add_heading("Research Experience", level=1)
    doc.add_heading("PHASE I", level=2)
    doc.add_heading("Healthy Adults", level=3)

    _add_study_para(doc, 2025, "ELI LILLY AND COMPANY", "LY246736",
                    "A single-ascending and multiple-ascending dose study "
                    "of XXXX in healthy participants")
    _add_study_para(doc, 2025, "VISTERRA", "VIS-001",
                    "A Phase I, randomized, placebo-controlled, double blind, "
                    "multiple ascending dose trial")

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")
    doc.save(path)
    return path


def _make_cv_table_heading(path: Path):
    """CV where headings are inside a table cell — edge case.
    Note: parse_research_experience only iterates doc.paragraphs, so
    table cells are NOT reached.  This test verifies no crash and that
    the section boundary detection works regardless."""
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Research Experience", level=1)

    doc.add_paragraph("Phase I")
    doc.add_paragraph("Healthy Adults")
    _add_study_para(doc, 2025, "ELI LILLY AND COMPANY", "LY246736",
                    "A single-ascending dose study of XXXX in healthy participants")

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")
    doc.save(path)
    return path


def _make_cv_split_runs(path: Path):
    """Phase heading split across multiple runs: bold 'Phase ' + non-bold 'I'."""
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Research Experience", level=1)

    p_phase = doc.add_paragraph()
    r1 = p_phase.add_run("Phase ")
    r1.bold = True
    r2 = p_phase.add_run("I")
    r2.bold = False

    doc.add_paragraph("Healthy Adults")
    _add_study_para(doc, 2025, "ELI LILLY AND COMPANY", "LY246736",
                    "A single-ascending dose study of XXXX in healthy participants")

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")
    doc.save(path)
    return path


def _make_cv_multispace_subcat(path: Path):
    """Subcategory heading with multiple spaces: 'Healthy  Adults'."""
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Research Experience", level=1)

    doc.add_paragraph("Phase I")
    doc.add_paragraph("Healthy  Adults")
    _add_study_para(doc, 2025, "ELI LILLY AND COMPANY", "LY246736",
                    "A single-ascending and multiple-ascending dose study "
                    "of XXXX in healthy participants, participants with "
                    "obesity and hypertension, and participants with "
                    "decreased estimated glomerular filtration rate")

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Extraction helpers
# ---------------------------------------------------------------------------

def _extract_items(docx_path: Path):
    """Extract (type, text) from a generated docx.

    type is 'phase', 'subcat', 'study', or 'other'.
    """
    from normalizer import is_phase_heading
    doc = Document(docx_path)
    items = []
    in_re = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if "research experience" in text.lower():
            in_re = True
            continue
        if not in_re:
            continue
        if is_phase_heading(text):
            items.append(("phase", text))
        elif text[:4].isdigit() and "\t" in para.text:
            items.append(("study", text))
        else:
            bold_runs = [r for r in para.runs if r.bold]
            if bold_runs and len(text.split()) <= 5:
                items.append(("subcat", text))
            else:
                items.append(("other", text))
    return items


def _count_para_with_text(docx_path, needle):
    """Count paragraphs containing *needle* (case-insensitive)."""
    doc = Document(docx_path)
    needle_lower = needle.lower()
    count = 0
    for para in doc.paragraphs:
        if needle_lower in para.text.lower():
            count += 1
    return count


# ===========================================================================
# FIX 1 — Style-agnostic heading detection
# ===========================================================================

class TestStyleAgnosticHeading:
    """Phase/Subcategory headings styled as Word Heading 2/3 must be
    detected properly and must NOT terminate the Research Experience
    section."""

    def test_heading_style_phase_detected(self, app_config, tmp_dir):
        """Phase I styled as Heading 2 is still detected as phase."""
        cv_path = _make_cv_heading_style(tmp_dir / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()

        phase_names = [p.name for p in research.phases]
        assert "Phase I" in phase_names, (
            f"Phase I not found in phases: {phase_names}"
        )

    def test_heading_style_subcat_detected(self, app_config, tmp_dir):
        """Subcategory 'Healthy Adults' styled as Heading 3 is detected."""
        cv_path = _make_cv_heading_style(tmp_dir / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()

        for phase in research.phases:
            if phase.name == "Phase I":
                subcat_names = [sc.name for sc in phase.subcategories]
                assert any("Healthy" in n for n in subcat_names), (
                    f"Healthy Adults not found in subcategories: {subcat_names}"
                )
                return
        pytest.fail("Phase I not found")

    def test_heading_style_studies_parsed(self, app_config, tmp_dir):
        """Studies after heading-styled phase/subcat are parsed."""
        cv_path = _make_cv_heading_style(tmp_dir / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()

        all_studies = research.get_all_studies()
        assert len(all_studies) >= 2, (
            f"Expected >=2 studies, got {len(all_studies)}"
        )

    def test_heading_style_section_not_truncated(self, app_config, tmp_dir):
        """Research Experience section includes all content up to
        the next real section ('Publications'), not truncated at a
        heading-styled Phase I."""
        cv_path = _make_cv_heading_style(tmp_dir / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        start, end = handler.find_research_experience_section()

        assert start is not None
        assert end is not None
        span = end - start
        assert span >= 4, (
            f"Section span too small ({span}); likely truncated at "
            f"heading-styled Phase I"
        )

    def test_uppercase_heading_style(self, app_config, tmp_dir):
        """'PHASE I' in Heading 2 is detected as Phase I."""
        cv_path = _make_cv_mixed_heading(tmp_dir / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()

        phase_names = [p.name for p in research.phases]
        assert "Phase I" in phase_names

    def test_end_to_end_heading_style_no_duplicates(self, app_config, tmp_dir):
        """Full update/inject with heading-styled CV produces no duplicate
        Phase or Subcategory blocks."""
        cv_path = _make_cv_heading_style(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=True,
        )
        assert result.success is True

        items = _extract_items(output)
        phase_count = sum(1 for t, txt in items if t == "phase")
        assert phase_count == 1, (
            f"Expected 1 Phase I heading, got {phase_count}. Items: {items}"
        )

    @pytest.mark.parametrize("level", [1, 2, 3])
    def test_various_heading_levels(self, app_config, tmp_dir, level):
        """Phase headings at Heading 1, 2, or 3 all detected."""
        cv_path = _make_cv_heading_style(
            tmp_dir / f"cv_h{level}.docx", phase_level=level,
        )
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()

        phase_names = [p.name for p in research.phases]
        assert "Phase I" in phase_names

    def test_split_runs_still_detected(self, app_config, tmp_dir):
        """Phase heading split across runs is still detected."""
        cv_path = _make_cv_split_runs(tmp_dir / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        research = handler.parse_research_experience()

        phase_names = [p.name for p in research.phases]
        assert "Phase I" in phase_names

    def test_multispace_subcategory(self, app_config, tmp_dir):
        """'Healthy  Adults' (double space) matches 'Healthy Adults'."""
        cv_path = _make_cv_multispace_subcat(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=True,
        )
        assert result.success is True

        items = _extract_items(output)
        # The key check: no duplicate "Healthy Adults" headings.
        # "Healthy  Adults" (double space) must normalise to the same
        # canonical key as "Healthy Adults" from master, so only ONE
        # Healthy Adults heading should appear.
        healthy_count = sum(
            1 for t, txt in items
            if t == "subcat" and "healthy" in txt.lower()
        )
        assert healthy_count == 1, (
            f"Expected exactly 1 Healthy Adults heading, "
            f"got {healthy_count}. Items: {items}"
        )


# ===========================================================================
# FIX 2 — Sort-disabled preserves existing paragraphs
# ===========================================================================

class TestSortDisabledPreservesExisting:
    """When enable_sort_existing=False, all existing study paragraphs
    must remain exactly as they appear in the original document."""

    def test_existing_paragraphs_unchanged(self, app_config, tmp_dir):
        """Existing study paragraphs must not be rewritten."""
        cv_path = _make_cv_normal_style(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        # Read original study texts
        orig_doc = Document(cv_path)
        orig_studies = []
        for para in orig_doc.paragraphs:
            text = para.text.strip()
            if text[:4].isdigit() and "\t" in para.text:
                orig_studies.append(text)

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success is True

        out_doc = Document(output)
        out_studies = []
        for para in out_doc.paragraphs:
            text = para.text.strip()
            if text[:4].isdigit() and "\t" in para.text:
                out_studies.append(text)

        # Every original study text must still appear in output
        for orig in orig_studies:
            assert orig in out_studies, (
                f"Original study '{orig[:60]}...' not found in output"
            )

    def test_existing_order_preserved(self, app_config, tmp_dir):
        """Relative order of existing studies is preserved."""
        cv_path = _make_cv_normal_style(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        # Record original order
        orig_doc = Document(cv_path)
        orig_order = []
        for para in orig_doc.paragraphs:
            text = para.text.strip()
            if text[:4].isdigit() and "\t" in para.text:
                orig_order.append(text)

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success is True

        out_doc = Document(output)
        out_all = []
        for para in out_doc.paragraphs:
            text = para.text.strip()
            if text[:4].isdigit() and "\t" in para.text:
                out_all.append(text)

        # Filter output to only original studies and check order
        out_orig_only = [s for s in out_all if s in orig_order]
        assert out_orig_only == orig_order, (
            f"Original order not preserved.\n"
            f"Expected: {orig_order}\nGot: {out_orig_only}"
        )

    def test_existing_run_formatting_preserved(self, app_config, tmp_dir):
        """Existing study paragraphs retain their original run count and
        bold/color properties (not re-created from scratch)."""
        cv_path = _make_cv_normal_style(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        # Record run counts of original study paragraphs
        orig_doc = Document(cv_path)
        orig_run_counts = {}
        for para in orig_doc.paragraphs:
            text = para.text.strip()
            if text[:4].isdigit() and "\t" in para.text:
                orig_run_counts[text] = len(para.runs)

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success is True

        out_doc = Document(output)
        for para in out_doc.paragraphs:
            text = para.text.strip()
            if text in orig_run_counts:
                assert len(para.runs) == orig_run_counts[text], (
                    f"Run count changed for '{text[:50]}...': "
                    f"was {orig_run_counts[text]}, now {len(para.runs)}"
                )

    def test_new_studies_injected(self, app_config, tmp_dir):
        """New studies (above year_bound) are injected."""
        cv_path = _make_cv_normal_style(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success is True

        out_doc = Document(output)
        all_text = " ".join(p.text for p in out_doc.paragraphs)
        assert "MERCK" in all_text, "MERCK 2026 study was not injected"

    def test_no_duplicate_headings(self, app_config, tmp_dir):
        """Sort-disabled path must not create duplicate Phase/Subcat headings."""
        cv_path = _make_cv_normal_style(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success is True

        items = _extract_items(output)
        phase_count = sum(1 for t, _ in items if t == "phase")
        assert phase_count == 1, f"Expected 1 Phase heading, got {phase_count}"

    def test_idempotent_sort_false(self, app_config, tmp_dir):
        """Running sort-false twice produces the same output."""
        cv_path = _make_cv_normal_style(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)

        out1 = tmp_dir / "out1.docx"
        r1 = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=out1,
            enable_sort_existing=False,
        )
        assert r1.success

        out2 = tmp_dir / "out2.docx"
        r2 = processor.mode_a_update_inject(
            out1, master_path=master_path, output_path=out2,
            enable_sort_existing=False,
        )
        assert r2.success

        doc1 = Document(out1)
        doc2 = Document(out2)
        texts1 = [p.text for p in doc1.paragraphs if p.text.strip()]
        texts2 = [p.text for p in doc2.paragraphs if p.text.strip()]
        assert texts1 == texts2, "Second run changed output"

    def test_sort_true_still_rewrites(self, app_config, tmp_dir):
        """Sort-true path still fully rewrites the section (baseline)."""
        cv_path = _make_cv_normal_style(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=True,
        )
        assert result.success is True

        out_doc = Document(output)
        all_text = " ".join(p.text for p in out_doc.paragraphs)
        assert "MERCK" in all_text


# ===========================================================================
# FIX 3 — Per-CV result folder output routing
# ===========================================================================

class TestDeriveOriginalCvName:
    """Unit tests for _derive_original_cv_name."""

    def test_plain_name(self):
        assert CVProcessor._derive_original_cv_name(
            Path("Jane Doe CV.docx")
        ) == "Jane Doe CV"

    def test_updated_suffix(self):
        assert CVProcessor._derive_original_cv_name(
            Path("Jane Doe CV (Updated 2025-03-16).docx")
        ) == "Jane Doe CV"

    def test_redacted_suffix(self):
        assert CVProcessor._derive_original_cv_name(
            Path("Jane Doe CV (Redacted 2025-03-16).docx")
        ) == "Jane Doe CV"

    def test_chained_suffixes(self):
        assert CVProcessor._derive_original_cv_name(
            Path("Jane Doe CV (Updated 2025-03-16) (Redacted 2025-03-16).docx")
        ) == "Jane Doe CV"

    def test_spaces_in_name(self):
        assert CVProcessor._derive_original_cv_name(
            Path("Dr Jane A Doe CV (Updated 2025-01-01).docx")
        ) == "Dr Jane A Doe CV"


class TestResultSuffixRegex:
    """Ensure _RESULT_SUFFIX_RE matches the expected patterns."""

    def test_updated(self):
        assert _RESULT_SUFFIX_RE.search(" (Updated 2025-03-16)")

    def test_redacted(self):
        assert _RESULT_SUFFIX_RE.search(" (Redacted 2025-03-16)")

    def test_no_match_plain(self):
        assert _RESULT_SUFFIX_RE.search("Jane Doe CV") is None


class TestOutputRouting:
    """mode_a and mode_b output files go into result/<cv_name>/."""

    def test_mode_a_creates_subfolder(self, app_config, tmp_dir):
        cv_path = _make_cv_normal_style(tmp_dir / "Jane Doe CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, enable_sort_existing=True,
        )
        assert result.success

        output_path = Path(result.output_path)
        assert output_path.exists()
        assert output_path.parent.name == "Jane Doe CV"
        assert output_path.parent.parent.name == "result"

    def test_mode_b_creates_subfolder(self, app_config, tmp_dir):
        cv_path = _make_cv_normal_style(tmp_dir / "Jane Doe CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path,
        )
        assert result.success

        output_path = Path(result.output_path)
        assert output_path.exists()
        assert output_path.parent.name == "Jane Doe CV"

    def test_mode_b_after_mode_a_same_folder(self, app_config, tmp_dir):
        """Mode B processing Mode A output writes to the same CV folder."""
        cv_path = _make_cv_normal_style(tmp_dir / "Jane Doe CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)

        result_a = processor.mode_a_update_inject(
            cv_path, master_path=master_path, enable_sort_existing=True,
        )
        assert result_a.success
        mode_a_output = Path(result_a.output_path)
        mode_a_folder = mode_a_output.parent

        result_b = processor.mode_b_redact_protocols(
            mode_a_output, master_path=master_path,
        )
        assert result_b.success
        mode_b_output = Path(result_b.output_path)

        assert mode_b_output.parent == mode_a_folder, (
            f"Mode B folder '{mode_b_output.parent}' != "
            f"Mode A folder '{mode_a_folder}'"
        )

    def test_explicit_output_path_honoured(self, app_config, tmp_dir):
        """When user specifies --output, that path is used as-is."""
        cv_path = _make_cv_normal_style(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")
        explicit = tmp_dir / "custom_output.docx"

        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=explicit,
            enable_sort_existing=True,
        )
        assert result.success
        assert Path(result.output_path) == explicit


class TestNoLogsInResultFolder:
    """Result folder must contain only .docx files, no JSON/CSV logs."""

    def test_no_log_files_in_result_folder(self, app_config, tmp_dir):
        cv_path = _make_cv_normal_style(tmp_dir / "Jane Doe CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, enable_sort_existing=True,
        )
        assert result.success

        result_folder = Path(result.output_path).parent
        log_files = list(result_folder.glob("*.json")) + list(
            result_folder.glob("*.csv")
        )
        assert len(log_files) == 0, (
            f"Expected NO log files in {result_folder}, "
            f"found {[f.name for f in log_files]}"
        )

        docx_files = list(result_folder.glob("*.docx"))
        assert len(docx_files) >= 1


# ===========================================================================
# Combined / integration tests
# ===========================================================================

class TestCombinedFixes:
    """Tests that exercise multiple fixes together."""

    def test_heading_style_sort_false_no_duplicates(self, app_config, tmp_dir):
        """Heading-styled CV + sort-disabled = no duplicate headings."""
        cv_path = _make_cv_heading_style(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success

        items = _extract_items(output)
        phase_count = sum(1 for t, _ in items if t == "phase")
        assert phase_count == 1

    def test_heading_style_output_routing(self, app_config, tmp_dir):
        """Heading-styled CV + output routing creates correct subfolder."""
        cv_path = _make_cv_heading_style(tmp_dir / "My CV.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, enable_sort_existing=True,
        )
        assert result.success
        assert Path(result.output_path).parent.name == "My CV"

    def test_sort_false_heading_style_preserves_runs(self, app_config, tmp_dir):
        """Sort-false + heading style: existing runs are preserved."""
        cv_path = _make_cv_heading_style(tmp_dir / "cv.docx")
        master_path = _make_master(tmp_dir / "master.xlsx")

        orig_doc = Document(cv_path)
        orig_run_counts = {}
        for para in orig_doc.paragraphs:
            text = para.text.strip()
            if text[:4].isdigit() and "\t" in para.text:
                orig_run_counts[text] = len(para.runs)

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path, output_path=output,
            enable_sort_existing=False,
        )
        assert result.success

        out_doc = Document(output)
        for para in out_doc.paragraphs:
            text = para.text.strip()
            if text in orig_run_counts:
                assert len(para.runs) == orig_run_counts[text], (
                    f"Runs changed for '{text[:40]}'"
                )
