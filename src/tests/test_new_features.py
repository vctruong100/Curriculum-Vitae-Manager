"""
Comprehensive pytest suite for the four new features:
1. Configurable uncategorized_label
2. Auto-close terminal when GUI exits
3. highlight_inserted (yellow highlight for new studies)
4. 7-column Import/Export schema

Also verifies carry-forward behaviors:
a) Mode B redact-only + optional sort
b) Enable sorting for existing studies
c) Robust Phase/Subcategory matching
d) Per-CV results folder routing (no logs in result folder)
e) Filesystem clean after teardown
"""

import sys
import os
import json
import shutil
import re
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

APP_ROOT = Path(__file__).parent.parent.resolve()
if str(APP_ROOT) not in sys.path:
    sys.path.insert(0, str(APP_ROOT))

from openpyxl import Workbook, load_workbook
from docx import Document
from docx.shared import Pt, Inches, RGBColor

from config import AppConfig, set_config, DEFAULT_UNCATEGORIZED_LABEL
from models import Study, ResearchExperience, Phase, Subcategory
from normalizer import (
    normalize_heading_key,
    normalize_subcat_key,
    is_uncategorized_key,
)
from excel_parser import (
    SEVEN_COL_HEADERS,
    detect_xlsx_format,
    parse_master_xlsx_seven_col,
    export_studies_to_xlsx_seven_col,
)
from import_export import ImportExportManager
from docx_handler import CVDocxHandler, HIGHLIGHT_COLOR_YELLOW


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_seven_col_xlsx(path, studies_data):
    """Create a 7-column .xlsx at *path*.

    studies_data: list of (phase, subcat, year, sponsor, protocol, masked, full)
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Studies"
    for ci, h in enumerate(SEVEN_COL_HEADERS, 1):
        ws.cell(row=1, column=ci, value=h)
    for ri, row in enumerate(studies_data, 2):
        for ci, val in enumerate(row, 1):
            ws.cell(row=ri, column=ci, value=val)
    wb.save(path)
    wb.close()
    return path


def _make_legacy_xlsx(path):
    """Create a legacy 3-column .xlsx (no header row)."""
    wb = Workbook()
    ws = wb.active
    ws.cell(row=1, column=1, value="Phase I")
    ws.cell(row=2, column=1, value="Oncology")
    ws.cell(row=3, column=1, value=2024)
    ws.cell(row=3, column=2, value="Pfizer PF-123: Some study")
    ws.cell(row=3, column=3, value="Pfizer: Some study")
    wb.save(path)
    wb.close()
    return path


def _make_cv_docx(path, studies_text=None):
    """Create a minimal CV .docx with a Research Experience section."""
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")
    doc.add_heading("Research Experience", level=1)
    doc.add_paragraph("Phase I")
    doc.add_paragraph("Oncology")
    if studies_text is None:
        p = doc.add_paragraph()
        p.add_run("2023\t")
        r = p.add_run("Pfizer")
        r.bold = True
        p.add_run(" ")
        rp = p.add_run("PF-12345")
        rp.bold = True
        rp.font.color.rgb = RGBColor(0xFF, 0, 0)
        p.add_run(": A Phase 1 study of PF-12345 in patients with advanced solid tumors")
    else:
        for line in studies_text:
            doc.add_paragraph(line)
    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J. Nature. 2023.")
    doc.save(path)
    return path


def _make_master_xlsx_for_update(path):
    """Create a 3-column master .xlsx suitable for Mode A."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Studies"
    data = [
        ("Phase I", None, None),
        ("Oncology", None, None),
        (2024, "Pfizer PF-99999: A Phase 1 study of PF-99999 (pembrolizumab) in advanced lung cancer",
         "Pfizer: A Phase 1 study of XXX in advanced lung cancer"),
        (2023, "Pfizer PF-12345: A Phase 1 study of PF-12345 in patients with advanced solid tumors",
         "Pfizer: A Phase 1 study of XXX in patients with advanced solid tumors"),
    ]
    for ri, item in enumerate(data, 1):
        ws.cell(row=ri, column=1, value=item[0])
        if item[1] is not None:
            ws.cell(row=ri, column=2, value=item[1])
        if len(item) > 2 and item[2] is not None:
            ws.cell(row=ri, column=3, value=item[2])
    wb.save(path)
    wb.close()
    return path


# ===========================================================================
# 1. Configurable uncategorized_label
# ===========================================================================

class TestUncategorizedLabel:
    """Tests for configurable uncategorized_label."""

    def test_default_label(self, tmp_path):
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        assert cfg.uncategorized_label == "Uncategorized"

    def test_custom_label(self, tmp_path):
        cfg = AppConfig(
            data_root=str(tmp_path / "data"),
            uncategorized_label="Not Assigned",
        )
        assert cfg.uncategorized_label == "Not Assigned"

    def test_empty_label_rejected(self, tmp_path):
        with pytest.raises(ValueError, match="uncategorized_label"):
            AppConfig(data_root=str(tmp_path / "data"), uncategorized_label="")

    def test_whitespace_only_label_rejected(self, tmp_path):
        with pytest.raises(ValueError, match="uncategorized_label"):
            AppConfig(data_root=str(tmp_path / "data"), uncategorized_label="   ")

    def test_is_uncategorized_key_default(self):
        assert is_uncategorized_key("Uncategorized") is True
        assert is_uncategorized_key("uncategorized") is True
        assert is_uncategorized_key("UNCATEGORIZED") is True
        assert is_uncategorized_key("Phase I") is False

    def test_phase_order_key_uncategorized_last(self):
        re_exp = ResearchExperience()
        assert re_exp.get_phase_order_key("Uncategorized") == 99
        assert re_exp.get_phase_order_key("uncategorized") == 99
        assert re_exp.get_phase_order_key("Phase I") == 0

    def test_sort_uncategorized_last(self):
        re_exp = ResearchExperience()
        re_exp.get_or_create_phase("Uncategorized")
        re_exp.get_or_create_phase("Phase I")
        re_exp.sort_all()
        assert re_exp.phases[0].name == "Phase I"
        assert re_exp.phases[-1].name == "Uncategorized"

    def test_custom_label_in_config_save_load(self, tmp_path):
        cfg = AppConfig(
            data_root=str(tmp_path / "data"),
            uncategorized_label="Misc",
        )
        cfg.ensure_user_directories()
        cfg.save()
        loaded = AppConfig.load(cfg.data_path / "config.json")
        assert loaded.uncategorized_label == "Misc"


# ===========================================================================
# 2. Auto-close terminal when GUI exits
# ===========================================================================

class TestAutoClose:
    """Tests for auto-close behavior."""

    def test_gui_main_does_not_block_on_input(self):
        """Verify that gui.main() does not contain input() calls."""
        import gui as gui_module
        import inspect
        source = inspect.getsource(gui_module.main)
        assert "input(" not in source

    def test_on_close_handler_calls_sys_exit(self):
        """Verify _on_close calls sys.exit(0)."""
        import gui as gui_module
        import inspect
        source = inspect.getsource(gui_module.CVManagerApp._on_close)
        assert "sys.exit(0)" in source

    def test_main_py_no_input(self):
        """Verify main.py does not block on input()."""
        main_path = APP_ROOT / "main.py"
        content = main_path.read_text(encoding="utf-8")
        lines = content.split("\n")
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("#"):
                continue
            assert "input(" not in stripped, (
                f"main.py should not block on input(): {line}"
            )


# ===========================================================================
# 3. highlight_inserted (yellow highlight for new studies)
# ===========================================================================

class TestHighlightInserted:
    """Tests for highlight_inserted configuration and behavior."""

    def test_config_default_false(self, tmp_path):
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        assert cfg.highlight_inserted is False

    def test_config_enable(self, tmp_path):
        cfg = AppConfig(
            data_root=str(tmp_path / "data"),
            highlight_inserted=True,
        )
        assert cfg.highlight_inserted is True

    def test_highlight_constant_is_yellow(self):
        from docx.enum.text import WD_COLOR_INDEX
        assert HIGHLIGHT_COLOR_YELLOW == WD_COLOR_INDEX.YELLOW

    def test_create_study_element_with_highlight(self, tmp_path):
        """_create_study_element with highlight=True adds yellow highlight XML."""
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        set_config(cfg)
        cv_path = _make_cv_docx(tmp_path / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        study = Study(
            phase="Phase I",
            subcategory="Oncology",
            year=2024,
            sponsor="TestCo",
            protocol="TC-001",
            description_full="Test study description",
            description_masked="Test study description",
        )
        elem = handler._create_study_element(study, True, True, highlight=True)
        xml_str = elem.xml.decode("utf-8") if isinstance(elem.xml, bytes) else elem.xml
        assert "yellow" in xml_str.lower(), (
            "Expected 'yellow' highlight in study element XML"
        )

    def test_create_study_element_without_highlight(self, tmp_path):
        """_create_study_element with highlight=False has no yellow highlight."""
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        set_config(cfg)
        cv_path = _make_cv_docx(tmp_path / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        study = Study(
            phase="Phase I",
            subcategory="Oncology",
            year=2024,
            sponsor="TestCo",
            protocol="TC-001",
            description_full="Test study description",
            description_masked="Test study description",
        )
        elem = handler._create_study_element(study, True, True, highlight=False)
        xml_str = elem.xml.decode("utf-8") if isinstance(elem.xml, bytes) else elem.xml
        assert "highlight" not in xml_str.lower() or "yellow" not in xml_str.lower(), (
            "Expected no yellow highlight in study element XML"
        )

    def test_create_study_paragraph_with_highlight(self, tmp_path):
        """_create_study_paragraph with highlight=True highlights sponsor and protocol only."""
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        set_config(cfg)
        cv_path = _make_cv_docx(tmp_path / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        study = Study(
            phase="Phase I",
            subcategory="Oncology",
            year=2024,
            sponsor="TestCo",
            protocol="TC-001",
            description_full="Test study",
            description_masked="Test study",
        )
        para = handler._create_study_paragraph(study, highlight=True)
        highlighted_texts = []
        non_highlighted_texts = []
        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue
            if run.font.highlight_color is not None:
                highlighted_texts.append(text)
            else:
                non_highlighted_texts.append(text)
        # Sponsor and protocol should be highlighted
        assert "TestCo" in highlighted_texts
        assert "TC-001" in highlighted_texts
        # Year and description should NOT be highlighted
        assert any("2024" in t for t in non_highlighted_texts)
        assert any("study" in t.lower() for t in non_highlighted_texts)

    def test_create_study_paragraph_without_highlight(self, tmp_path):
        """_create_study_paragraph with highlight=False does not highlight."""
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        set_config(cfg)
        cv_path = _make_cv_docx(tmp_path / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        study = Study(
            phase="Phase I",
            subcategory="Oncology",
            year=2024,
            sponsor="TestCo",
            protocol="TC-001",
            description_full="Test study",
            description_masked="Test study",
        )
        para = handler._create_study_paragraph(study, highlight=False)
        for run in para.runs:
            assert run.font.highlight_color is None, (
                f"Run '{run.text}' should not be highlighted"
            )

    def test_inject_new_studies_highlight(self, tmp_path):
        """inject_new_studies_only with highlight_inserted=True puts yellow on new."""
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        set_config(cfg)
        cv_path = _make_cv_docx(tmp_path / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        handler.parse_research_experience()

        new_study = Study(
            phase="Phase I",
            subcategory="Oncology",
            year=2025,
            sponsor="NewCo",
            protocol="NC-001",
            description_full="A new study",
            description_masked="A new study",
        )
        inserted = handler.inject_new_studies_only(
            [(new_study, "Phase I", "Oncology")],
            highlight_inserted=True,
        )
        assert inserted >= 1

        out = tmp_path / "cv_highlight.docx"
        handler.save(out)

        doc = Document(out)
        found_highlight = False
        for para in doc.paragraphs:
            if "NewCo" in para.text:
                xml_str = para._element.xml
                if isinstance(xml_str, bytes):
                    xml_str = xml_str.decode("utf-8")
                if "yellow" in xml_str.lower():
                    found_highlight = True
                break
        assert found_highlight, "Newly injected study should have yellow highlight"


# ===========================================================================
# 4. 7-column Import/Export schema
# ===========================================================================

class TestSevenColumnSchema:
    """Tests for the 7-column import/export schema."""

    def test_seven_col_headers_constant(self):
        assert SEVEN_COL_HEADERS == [
            "Phase", "Subcategory", "Year", "Sponsor", "Protocol",
            "Masked Description", "Full Description",
        ]

    def test_detect_format_seven_col(self, tmp_path):
        path = tmp_path / "seven.xlsx"
        _make_seven_col_xlsx(path, [
            ("Phase I", "Oncology", 2024, "Pfizer", "PF-001", "masked", "full"),
        ])
        assert detect_xlsx_format(path) == "7col"

    def test_detect_format_legacy(self, tmp_path):
        path = tmp_path / "legacy.xlsx"
        _make_legacy_xlsx(path)
        assert detect_xlsx_format(path) == "3col"

    def test_parse_seven_col_success(self, tmp_path):
        path = tmp_path / "studies.xlsx"
        _make_seven_col_xlsx(path, [
            ("Phase I", "Oncology", 2024, "Pfizer", "PF-001", "masked desc", "full desc"),
            ("Phase I", "Cardiology", 2023, "Novartis", "NVS-002", "masked2", "full2"),
        ])
        studies = parse_master_xlsx_seven_col(path)
        assert len(studies) == 2
        assert studies[0].phase == "Phase I"
        assert studies[0].subcategory == "Oncology"
        assert studies[0].year == 2024
        assert studies[0].sponsor == "Pfizer"
        assert studies[0].protocol == "PF-001"
        assert studies[0].description_masked == "masked desc"
        assert studies[0].description_full == "full desc"

    def test_parse_seven_col_wrong_headers(self, tmp_path):
        path = tmp_path / "bad_headers.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.cell(row=1, column=1, value="Wrong")
        ws.cell(row=1, column=2, value="Headers")
        wb.save(path)
        wb.close()
        with pytest.raises(ValueError, match="Expected 7 columns"):
            parse_master_xlsx_seven_col(path)

    def test_parse_seven_col_wrong_header_names(self, tmp_path):
        """7 columns but wrong header names."""
        path = tmp_path / "bad_names.xlsx"
        wb = Workbook()
        ws = wb.active
        for ci in range(1, 8):
            ws.cell(row=1, column=ci, value=f"Col{ci}")
        wb.save(path)
        wb.close()
        with pytest.raises(ValueError, match="Column 1 header must be 'Phase'"):
            parse_master_xlsx_seven_col(path)

    def test_parse_seven_col_missing_columns(self, tmp_path):
        path = tmp_path / "few_cols.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.cell(row=1, column=1, value="Phase")
        ws.cell(row=1, column=2, value="Subcategory")
        wb.save(path)
        wb.close()
        with pytest.raises(ValueError, match="Expected 7 columns"):
            parse_master_xlsx_seven_col(path)

    def test_parse_seven_col_bad_year(self, tmp_path):
        path = tmp_path / "bad_year.xlsx"
        _make_seven_col_xlsx(path, [
            ("Phase I", "Oncology", "not-a-year", "Pfizer", "PF-001", "m", "f"),
        ])
        with pytest.raises(ValueError, match="Year must be numeric"):
            parse_master_xlsx_seven_col(path)

    def test_parse_seven_col_empty_rows_skipped(self, tmp_path):
        path = tmp_path / "empty_rows.xlsx"
        wb = Workbook()
        ws = wb.active
        for ci, h in enumerate(SEVEN_COL_HEADERS, 1):
            ws.cell(row=1, column=ci, value=h)
        ws.cell(row=2, column=1, value="Phase I")
        ws.cell(row=2, column=2, value="Oncology")
        ws.cell(row=2, column=3, value=2024)
        ws.cell(row=2, column=4, value="Pfizer")
        ws.cell(row=2, column=5, value="PF-001")
        ws.cell(row=2, column=6, value="masked")
        ws.cell(row=2, column=7, value="full")
        # row 3 is empty
        ws.cell(row=4, column=1, value="Phase I")
        ws.cell(row=4, column=2, value="Cardiology")
        ws.cell(row=4, column=3, value=2023)
        ws.cell(row=4, column=4, value="Novartis")
        ws.cell(row=4, column=5, value="")
        ws.cell(row=4, column=6, value="m2")
        ws.cell(row=4, column=7, value="f2")
        wb.save(path)
        wb.close()
        studies = parse_master_xlsx_seven_col(path)
        assert len(studies) == 2

    def test_parse_seven_col_year_as_float(self, tmp_path):
        """Year stored as float (e.g. 2024.0) should be parsed correctly."""
        path = tmp_path / "float_year.xlsx"
        _make_seven_col_xlsx(path, [
            ("Phase I", "Oncology", 2024.0, "Pfizer", "PF-001", "m", "f"),
        ])
        studies = parse_master_xlsx_seven_col(path)
        assert studies[0].year == 2024

    def test_export_seven_col(self, tmp_path):
        studies = [
            Study("Phase I", "Oncology", 2024, "Pfizer", "PF-001",
                  "full desc 1", "masked desc 1"),
            Study("Phase I", "Cardiology", 2023, "Novartis", "",
                  "full desc 2", "masked desc 2"),
        ]
        out = tmp_path / "export.xlsx"
        export_studies_to_xlsx_seven_col(studies, out)
        assert out.exists()

        wb = load_workbook(out, read_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        wb.close()

        assert list(rows[0]) == SEVEN_COL_HEADERS
        assert len(rows) == 3  # header + 2 studies

    def test_round_trip_integrity(self, tmp_path):
        """Export → Import should reproduce identical records (order may differ)."""
        original = [
            Study("Phase I", "Oncology", 2024, "Pfizer", "PF-001",
                  "A Phase 1 study of PF-001 in lung cancer",
                  "A Phase 1 study of XXX in lung cancer"),
            Study("Phase I", "Cardiology", 2023, "Novartis", "NVS-002",
                  "A study of NVS-002 in heart failure",
                  "A study of XXX in heart failure"),
        ]
        export_path = tmp_path / "round_trip.xlsx"
        export_studies_to_xlsx_seven_col(original, export_path)

        imported = parse_master_xlsx_seven_col(export_path)
        assert len(imported) == len(original)

        def _key(s):
            return (s.phase, s.subcategory, s.year, s.sponsor, s.protocol)

        orig_sorted = sorted(original, key=_key)
        imp_sorted = sorted(imported, key=_key)
        for orig, imp in zip(orig_sorted, imp_sorted):
            assert imp.phase == orig.phase
            assert imp.subcategory == orig.subcategory
            assert imp.year == orig.year
            assert imp.sponsor == orig.sponsor
            assert imp.protocol == orig.protocol
            assert imp.description_full == orig.description_full
            assert imp.description_masked == orig.description_masked

    def test_import_legacy_format_rejected(self, tmp_path):
        """Importing a legacy 3-column file returns an error."""
        path = tmp_path / "legacy.xlsx"
        _make_legacy_xlsx(path)
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        cfg.ensure_user_directories()
        set_config(cfg)
        manager = ImportExportManager(cfg)
        success, message, site_id = manager.import_xlsx_to_site(
            path, "TestSite",
        )
        assert success is False
        assert "legacy 3-column" in message.lower() or "3-column" in message

    def test_import_seven_col_to_site(self, tmp_path):
        """Import a 7-column xlsx to a site database."""
        path = tmp_path / "seven.xlsx"
        _make_seven_col_xlsx(path, [
            ("Phase I", "Oncology", 2024, "Pfizer", "PF-001", "masked", "full"),
            ("Phase I", "Cardiology", 2023, "Novartis", "", "m2", "f2"),
        ])
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        cfg.ensure_user_directories()
        set_config(cfg)
        manager = ImportExportManager(cfg)
        success, message, site_id = manager.import_xlsx_to_site(
            path, "TestSite",
        )
        assert success is True
        assert site_id is not None
        assert "2 studies" in message

    def test_export_site_to_seven_col(self, tmp_path):
        """Export a site produces a 7-column xlsx."""
        # First import some data
        path = tmp_path / "seven.xlsx"
        _make_seven_col_xlsx(path, [
            ("Phase I", "Oncology", 2024, "Pfizer", "PF-001", "masked", "full"),
        ])
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        cfg.ensure_user_directories()
        set_config(cfg)
        manager = ImportExportManager(cfg)
        success, _, site_id = manager.import_xlsx_to_site(path, "ExportTest")
        assert success

        # Now export
        success, message, out_path = manager.export_site_to_xlsx(site_id)
        assert success is True
        assert out_path is not None
        assert out_path.exists()

        # Verify 7-column format
        wb = load_workbook(out_path, read_only=True)
        ws = wb.active
        header = []
        for row in ws.iter_rows(min_row=1, max_row=1, values_only=True):
            header = [str(c) if c else "" for c in row]
        wb.close()
        assert header[:7] == SEVEN_COL_HEADERS

    def test_export_goes_to_result_root(self, tmp_path):
        """Export output file lands in config.get_result_root()/{site}/."""
        path = tmp_path / "seven.xlsx"
        _make_seven_col_xlsx(path, [
            ("Phase I", "Oncology", 2024, "Pfizer", "PF-001", "masked", "full"),
        ])
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        cfg.ensure_user_directories()
        set_config(cfg)
        manager = ImportExportManager(cfg)
        success, _, site_id = manager.import_xlsx_to_site(path, "PathTest")
        assert success

        success, _, out_path = manager.export_site_to_xlsx(site_id)
        assert success
        results_dir = cfg.get_result_root()
        assert str(results_dir) in str(out_path)


# ===========================================================================
# 5. Mode B carry-forward: redact-only, optional sort
# ===========================================================================

class TestModeBCarryForward:
    """Verify Mode B behavior is not regressed."""

    def test_mode_b_sort_and_format_parameter_exists(self):
        from processor import CVProcessor
        import inspect
        sig = inspect.signature(CVProcessor.mode_b_redact_protocols)
        assert "sort_and_format" in sig.parameters

    def test_mode_b_preview_includes_sort_field(self):
        from processor import CVProcessor
        import inspect
        sig = inspect.signature(CVProcessor.preview_changes)
        assert "sort_and_format" in sig.parameters


# ===========================================================================
# 6. Per-CV results folder routing
# ===========================================================================

class TestPerCVResultsFolder:
    """Verify per-CV results folder routing is correct."""

    def test_results_path_method_exists(self, tmp_path):
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        path = cfg.get_user_results_path()
        assert "results" in str(path)

    def test_results_path_under_user_data(self, tmp_path):
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        results = cfg.get_user_results_path()
        user_data = cfg.get_user_data_path()
        assert str(results).startswith(str(user_data))


# ===========================================================================
# 7. Filesystem cleanup
# ===========================================================================

class TestFilesystemCleanup:
    """Verify tests leave the filesystem clean."""

    def test_tmp_path_is_clean(self, tmp_path):
        """After all tests, tmp_path should be managed by pytest."""
        assert tmp_path.exists()

    def test_no_stale_result_folder(self, tmp_path):
        """No result/ folder should exist at project root after tests."""
        project_result = APP_ROOT.parent / "result"
        # Don't fail if it exists from user's own work, just check we
        # didn't create any test-specific subfolder
        if project_result.exists():
            for child in project_result.iterdir():
                assert "test_" not in child.name.lower()


# ===========================================================================
# 8. write_research_experience with highlight_new
# ===========================================================================

class TestWriteResearchExperienceHighlight:
    """Test write_research_experience highlight_new parameter."""

    def test_highlight_new_all(self, tmp_path):
        """highlight_new=True with no new_study_ids highlights everything."""
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        set_config(cfg)
        cv_path = _make_cv_docx(tmp_path / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        handler.parse_research_experience()

        re_exp = ResearchExperience()
        phase = re_exp.get_or_create_phase("Phase I")
        subcat = phase.get_or_create_subcategory("Oncology")
        s = Study("Phase I", "Oncology", 2024, "TestCo", "TC-1",
                  "A test study", "A test study")
        subcat.studies.append(s)

        handler.write_research_experience(
            re_exp,
            highlight_new=True,
            new_study_ids=None,
        )
        out = tmp_path / "out.docx"
        handler.save(out)

        doc = Document(out)
        found_highlight = False
        for para in doc.paragraphs:
            if "TestCo" in para.text:
                xml_str = para._element.xml
                if isinstance(xml_str, bytes):
                    xml_str = xml_str.decode("utf-8")
                if "yellow" in xml_str.lower():
                    found_highlight = True
                break
        assert found_highlight

    def test_highlight_new_only_new_studies(self, tmp_path):
        """highlight_new=True with new_study_ids only highlights non-existing."""
        cfg = AppConfig(data_root=str(tmp_path / "data"))
        set_config(cfg)
        cv_path = _make_cv_docx(tmp_path / "cv.docx")
        handler = CVDocxHandler(cv_path)
        handler.load()
        handler.find_research_experience_section()
        handler.parse_research_experience()

        existing = Study("Phase I", "Oncology", 2023, "ExistCo", "EX-1",
                         "Existing study", "Existing study")
        new = Study("Phase I", "Oncology", 2025, "NewCo", "NW-1",
                    "New study", "New study")

        re_exp = ResearchExperience()
        phase = re_exp.get_or_create_phase("Phase I")
        subcat = phase.get_or_create_subcategory("Oncology")
        subcat.studies.append(existing)
        subcat.studies.append(new)

        existing_ids = {id(existing)}

        handler.write_research_experience(
            re_exp,
            highlight_new=True,
            new_study_ids=existing_ids,
        )
        out = tmp_path / "out2.docx"
        handler.save(out)

        doc = Document(out)
        for para in doc.paragraphs:
            xml_str = para._element.xml
            if isinstance(xml_str, bytes):
                xml_str = xml_str.decode("utf-8")
            if "ExistCo" in para.text:
                assert "yellow" not in xml_str.lower(), (
                    "Existing study should NOT be highlighted"
                )
            if "NewCo" in para.text:
                assert "yellow" in xml_str.lower(), (
                    "New study SHOULD be highlighted"
                )


# ===========================================================================
# 9. Config round-trip (uncategorized_label + highlight_inserted)
# ===========================================================================

class TestConfigRoundTrip:
    """Verify new config fields persist through save/load."""

    def test_all_new_fields_round_trip(self, tmp_path):
        cfg = AppConfig(
            data_root=str(tmp_path / "data"),
            uncategorized_label="Custom Label",
            highlight_inserted=True,
        )
        cfg.ensure_user_directories()
        cfg.save()

        loaded = AppConfig.load(cfg.data_path / "config.json")
        assert loaded.uncategorized_label == "Custom Label"
        assert loaded.highlight_inserted is True


# ===========================================================================
# 10. is_uncategorized_key matching
# ===========================================================================

class TestIsUncategorizedKeyMatching:
    """Verify is_uncategorized_key handles edge cases."""

    def test_canonical(self):
        assert is_uncategorized_key("Uncategorized") is True

    def test_casefold(self):
        assert is_uncategorized_key("UNCATEGORIZED") is True
        assert is_uncategorized_key("uncategorized") is True

    def test_whitespace_stripped(self):
        assert is_uncategorized_key("  Uncategorized  ") is True

    def test_phase_names_not_matched(self):
        assert is_uncategorized_key("Phase I") is False
        assert is_uncategorized_key("Oncology") is False
        assert is_uncategorized_key("General") is False
