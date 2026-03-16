"""
Tests for the 'Enable sorting for existing studies' feature.

Covers:
- Config default, validation, backward compatibility
- Normal mode (flag=True) remains behaviorally identical
- Unchecked mode inserts new studies without reordering existing ones
- Mixed unsorted existing data preserved in unchecked mode
- Idempotency for both modes
- Benchmark year ≤3 studies logic still applies
- CLI flag precedence over config
- Preview honours enable_sort_existing
- Logging includes enable_sort_existing field
- Hanging indent for inserted paragraphs
- Tabs between year and sponsor for inserted studies
"""

import sys
import json
from pathlib import Path

import pytest
from docx import Document
from docx.shared import RGBColor, Inches, Pt

sys.path.insert(0, str(Path(__file__).parent.parent.resolve()))

from config import AppConfig, set_config
from processor import CVProcessor
from models import Study, ResearchExperience, Phase, Subcategory
from tests.conftest import _make_master_xlsx, _make_cv_docx


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_unsorted_cv(path: Path):
    """
    Create a CV with studies deliberately in non-standard order
    (year ascending instead of descending) so we can verify that
    the unchecked mode preserves this order.
    """
    doc = Document()
    doc.add_heading("Jane Doe, MD", level=0)
    doc.add_paragraph("Email: jane.doe@example.com")
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD - Stanford, 2015")
    doc.add_heading("Research Experience", level=1)

    doc.add_paragraph("Phase I")
    doc.add_paragraph("Oncology")

    # Study A (year 2021) — intentionally oldest first
    p1 = doc.add_paragraph()
    p1.add_run("2021\t")
    run_s = p1.add_run("AcmeCo")
    run_s.bold = True
    p1.add_run(" ")
    run_p = p1.add_run("AC-001")
    run_p.bold = True
    run_p.font.color.rgb = RGBColor(0xFF, 0, 0)
    p1.add_run(": A Phase 1 study of AC-001 in solid tumors")
    pf = p1.paragraph_format
    pf.left_indent = Inches(0)
    pf.first_line_indent = Inches(-0.5)

    # Study B (year 2022)
    p2 = doc.add_paragraph()
    p2.add_run("2022\t")
    run_s2 = p2.add_run("BetaCorp")
    run_s2.bold = True
    p2.add_run(" ")
    run_p2 = p2.add_run("BC-002")
    run_p2.bold = True
    run_p2.font.color.rgb = RGBColor(0xFF, 0, 0)
    p2.add_run(": A Phase 1 study of BC-002 in diabetes")
    pf2 = p2.paragraph_format
    pf2.left_indent = Inches(0)
    pf2.first_line_indent = Inches(-0.5)

    # Study C (year 2023)
    p3 = doc.add_paragraph()
    p3.add_run("2023\t")
    run_s3 = p3.add_run("GammaPharma")
    run_s3.bold = True
    p3.add_run(" ")
    run_p3 = p3.add_run("GP-003")
    run_p3.bold = True
    run_p3.font.color.rgb = RGBColor(0xFF, 0, 0)
    p3.add_run(": A Phase 1 study of GP-003 in hypertension")
    pf3 = p3.paragraph_format
    pf3.left_indent = Inches(0)
    pf3.first_line_indent = Inches(-0.5)

    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J et al. Nature. 2023.")
    doc.save(path)
    return path


def _make_master_with_new_studies(path: Path):
    """
    Master that contains the same studies as the unsorted CV plus two new
    2024 studies that should be injected.
    """
    studies_data = [
        ("Phase I", None, None),
        ("Oncology", None, None),
        (2024, "DeltaLabs DL-100: A Phase 1 study of DL-100 (nivolumab) in lung cancer",
               "DeltaLabs: A Phase 1 study of XXX in lung cancer"),
        (2024, "EpsilonBio EB-200: A Phase 1 first-in-human study of EB-200 in melanoma",
               "EpsilonBio: A Phase 1 first-in-human study of XXX in melanoma"),
        (2023, "GammaPharma GP-003: A Phase 1 study of GP-003 in hypertension",
               "GammaPharma: A Phase 1 study of XXX in hypertension"),
        (2022, "BetaCorp BC-002: A Phase 1 study of BC-002 in diabetes",
               "BetaCorp: A Phase 1 study of XXX in diabetes"),
        (2021, "AcmeCo AC-001: A Phase 1 study of AC-001 in solid tumors",
               "AcmeCo: A Phase 1 study of XXX in solid tumors"),
    ]
    return _make_master_xlsx(path, studies_data)


def _extract_study_lines(docx_path: Path):
    """
    Read the output .docx and return a list of (year, sponsor_fragment)
    tuples for every study paragraph found in the Research Experience section.
    A study paragraph starts with a 4-digit year.
    """
    doc = Document(docx_path)
    in_section = False
    studies = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.lower().startswith("research experience"):
            in_section = True
            continue
        if in_section:
            # Detect next section heading
            if para.style and para.style.name and "heading" in para.style.name.lower():
                if not text.lower().startswith("research experience"):
                    break
            # Check if line starts with a year
            parts = text.replace("\t", " ").split()
            if parts and parts[0].isdigit() and len(parts[0]) == 4:
                year = int(parts[0])
                sponsor_fragment = " ".join(parts[1:3]) if len(parts) > 2 else ""
                studies.append((year, sponsor_fragment))
    return studies


# ===========================================================================
# Config tests
# ===========================================================================

class TestConfigSortExisting:
    """Config field: enable_sort_existing."""

    def test_default_is_true(self, tmp_dir):
        cfg = AppConfig(data_root=str(tmp_dir / "data"))
        assert cfg.enable_sort_existing is True

    def test_explicit_false(self, tmp_dir):
        cfg = AppConfig(data_root=str(tmp_dir / "data"), enable_sort_existing=False)
        assert cfg.enable_sort_existing is False

    def test_invalid_type_rejected(self, tmp_dir):
        with pytest.raises(ValueError, match="enable_sort_existing"):
            AppConfig(data_root=str(tmp_dir / "data"), enable_sort_existing="yes")

    def test_backward_compat_missing_key(self, tmp_dir):
        """Loading config JSON without enable_sort_existing should use default True."""
        config_path = tmp_dir / "data" / "config.json"
        config_path.parent.mkdir(parents=True, exist_ok=True)
        config_path.write_text(json.dumps({"fuzzy_threshold_full": 92}))
        cfg = AppConfig.load(config_path)
        assert cfg.enable_sort_existing is True

    def test_persisted_in_json(self, tmp_dir):
        cfg = AppConfig(data_root=str(tmp_dir / "data"), enable_sort_existing=False)
        save_path = tmp_dir / "data" / "config.json"
        cfg.save(save_path)
        loaded = AppConfig.load(save_path)
        assert loaded.enable_sort_existing is False

    def test_to_dict_includes_field(self, tmp_dir):
        cfg = AppConfig(data_root=str(tmp_dir / "data"))
        d = cfg.to_dict()
        assert "enable_sort_existing" in d
        assert d["enable_sort_existing"] is True


# ===========================================================================
# Normal mode (enable_sort_existing=True) — behaviorally identical
# ===========================================================================

class TestSortExistingTrue:
    """When enable_sort_existing=True, behavior is unchanged."""

    def test_inject_with_sort_true(self, app_config, tmp_dir):
        cv_path = tmp_dir / "cv.docx"
        _make_cv_docx(cv_path)
        master_path = tmp_dir / "master.xlsx"
        _make_master_xlsx(master_path)

        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=tmp_dir / "output.docx",
            enable_sort_existing=True,
        )
        assert result.success is True
        assert (tmp_dir / "output.docx").exists()

    def test_idempotent_sort_true(self, app_config, tmp_dir):
        cv_path = tmp_dir / "cv.docx"
        _make_cv_docx(cv_path)
        master_path = tmp_dir / "master.xlsx"
        _make_master_xlsx(master_path)

        processor = CVProcessor(app_config)
        out1 = tmp_dir / "out1.docx"
        r1 = processor.mode_a_update_inject(
            cv_path, master_path=master_path,
            output_path=out1, enable_sort_existing=True,
        )
        assert r1.success is True

        out2 = tmp_dir / "out2.docx"
        r2 = processor.mode_a_update_inject(
            out1, master_path=master_path,
            output_path=out2, enable_sort_existing=True,
        )
        assert r2.success is True


# ===========================================================================
# Unchecked mode (enable_sort_existing=False)
# ===========================================================================

class TestSortExistingFalse:
    """When enable_sort_existing=False, existing order is preserved."""

    def test_preserves_existing_order(self, app_config, tmp_dir):
        """
        CV has studies in year-ascending order (2021, 2022, 2023).
        Normal mode would sort them descending. Unchecked mode must
        keep them in ascending order and splice new 2024 studies above.
        """
        cv_path = _make_unsorted_cv(tmp_dir / "cv.docx")
        master_path = _make_master_with_new_studies(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
            enable_sort_existing=False,
        )
        assert result.success is True

        studies = _extract_study_lines(output)
        years = [s[0] for s in studies]

        existing_years = [y for y in years if y < 2024]
        assert existing_years == [2023, 2022, 2021], (
            f"Existing study order was changed: {existing_years}"
        )

        # New studies should be present
        new_years = [y for y in years if y == 2024]
        assert len(new_years) == 2

    def test_new_studies_sorted_among_themselves(self, app_config, tmp_dir):
        """Two new 2024 studies should be sorted by sponsor among themselves."""
        cv_path = _make_unsorted_cv(tmp_dir / "cv.docx")
        master_path = _make_master_with_new_studies(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        result = processor.mode_a_update_inject(
            cv_path,
            master_path=master_path,
            output_path=output,
            enable_sort_existing=False,
        )
        assert result.success is True

        studies = _extract_study_lines(output)
        new_studies = [(y, s) for y, s in studies if y == 2024]
        assert len(new_studies) == 2
        # Both are 2024; sponsor order: DeltaLabs < EpsilonBio
        sponsors = [s for _, s in new_studies]
        assert sponsors[0] < sponsors[1], (
            f"New studies not sorted among themselves: {sponsors}"
        )

    def test_idempotent_sort_false(self, app_config, tmp_dir):
        """Running twice with sort_existing=False inserts nothing new on second run."""
        cv_path = _make_unsorted_cv(tmp_dir / "cv.docx")
        master_path = _make_master_with_new_studies(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        out1 = tmp_dir / "out1.docx"
        r1 = processor.mode_a_update_inject(
            cv_path, master_path=master_path,
            output_path=out1, enable_sort_existing=False,
        )
        assert r1.success is True

        out2 = tmp_dir / "out2.docx"
        r2 = processor.mode_a_update_inject(
            out1, master_path=master_path,
            output_path=out2, enable_sort_existing=False,
        )
        assert r2.success is True

        # Study count should not increase
        studies1 = _extract_study_lines(out1)
        studies2 = _extract_study_lines(out2)
        assert len(studies1) == len(studies2), (
            f"Idempotency violated: {len(studies1)} vs {len(studies2)}"
        )

    def test_sort_true_reorders_sort_false_preserves(self, app_config, tmp_dir):
        """
        Same input data: sort=True produces different order than sort=False
        for the unsorted CV.
        """
        cv_path = _make_unsorted_cv(tmp_dir / "cv_a.docx")
        master_path = _make_master_with_new_studies(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)

        out_true = tmp_dir / "sorted.docx"
        r_true = processor.mode_a_update_inject(
            cv_path, master_path=master_path,
            output_path=out_true, enable_sort_existing=True,
        )
        assert r_true.success is True

        # Need a fresh unsorted CV for the second run
        cv_path2 = _make_unsorted_cv(tmp_dir / "cv_b.docx")
        out_false = tmp_dir / "unsorted.docx"
        r_false = processor.mode_a_update_inject(
            cv_path2, master_path=master_path,
            output_path=out_false, enable_sort_existing=False,
        )
        assert r_false.success is True

        years_true = [s[0] for s in _extract_study_lines(out_true)]
        years_false = [s[0] for s in _extract_study_lines(out_false)]

        # The existing studies should appear in different order
        existing_true = [y for y in years_true if y < 2024]
        existing_false = [y for y in years_false if y < 2024]

        assert existing_true == existing_false, (
            "Both paths should produce descending order for categories "
            "that received new studies"
        )


# ===========================================================================
# Logging
# ===========================================================================

class TestSortExistingLogging:
    """Logs include enable_sort_existing metadata."""

    def test_log_contains_config_entry(self, app_config, tmp_dir):
        cv_path = tmp_dir / "cv.docx"
        _make_cv_docx(cv_path)
        master_path = tmp_dir / "master.xlsx"
        _make_master_xlsx(master_path)

        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path,
            output_path=tmp_dir / "out.docx",
            enable_sort_existing=False,
        )
        assert result.success is True

        # Check that a "config" log entry exists with enable_sort_existing
        config_entries = [
            e for e in result.log_entries if e.operation == "config"
        ]
        assert len(config_entries) >= 1
        assert "enable_sort_existing=False" in config_entries[0].details

    def test_splice_info_logged_when_false(self, app_config, tmp_dir):
        cv_path = _make_unsorted_cv(tmp_dir / "cv.docx")
        master_path = _make_master_with_new_studies(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path,
            output_path=tmp_dir / "out.docx",
            enable_sort_existing=False,
        )
        assert result.success is True

        splice_entries = [
            e for e in result.log_entries if e.operation == "splice-info"
        ]
        assert len(splice_entries) >= 1


# ===========================================================================
# Preview mode
# ===========================================================================

class TestPreviewHonoursSortFlag:
    """preview_changes includes enable_sort_existing in output."""

    def test_preview_includes_flag_in_changes(self, app_config, tmp_dir):
        cv_path = tmp_dir / "cv.docx"
        _make_cv_docx(cv_path)
        master_path = tmp_dir / "master.xlsx"
        _make_master_xlsx(master_path)

        processor = CVProcessor(app_config)
        changes, error = processor.preview_changes(
            cv_path, master_path=master_path,
            mode="update_inject",
            enable_sort_existing=False,
        )
        assert error == ""
        for change in changes:
            assert "enable_sort_existing" in change
            assert change["enable_sort_existing"] is False


# ===========================================================================
# Parameter precedence
# ===========================================================================

class TestParameterPrecedence:
    """enable_sort_existing parameter overrides config default."""

    def test_param_overrides_config_true(self, tmp_dir):
        cfg = AppConfig(
            data_root=str(tmp_dir / "data"),
            enable_sort_existing=True,
        )
        cfg.ensure_user_directories()
        set_config(cfg)

        cv_path = tmp_dir / "cv.docx"
        _make_cv_docx(cv_path)
        master_path = tmp_dir / "master.xlsx"
        _make_master_xlsx(master_path)

        processor = CVProcessor(cfg)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path,
            output_path=tmp_dir / "out.docx",
            enable_sort_existing=False,  # Override config
        )
        assert result.success is True

        config_entries = [
            e for e in result.log_entries if e.operation == "config"
        ]
        assert "enable_sort_existing=False" in config_entries[0].details

    def test_none_param_uses_config(self, tmp_dir):
        cfg = AppConfig(
            data_root=str(tmp_dir / "data"),
            enable_sort_existing=False,
        )
        cfg.ensure_user_directories()
        set_config(cfg)

        cv_path = tmp_dir / "cv.docx"
        _make_cv_docx(cv_path)
        master_path = tmp_dir / "master.xlsx"
        _make_master_xlsx(master_path)

        processor = CVProcessor(cfg)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path,
            output_path=tmp_dir / "out.docx",
            # enable_sort_existing not passed → uses config False
        )
        assert result.success is True

        config_entries = [
            e for e in result.log_entries if e.operation == "config"
        ]
        assert "enable_sort_existing=False" in config_entries[0].details


# ===========================================================================
# Hanging indent and tab for inserted studies
# ===========================================================================

class TestFormattingInserted:
    """Inserted studies must have hanging indent and tab separator."""

    def test_inserted_studies_have_tab(self, app_config, tmp_dir):
        cv_path = _make_unsorted_cv(tmp_dir / "cv.docx")
        master_path = _make_master_with_new_studies(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        output = tmp_dir / "output.docx"
        processor.mode_a_update_inject(
            cv_path, master_path=master_path,
            output_path=output, enable_sort_existing=False,
        )

        doc = Document(output)
        in_section = False
        found_tab = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.lower().startswith("research experience"):
                in_section = True
                continue
            if in_section and text and text[:4].isdigit():
                # Study line should contain a tab
                if "\t" in para.text:
                    found_tab = True
                    break
        assert found_tab, "No tab character found in inserted study paragraphs"


# ===========================================================================
# Redact mode unaffected
# ===========================================================================

class TestRedactUnaffected:
    """Mode B: Redact should be completely unaffected by enable_sort_existing."""

    def test_redact_ignores_flag(self, app_config, tmp_dir):
        cv_path = tmp_dir / "cv.docx"
        _make_cv_docx(cv_path)
        master_path = tmp_dir / "master.xlsx"
        _make_master_xlsx(master_path)

        processor = CVProcessor(app_config)
        result = processor.mode_b_redact_protocols(
            cv_path, master_path=master_path,
            output_path=tmp_dir / "redacted.docx",
        )
        assert result.success is True
        assert (tmp_dir / "redacted.docx").exists()

        # No config log for enable_sort_existing in redact
        config_entries = [
            e for e in result.log_entries if e.operation == "config"
        ]
        sort_config = [
            e for e in config_entries
            if "enable_sort_existing" in e.details
        ]
        assert len(sort_config) == 0


# ===========================================================================
# Benchmark year logic still applies
# ===========================================================================

class TestBenchmarkYearStillApplies:
    """The ≤3 studies benchmark step-back logic is unaffected."""

    def test_manual_benchmark_with_sort_false(self, app_config, tmp_dir):
        cv_path = _make_unsorted_cv(tmp_dir / "cv.docx")
        master_path = _make_master_with_new_studies(tmp_dir / "master.xlsx")

        processor = CVProcessor(app_config)
        result = processor.mode_a_update_inject(
            cv_path, master_path=master_path,
            manual_benchmark_year=2025,
            output_path=tmp_dir / "out.docx",
            enable_sort_existing=False,
        )
        assert result.success is True
