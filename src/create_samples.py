"""
Create sample CV and master .xlsx files for testing.
Run this script to generate test files in the ./samples/ directory.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from openpyxl import Workbook


def create_sample_cv(output_path: Path):
    """Create a sample CV document with Research Experience section."""
    doc = Document()
    
    # Title
    title = doc.add_heading("John Doe, MD, PhD", level=0)
    
    # Contact
    doc.add_paragraph("Email: john.doe@example.com | Phone: (555) 123-4567")
    doc.add_paragraph("")
    
    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("MD, PhD - Harvard Medical School, 2015")
    doc.add_paragraph("BS Biology - MIT, 2008")
    doc.add_paragraph("")
    
    # Research Experience
    doc.add_heading("Research Experience", level=1)
    
    # Phase I
    doc.add_paragraph("Phase I")
    
    # Oncology subcategory
    doc.add_paragraph("Oncology")
    
    # Sample studies
    p1 = doc.add_paragraph()
    p1.add_run("2023\t")
    run = p1.add_run("Pfizer")
    run.bold = True
    p1.add_run(" ")
    run2 = p1.add_run("PF-12345")
    run2.bold = True
    p1.add_run(": A Phase 1 study of PF-12345 in patients with advanced solid tumors")
    
    p2 = doc.add_paragraph()
    p2.add_run("2022\t")
    run = p2.add_run("Novartis")
    run.bold = True
    p2.add_run(" ")
    run2 = p2.add_run("NVS-789")
    run2.bold = True
    p2.add_run(": First-in-human study of NVS-789 for metastatic breast cancer")
    
    # Cardiology subcategory
    doc.add_paragraph("Cardiology")
    
    p3 = doc.add_paragraph()
    p3.add_run("2021\t")
    run = p3.add_run("Merck")
    run.bold = True
    p3.add_run(" ")
    run2 = p3.add_run("MK-001")
    run2.bold = True
    p3.add_run(": Phase 1 dose escalation study of MK-001 in heart failure patients")
    
    # Phase II-IV
    doc.add_paragraph("Phase II–IV")
    
    # Oncology
    doc.add_paragraph("Oncology")
    
    p4 = doc.add_paragraph()
    p4.add_run("2023\t")
    run = p4.add_run("Roche")
    run.bold = True
    p4.add_run(" ")
    run2 = p4.add_run("RO-555")
    run2.bold = True
    p4.add_run(": Phase 3 randomized trial of RO-555 vs standard of care in NSCLC")
    
    p5 = doc.add_paragraph()
    p5.add_run("2020\t")
    run = p5.add_run("BMS")
    run.bold = True
    p5.add_run(" ")
    run2 = p5.add_run("BMS-222")
    run2.bold = True
    p5.add_run(": Phase 2 study of BMS-222 immunotherapy in melanoma")
    
    doc.add_paragraph("")
    
    # Publications
    doc.add_heading("Publications", level=1)
    doc.add_paragraph("1. Doe J, et al. Novel biomarkers in oncology. Nature Medicine. 2023.")
    doc.add_paragraph("2. Doe J, et al. Advances in immunotherapy. NEJM. 2022.")
    
    # Save
    doc.save(output_path)
    print(f"Created sample CV: {output_path}")


def create_sample_master(output_path: Path):
    """Create a sample master .xlsx file with studies."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Studies"
    
    # Headers (optional - will be skipped during parsing)
    ws['A1'] = "Hierarchy / Year"
    ws['B1'] = "Description Full"
    ws['C1'] = "Description Masked"
    
    data = [
        # Phase I
        ("Phase I", "", ""),
        ("Oncology", "", ""),
        (2024, "Pfizer PF-99999: A Phase 1 study of PF-99999 (pembrolizumab) in advanced lung cancer",
               "Pfizer: A Phase 1 study of XXX in advanced lung cancer"),
        (2024, "Novartis NVS-888: First-in-human study of NVS-888 (ribociclib) for breast cancer",
               "Novartis: First-in-human study of XXX for breast cancer"),
        (2023, "Pfizer PF-12345: A Phase 1 study of PF-12345 in patients with advanced solid tumors",
               "Pfizer: A Phase 1 study of XXX in patients with advanced solid tumors"),
        (2022, "Novartis NVS-789: First-in-human study of NVS-789 for metastatic breast cancer",
               "Novartis: First-in-human study of XXX for metastatic breast cancer"),
        ("Cardiology", "", ""),
        (2024, "AstraZeneca AZ-111: Phase 1 trial of AZ-111 (dapagliflozin) in heart failure",
               "AstraZeneca: Phase 1 trial of XXX in heart failure"),
        (2021, "Merck MK-001: Phase 1 dose escalation study of MK-001 in heart failure patients",
               "Merck: Phase 1 dose escalation study of XXX in heart failure patients"),
        ("Neurology", "", ""),
        (2024, "Biogen BIO-333: Phase 1 study of BIO-333 (lecanemab) in early Alzheimer's",
               "Biogen: Phase 1 study of XXX in early Alzheimer's"),
        # Phase II-IV
        ("Phase II–IV", "", ""),
        ("Oncology", "", ""),
        (2024, "Roche RO-777: Phase 3 study of RO-777 (atezolizumab) vs placebo in TNBC",
               "Roche: Phase 3 study of XXX vs placebo in TNBC"),
        (2023, "Roche RO-555: Phase 3 randomized trial of RO-555 vs standard of care in NSCLC",
               "Roche: Phase 3 randomized trial of XXX vs standard of care in NSCLC"),
        (2020, "BMS BMS-222: Phase 2 study of BMS-222 immunotherapy in melanoma",
               "BMS: Phase 2 study of XXX immunotherapy in melanoma"),
        ("Cardiology", "", ""),
        (2024, "Bayer BAY-444: Phase 2/3 trial of BAY-444 (rivaroxaban) in AF patients",
               "Bayer: Phase 2/3 trial of XXX in AF patients"),
    ]
    
    row = 2
    for item in data:
        ws.cell(row=row, column=1, value=item[0])
        if item[1]:
            ws.cell(row=row, column=2, value=item[1])
        if item[2]:
            ws.cell(row=row, column=3, value=item[2])
        row += 1
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 80
    ws.column_dimensions['C'].width = 70
    
    wb.save(output_path)
    print(f"Created sample master: {output_path}")


def main():
    """Create sample files."""
    # Create samples directory
    samples_dir = Path(__file__).parent / "samples"
    samples_dir.mkdir(exist_ok=True)
    
    # Create sample CV
    create_sample_cv(samples_dir / "sample_cv.docx")
    
    # Create sample master
    create_sample_master(samples_dir / "sample_master.xlsx")
    
    print("\nSample files created successfully!")
    print(f"Location: {samples_dir}")
    print("\nTo test the application:")
    print("  1. Run: python main.py")
    print("  2. Select the sample CV and master files")
    print("  3. Try Mode A (Update/Inject) to add 2024 studies")
    print("  4. Try Mode B (Redact) to mask protocols")


if __name__ == "__main__":
    main()
